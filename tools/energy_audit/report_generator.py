"""
能源审计报告生成工具
生成符合《能源审计报告编写格式规范标准》的 Word + Markdown 报告

格式规范（见 skills/productivity/energy-audit-imitate/references/report-format-spec.md）：
- 封面：36pt/14pt 宋体 加粗 居中
- 正文：12pt 宋体 + Times New Roman, 1.5倍行距, 首行缩进2字符, 两端对齐
- 标题：15pt 宋体 居中 / 14pt 宋体 左对齐 / 12pt 宋体 左对齐
- 目录：黑体 18pt 居中 + TOC 域；水印为被审计单位全称（DrawingML）
- 表格：12pt 宋体，行高 1.01cm
- 数字/单位：数字 + 空格 + 单位（如 1234 tce）
"""

from dataclasses import dataclass, field
from typing import Dict, List, Optional
from datetime import datetime
import json
import re, os

from tools.energy_audit.chart_utils import setup_chart_font, chart_text


# ============================================================
# 格式常量
# ============================================================

@dataclass
class FormatSpec:
    """Word 格式规范（基于19份省直报告统计）"""
    # 正文字体
    body_font_cn: str = "宋体"
    body_font_en: str = "Times New Roman"
    body_size: int = 12  # pt（小四号）
    body_line_spacing: float = 1.5
    body_first_line_indent: float = 2.0  # 字符
    body_alignment: int = 3  # 3=两端对齐(JUSTIFY)

    # 一级标题（第X章）— 小三号(15pt) 宋体 居中
    h1_font: str = "宋体"
    h1_size: int = 15  # 小三号
    h1_bold: bool = True
    h1_alignment: int = 1  # 1=居中
    h1_line_spacing: float = 1.5
    h1_space_before: int = 24
    h1_space_after: int = 12

    # 二级标题（X.X）— 四号(14pt) 宋体
    h2_font: str = "宋体"
    h2_size: int = 14  # 四号
    h2_bold: bool = True
    h2_alignment: int = 0  # 左对齐
    h2_line_spacing: float = 1.5
    h2_space_before: int = 12
    h2_space_after: int = 6

    # 三级标题（X.X.X）
    h3_font: str = "宋体"
    h3_size: int = 12
    h3_bold: bool = True
    h3_alignment: int = 0
    h3_line_spacing: float = 1.5

    # 封面
    cover_title_size: int = 36  # 被审计单位名称 + "能源审计报告"
    cover_info_size: int = 14   # 审计机构 + 日期

    # 表格
    table_header_size: int = 12      # 小四号
    table_content_size: int = 12     # 小四号
    table_line_spacing: float = 1.0

    # 计量单位格式
    unit_separator: str = " "  # 数字 + 空格 + 单位


FMT = FormatSpec()


# ============================================================
# 章节结构定义
# ============================================================

CHAPTER_STRUCTURES = {
    "公共机构": [
        ("第1章", "能源审计执行概要", [
            ("1.1", "审计目的"),
            ("1.2", "审计范围"),
            ("1.3", "审计周期"),
            ("1.4", "审计内容"),
            ("1.5", "审计过程"),
            ("1.6", "审计依据"),
        ]),
        ("第2章", "公共机构概况", [
            ("2.1", "公共机构基本情况"),
            ("2.2", "建筑物概况"),
            ("2.3", "能源资源利用情况"),
        ]),
        ("第3章", "能源资源管理状况", [
            ("3.1", "能源资源管理机构职责"),
            ("3.2", "能源资源管理目标和方针"),
            ("3.3", "能源资源管理成效与问题"),
        ]),
        ("第4章", "能源资源计量及统计状况", [
            ("4.1", "能源资源计量体系"),
            ("4.2", "计量器具配备及管理"),
            ("4.3", "能源资源统计情况"),
            ("4.4", "能源资源统计成效及问题"),
        ]),
        ("第5章", "能源资源消费/消耗指标分析", [
            ("5.1", "能源资源消费/消耗概况"),
            ("5.2", "能源资源消耗/消费数据"),
            ("5.3", "能耗资源消耗/消费指标"),
        ]),
        ("第6章", "主要能源资源利用系统分析", [
            ("6.1", "用电系统运行分析"),
            ("6.2", "用水系统运行分析"),
            ("6.3", "供暖系统运行分析"),
            ("6.4", "其他用能系统运行分析"),
            ("6.5", "室内环境检测"),
        ]),
        ("第7章", "节能效果与节能潜力分析", [
            ("7.1", "用能系统现状分析"),
            ("7.2", "节能潜力分析及建议"),
        ]),
        ("第8章", "审计结论", [
            ("8.1", "审计主要发现"),
            ("8.2", "能源利用评价"),
            ("8.3", "节能建议"),
            ("8.4", "后续工作建议"),
        ]),
    ],
    "公共建筑": [
        ("第1章", "审计执行概要", [("1.1", "审计目的"), ("1.2", "审计范围"), ("1.3", "审计依据"), ("1.4", "审计方法"), ("1.5", "主要结论")]),
        ("第2章", "建筑概况", [("2.1", "基本信息"), ("2.2", "建筑信息"), ("2.3", "用能人数"), ("2.4", "主要用能设备")]),
        ("第3章", "能源管理状况", [("3.1", "能源管理组织架构"), ("3.2", "能源管理制度"), ("3.3", "能源管理人员"), ("3.4", "能源管理培训")]),
        ("第4章", "能源计量与统计", [("4.1", "能源计量器具配置"), ("4.2", "能源统计方法"), ("4.3", "能源数据记录"), ("4.4", "能源数据报送")]),
        ("第5章", "能源消耗分析", [("5.1", "能源消费结构"), ("5.2", "能源消耗总量"), ("5.3", "人均能耗指标"), ("5.4", "单位面积能耗指标"), ("5.5", "同比环比分析")]),
        ("第6章", "能源系统分析", [("6.1", "电力系统分析"), ("6.2", "供暖系统分析"), ("6.3", "空调系统分析"), ("6.4", "照明系统分析"), ("6.5", "其他用能系统分析")]),
        ("第7章", "节能潜力分析", [("7.1", "已实施节能措施效果"), ("7.2", "节能潜力分析"), ("7.3", "节能改造建议"), ("7.4", "预期节能效果")]),
        ("第8章", "审计结论与建议", [("8.1", "审计主要发现"), ("8.2", "能源利用评价"), ("8.3", "节能建议"), ("8.4", "后续工作建议")]),
    ],
    "工业企业": [
        ("第1章", "审计执行概要", [("1.1", "审计目的"), ("1.2", "审计范围"), ("1.3", "审计依据"), ("1.4", "审计方法"), ("1.5", "主要结论")]),
        ("第2章", "企业概况", [("2.1", "基本信息"), ("2.2", "建筑信息"), ("2.3", "用能人数"), ("2.4", "主要用能设备")]),
        ("第3章", "能源管理状况", [("3.1", "能源管理组织架构"), ("3.2", "能源管理制度"), ("3.3", "能源管理人员"), ("3.4", "能源管理培训")]),
        ("第4章", "能源计量与统计", [("4.1", "能源计量器具配置"), ("4.2", "能源统计方法"), ("4.3", "能源数据记录"), ("4.4", "能源数据报送")]),
        ("第5章", "能源消耗分析", [("5.1", "能源消费结构"), ("5.2", "能源消耗总量"), ("5.3", "人均能耗指标"), ("5.4", "单位面积能耗指标"), ("5.5", "同比环比分析")]),
        ("第6章", "主要用能系统分析", [("6.1", "电力系统分析"), ("6.2", "供暖系统分析"), ("6.3", "空调系统分析"), ("6.4", "照明系统分析"), ("6.5", "其他用能系统分析")]),
        ("第7章", "节能潜力分析", [("7.1", "已实施节能措施效果"), ("7.2", "节能潜力分析"), ("7.3", "节能改造建议"), ("7.4", "预期节能效果")]),
        ("第8章", "审计结论与建议", [("8.1", "审计主要发现"), ("8.2", "能源利用评价"), ("8.3", "节能建议"), ("8.4", "后续工作建议")]),
    ],
}


# ============================================================
# Word 报告生成器
# ============================================================

class WordReportBuilder:
    """生成符合格式规范的 Word (.docx) 能源审计报告"""

    def __init__(self, audit_type: str):
        self.audit_type = audit_type
        self.chapters = CHAPTER_STRUCTURES.get(audit_type, CHAPTER_STRUCTURES["公共机构"])
        self.report_data: Dict = {}
        self.doc = None
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        self.Document = Document
        self.Pt = Pt
        self.Cm = Cm
        self.WD_ALIGN = WD_ALIGN_PARAGRAPH
        self.RGBColor = RGBColor
        self.qn = qn

    def set_data(self, data: Dict):
        """设置报告数据"""
        self.report_data = data

    def _set_font(self, run, cn_font: str, en_font: str, size_pt: int, bold: bool = False):
        """设置 run 的字体属性（直接操作 XML 确保可靠）"""
        from lxml import etree as _etree
        from docx.oxml.ns import qn as _qn
        rPr = run._element.get_or_add_rPr()

        # 字号 (w:sz = half-points, w:szCs = complex script)
        sz = rPr.find(_qn('w:sz'))
        if sz is None:
            sz = _etree.SubElement(rPr, _qn('w:sz'))
        sz.set(_qn('w:val'), str(size_pt * 2))
        szCs = rPr.find(_qn('w:szCs'))
        if szCs is None:
            szCs = _etree.SubElement(rPr, _qn('w:szCs'))
        szCs.set(_qn('w:val'), str(size_pt * 2))

        # 加粗
        if bold:
            if rPr.find(_qn('w:b')) is None:
                _etree.SubElement(rPr, _qn('w:b'))

        # 字体
        rFonts = rPr.find(_qn('w:rFonts'))
        if rFonts is None:
            rFonts = _etree.SubElement(rPr, _qn('w:rFonts'))
        rFonts.set(_qn('w:ascii'), en_font)
        rFonts.set(_qn('w:hAnsi'), en_font)
        rFonts.set(_qn('w:eastAsia'), cn_font)

    def _set_paragraph_format(self, para, line_spacing: float, alignment: int,
                              first_line_indent_chars: float = 0,
                              space_before: int = 0, space_after: int = 0):
        """设置段落格式"""
        from docx.shared import Pt as _Pt
        pf = para.paragraph_format
        pf.line_spacing = line_spacing
        pf.alignment = alignment
        if first_line_indent_chars > 0:
            # 首行缩进 = 字号(pt) × 缩进字符数
            pf.first_line_indent = _Pt(FMT.body_size * first_line_indent_chars)
        if space_before > 0:
            pf.space_before = _Pt(space_before)
        if space_after > 0:
            pf.space_after = _Pt(space_after)

    def _add_body_text(self, text: str):
        """添加正文段落（12pt 宋体 + Times New Roman, 1.5倍行距, 首行缩进2字符, 两端对齐）"""
        if not text:
            return
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, FMT.body_line_spacing, FMT.body_alignment,
                                   FMT.body_first_line_indent)
        run = para.add_run(text)
        self._set_font(run, FMT.body_font_cn, FMT.body_font_en, FMT.body_size)

    def _add_formula(self, formula_text: str):
        """插入OMML公式（Word数学公式格式，分数线+斜体变量+下标）"""
        from lxml import etree
        from docx.oxml.ns import qn

        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, FMT.body_line_spacing, FMT.body_alignment,
                                   FMT.body_first_line_indent)

        M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        oMath = etree.SubElement(para._element, qn('m:oMath'), nsmap={'m': M})

        def _add_math_run(parent, text, is_var=False):
            r = etree.SubElement(parent, qn('m:r'))
            if is_var:
                rPr = etree.SubElement(r, qn('m:rPr'))
                sty = etree.SubElement(rPr, qn('m:sty'))
                sty.set(qn('m:val'), 'p')
            t = etree.SubElement(r, qn('m:t'))
            t.text = str(text)
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        def _add_sub(parent, base, sub):
            """添加带下标的变量：E_jrcn, E_gn, E_jt"""
            sSub = etree.SubElement(parent, qn('m:sSub'))
            e = etree.SubElement(sSub, qn('m:e'))
            _add_math_run(e, base, is_var=True)
            sub_e = etree.SubElement(sSub, qn('m:sub'))
            _add_math_run(sub_e, sub, is_var=True)

        # E_jrcn =
        _add_sub(oMath, 'E', 'jrcn')
        _add_math_run(oMath, ' = ')

        # 分数线 (E - E_gn - E_jt) / M
        frac = etree.SubElement(oMath, qn('m:f'))
        etree.SubElement(frac, qn('m:fPr'))
        num = etree.SubElement(frac, qn('m:num'))
        den = etree.SubElement(frac, qn('m:den'))

        _add_math_run(num, '(')
        _add_math_run(num, 'E', is_var=True)
        _add_math_run(num, ' - ')
        _add_sub(num, 'E', 'gn')
        _add_math_run(num, ' - ')
        _add_sub(num, 'E', 'jt')
        _add_math_run(num, ')')

        _add_math_run(den, 'M', is_var=True)

    def _add_simple_fraction_formula(self, result_var: str, num_var: str, den_var: str):
        """插入简单分数 OMML 公式：result_var = num_var / den_var（支持下标的变量名如 E_ja）"""
        from lxml import etree
        from docx.oxml.ns import qn

        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, FMT.body_line_spacing, FMT.body_alignment,
                                   FMT.body_first_line_indent)

        M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        oMath = etree.SubElement(para._element, qn('m:oMath'), nsmap={'m': M})

        def _add_math_run(parent, text, is_var=False):
            r = etree.SubElement(parent, qn('m:r'))
            if is_var:
                rPr = etree.SubElement(r, qn('m:rPr'))
                sty = etree.SubElement(rPr, qn('m:sty'))
                sty.set(qn('m:val'), 'p')
            t = etree.SubElement(r, qn('m:t'))
            t.text = str(text)
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        def _add_sub_var(parent, var_str):
            """解析 E_ja → E 带下标 ja"""
            parts = var_str.split('_', 1)
            if len(parts) == 2:
                sSub = etree.SubElement(parent, qn('m:sSub'))
                e = etree.SubElement(sSub, qn('m:e'))
                _add_math_run(e, parts[0], is_var=True)
                sub_e = etree.SubElement(sSub, qn('m:sub'))
                _add_math_run(sub_e, parts[1], is_var=True)
            else:
                _add_math_run(parent, var_str, is_var=True)

        # result_var =
        _add_sub_var(oMath, result_var)
        _add_math_run(oMath, ' = ')

        # 分数线 num_var / den_var
        frac = etree.SubElement(oMath, qn('m:f'))
        etree.SubElement(frac, qn('m:fPr'))
        num = etree.SubElement(frac, qn('m:num'))
        den = etree.SubElement(frac, qn('m:den'))

        _add_sub_var(num, num_var)
        _add_sub_var(den, den_var)

    def _add_formula_symbol(self, symbol: str, description: str):
        """插入公式符号说明行：斜体变量（含下标）+ 正体中文描述"""
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, FMT.body_line_spacing, FMT.body_alignment,
                                   FMT.body_first_line_indent)

        # 解析 symbol 中的下划线：E_jrcn——  → E(italic) + subscript jrcn + ——description
        import re
        parts = re.split(r'(——)', symbol, 1)
        var_part = parts[0]  # e.g. "E_jrcn" or "M"
        sep = parts[1] if len(parts) > 1 else ''

        # 解析变量名中的下标
        var_match = re.match(r'([A-Za-z]+)_([a-z]+)', var_part)
        if var_match:
            # 有下标
            base = var_match.group(1)
            sub = var_match.group(2)
            run_base = para.add_run(base)
            self._set_font(run_base, FMT.body_font_cn, FMT.body_font_en, FMT.body_size)
            run_base.italic = True
            run_sub = para.add_run(sub)
            self._set_font(run_sub, FMT.body_font_cn, FMT.body_font_en, FMT.body_size)
            run_sub.italic = True
            run_sub.font.subscript = True
        else:
            # 无下标
            run_var = para.add_run(var_part)
            self._set_font(run_var, FMT.body_font_cn, FMT.body_font_en, FMT.body_size)
            run_var.italic = True

        if sep:
            run_sep = para.add_run(sep)
            self._set_font(run_sep, FMT.body_font_cn, FMT.body_font_en, FMT.body_size)

        run_desc = para.add_run(description)
        self._set_font(run_desc, FMT.body_font_cn, FMT.body_font_en, FMT.body_size)

    def _add_bullet_text(self, text: str):
        """添加无序列表项：实心圆点 + 12pt 宋体，1.5倍行距"""
        if not text:
            return
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, FMT.body_line_spacing, FMT.body_alignment)
        run = para.add_run("● " + text)
        self._set_font(run, FMT.body_font_cn, FMT.body_font_en, FMT.body_size)

    def _set_outline_level(self, para, level: int):
        """Word 大纲级别（0=章），供目录域收录。"""
        from lxml import etree
        from docx.oxml.ns import qn as _qn
        pPr = para._p.get_or_add_pPr()
        el = pPr.find(_qn("w:outlineLvl"))
        if el is None:
            el = etree.SubElement(pPr, _qn("w:outlineLvl"))
        el.set(_qn("w:val"), str(level))

    def _add_heading_1(self, text: str):
        """一级标题：小三号(15pt) 宋体 加粗 居中。"""
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, FMT.h1_line_spacing, FMT.h1_alignment,
                                   space_before=FMT.h1_space_before, space_after=FMT.h1_space_after)
        run = para.add_run(text)
        self._set_font(run, FMT.h1_font, FMT.body_font_en, FMT.h1_size, bold=FMT.h1_bold)
        self._set_outline_level(para, 0)

    def _add_heading_2(self, text: str):
        """二级标题：四号(14pt) 宋体 加粗 左对齐。"""
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, FMT.h2_line_spacing, FMT.h2_alignment,
                                   space_before=FMT.h2_space_before, space_after=FMT.h2_space_after)
        run = para.add_run(text)
        self._set_font(run, FMT.h2_font, FMT.body_font_en, FMT.h2_size, bold=FMT.h2_bold)
        self._set_outline_level(para, 1)

    def _add_heading_3(self, text: str):
        """三级标题：12pt 宋体 加粗 左对齐。"""
        para = self.doc.add_paragraph()
        self._set_paragraph_format(para, FMT.h3_line_spacing, FMT.h3_alignment)
        run = para.add_run(text)
        self._set_font(run, FMT.h3_font, FMT.body_font_en, FMT.h3_size, bold=FMT.h3_bold)
        self._set_outline_level(para, 2)

    def _add_table(self, headers: List[str], rows: List[List[str]], title: str = None):
        """添加表格：标题12pt 宋体 加粗 居中，内容12pt 宋体 居中，行高 1.01cm，上下居中"""
        if title:
            para = self.doc.add_paragraph()
            para.alignment = self.WD_ALIGN.CENTER
            run = para.add_run(title)
            self._set_font(run, FMT.body_font_cn, FMT.body_font_en, FMT.table_header_size, bold=True)

        from docx.shared import Cm as _Cm
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 行高 1.01cm
        for row in table.rows:
            row.height = _Cm(1.01)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # Header
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.clear()
            p.alignment = self.WD_ALIGN.CENTER
            r = p.add_run(h)
            self._set_font(r, FMT.body_font_cn, FMT.body_font_en, FMT.table_header_size, bold=True)
        # Data rows
        for ri, row_data in enumerate(rows):
            row = table.rows[ri + 1]
            for ci, val in enumerate(row_data):
                cell = row.cells[ci]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.clear()
                p.alignment = self.WD_ALIGN.CENTER
                r = p.add_run(str(val))
                self._set_font(r, FMT.body_font_cn, FMT.body_font_en, FMT.table_content_size)

    def _add_table_title(self, text: str):
        """表格标题：宋体 小四号(12pt) 加粗 居中"""
        para = self.doc.add_paragraph()
        para.alignment = self.WD_ALIGN.CENTER
        run = para.add_run(text)
        self._set_font(run, FMT.body_font_cn, FMT.body_font_en, FMT.table_header_size, bold=True)

    def _add_page_break(self):
        from docx.oxml.ns import qn
        self.doc.add_page_break()

    # ============================================================
    # 各章节生成
    # ============================================================

    def build_audit_info_tables(self):
        """审计基本信息三张表：机构信息表、审计组人员名单、配合人员名单（10.5pt 宋体 单倍行距）"""
        tables_data = self.report_data.get('audit_info_tables', {})

        # ---- 表1：能源审计机构信息表 ----
        self._add_table_title("能源审计机构信息表")
        inst = tables_data.get('institution', {})
        self._add_two_col_table(
            ["项  目", "内  容"],
            [
                ["机构名称", inst.get('name', '⚠️ 请提供被审计单位全称。')],
                ["地  址", inst.get('address', '⚠️ 请提供单位详细地址。')],
                ["负 责 人", inst.get('contact', '⚠️ 请提供负责人姓名。')],
                ["联系方式", inst.get('phone', '⚠️ 请提供联系电话。')],
            ]
        )

        # ---- 表2：能源审计组人员名单 ----
        self._add_table_title("能源审计组人员名单")
        members = tables_data.get('team_members', [])
        headers = ["组内职务", "姓  名", "学  历", "所获资质", "专  业"]
        rows = [[m.get(k, '') for k in ['role','name','education','certification','major']] for m in members] if members else [
            ["⚠️ 请提供审计组人员名单（姓名/学历/资质/专业）","","","",""]
        ]
        self._add_table(headers, rows)

        # ---- 表3：能源审计配合人员名单 ----
        self._add_table_title("能源审计配合人员名单")
        coop = tables_data.get('cooperation', [])
        headers = ["组内职务", "部  门", "姓  名", "性  别", "职  务"]
        rows = [[c.get(k, '') for k in ['role','dept','name','gender','position']] for c in coop] if coop else [
            ["⚠️ 请提供配合人员名单（部门/姓名/性别/职务）","","","",""]
        ]
        self._add_table(headers, rows)

    def _add_two_col_table(self, headers, rows):
        """添加两列表格（标题+内容格式），标题列加粗居中，内容列居中，行高 1.01cm，上下居中"""
        from docx.shared import Cm as _Cm
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        table = self.doc.add_table(rows=len(rows), cols=2, style='Table Grid')
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 行高 1.01cm
        for row in table.rows:
            row.height = _Cm(1.01)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for ri, row_data in enumerate(rows):
            row = table.rows[ri]
            for ci, val in enumerate(row_data):
                cell = row.cells[ci]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.clear()
                p.alignment = self.WD_ALIGN.CENTER
                r = p.add_run(str(val))
                if ci == 0:
                    self._set_font(r, FMT.body_font_cn, FMT.body_font_en, FMT.table_content_size, bold=True)
                else:
                    self._set_font(r, FMT.body_font_cn, FMT.body_font_en, FMT.table_content_size)
        self.doc.add_paragraph()

    def build_cover(self):
        """封面"""
        cover = self.report_data.get('cover', {})
        unit_name = cover.get('title', '')
        # 提取被审计单位名称（去掉"能源审计报告"后缀）
        if '能源审计报告' in unit_name:
            unit_name = unit_name.replace('能源审计报告', '').strip()

        # 顶部留白
        for _ in range(2):
            self.doc.add_paragraph()
        para = self.doc.add_paragraph()
        para.alignment = self.WD_ALIGN.CENTER
        run = para.add_run(unit_name)
        self._set_font(run, FMT.body_font_cn, FMT.body_font_en, FMT.cover_title_size, bold=True)

        # "能源审计报告"：36pt 宋体 加粗 居中
        para = self.doc.add_paragraph()
        para.alignment = self.WD_ALIGN.CENTER
        run = para.add_run("能源审计报告")
        self._set_font(run, FMT.body_font_cn, FMT.body_font_en, FMT.cover_title_size, bold=True)

        # 底部信息：通过审计机构的段前间距推到页面下部
        audit_org = cover.get('audit_organization', '')
        para = self.doc.add_paragraph()
        para.alignment = self.WD_ALIGN.CENTER
        para.paragraph_format.space_before = self.Pt(420)  # 推到页面下部
        run = para.add_run(audit_org)
        self._set_font(run, FMT.body_font_cn, FMT.body_font_en, FMT.cover_info_size, bold=True)

        self.doc.add_paragraph()

        # 日期：14pt 宋体 加粗 居中（YYYY年M月）
        date_str = cover.get('report_date', datetime.now().strftime('%Y年%m月'))
        # 转换为 YYYY年M月 格式
        try:
            dt = datetime.strptime(date_str, '%Y年%m月%d日')
            date_str = dt.strftime('%Y年%m月').replace('年0', '年').replace('月0', '月')
        except Exception:
            pass
        para = self.doc.add_paragraph()
        para.alignment = self.WD_ALIGN.CENTER
        run = para.add_run(date_str)
        self._set_font(run, FMT.body_font_cn, FMT.body_font_en, FMT.cover_info_size, bold=True)

        self._add_page_break()

    def build_toc(self):
        """目录：黑体 小二号(18pt) 加粗 居中 + Word TOC 域。"""
        from lxml import etree
        from docx.oxml.ns import qn as _qn

        para = self.doc.add_paragraph()
        para.alignment = self.WD_ALIGN.CENTER
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run("目  录")
        self._set_font(run, "黑体", FMT.body_font_en, 18, bold=True)

        field = self.doc.add_paragraph()
        field.paragraph_format.line_spacing = 1.5
        r = field.add_run()._r
        begin = etree.SubElement(r, _qn("w:fldChar"))
        begin.set(_qn("w:fldCharType"), "begin")
        instr_run = field.add_run()._r
        instr = etree.SubElement(instr_run, _qn("w:instrText"))
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instr.text = ' TOC \\o "1-2" \\h \\z \\u '
        sep_run = field.add_run()._r
        separate = etree.SubElement(sep_run, _qn("w:fldChar"))
        separate.set(_qn("w:fldCharType"), "separate")
        hint = field.add_run("（目录将在 Word/WPS 打开时自动生成；若未显示请按 Ctrl+A 后 F9 更新域）")
        self._set_font(hint, FMT.body_font_cn, FMT.body_font_en, FMT.body_size)
        end_run = field.add_run()._r
        end = etree.SubElement(end_run, _qn("w:fldChar"))
        end.set(_qn("w:fldCharType"), "end")

        self._add_page_break()

    def build_all_chapters(self):
        """生成第1章（模板）及其余章节（通用内容填充）。

        仿写模式（report_data['imitated_chapters']）按 CHAPTER_STRUCTURES 标题写章，
        有仿写正文的章节不再走公共机构硬编码模板。缺章时回退原 build_chapterN。
        """
        if self.report_data.get("generation_mode") == "imitate" or self.report_data.get("imitated_chapters"):
            self._build_imitated_chapters()
            return
        self.build_chapter1()
        self.build_chapter2()
        self.build_chapter3()
        self.build_chapter4()
        self.build_chapter5()
        self.build_chapter6()
        self.build_chapter7()
        self.build_chapter8()
        sections = self.report_data.get('sections', {})
        chapters_to_skip = {'第1章', '第2章', '第3章', '第4章', '第5章', '第6章', '第7章', '第8章'}
        for ch_num, ch_title, sub_sections in self.chapters:
            if ch_num in chapters_to_skip:
                continue
            self._add_heading_1(f"{ch_num}  {ch_title}")
            for sec_num, sec_title in sub_sections:
                self._add_heading_2(f"{sec_num}  {sec_title}")
                content = sections.get(sec_num, sections.get(ch_title, {}))
                if isinstance(content, dict):
                    content = content.get(sec_title, '')
                if isinstance(content, list):
                    for item in content:
                        self._add_body_text(str(item))
                elif content:
                    self._add_body_text(str(content))
                else:
                    self._add_body_text(f"[{sec_title}内容待补充]")

    _IMITATED_H1 = re.compile(r"^第[1-8一二三四五六七八]章")
    _IMITATED_H3 = re.compile(r"^\d+\.\d+\.\d+\s+\S")
    _IMITATED_H2 = re.compile(r"^\d+\.\d+\s+\S")

    def _imitated_chapter_text(self, chapter_key: str) -> str:
        raw = (self.report_data.get("imitated_chapters") or {}).get(chapter_key)
        if isinstance(raw, dict):
            return (raw.get("text") or "").strip()
        return (raw or "").strip()

    def _write_imitated_body(self, text: str):
        """把仿写正文写入 Word：识别 1.1 / 1.1.1 为标题，跳过重复的章标题。

        支持 markdown 表格（`| a | b |` 行，首行为表头，分隔行 `|---|` 忽略），
        以及表格标题行（以"表X.Y"开头的行）——渲染为规范表格（12pt 居中、行高 1.01cm）。
        支持图表标记行 `[[图:类型|图注]]`：按 chart_data 渲染并嵌入图片。
        """
        lines = (text or "").replace("\r\n", "\n").split("\n")
        pending_title = None  # 待绑定的表格标题
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                pending_title = None
                i += 1
                continue
            # 图表标记：[[图:类型|图注]] —— 渲染图表并嵌入
            if line.startswith("[[图:") and line.endswith("]]"):
                img_path = self._render_imitated_chart(line)
                if img_path:
                    caption = line[line.find("|")+1:-2].strip() if "|" in line else ""
                    self._add_image_with_caption(img_path, caption)
                pending_title = None
                i += 1
                continue
            # 表格标题行（表X.Y 开头）——预留给后续表格块
            if re.match(r"^表\d+\.\d+", line) and len(line) < 80:
                # 若该行后紧跟表格块则作为表格标题，否则作为正文（如"表5.1 所示"）
                j = i + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and lines[j].strip().startswith("|"):
                    pending_title = line
                    i += 1
                    continue
            # markdown 表格块：连续 | 行
            if line.startswith("|") and line.endswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.strip("|").split("|")]
                    if all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                        continue  # markdown 分隔行
                    rows.append(cells)
                if rows:
                    self._add_table(rows[0], rows[1:], title=pending_title)
                    pending_title = None
                continue
            if self._IMITATED_H1.match(line) and len(line) < 80:
                pending_title = None
                i += 1
                continue
            if self._IMITATED_H3.match(line) and len(line) < 80:
                pending_title = None
                self._add_heading_3(line)
            elif self._IMITATED_H2.match(line) and len(line) < 80:
                pending_title = None
                self._add_heading_2(line)
            else:
                pending_title = None
                self._add_body_text(line)
            i += 1

    def _render_imitated_chart(self, marker: str):
        """渲染仿写正文中的图表标记 [[图:类型|图注]]。

        类型支持：trend（逐年趋势）、pie（能源结构饼图）、
        monthly_电/monthly_水/monthly_气（逐月柱状图）、flow（能源流向图）。
        数据来自 report_data['chart_data']，缺数据或渲染失败时返回 None（静默跳过）。
        """
        inner = marker[4:-2].strip()
        chart_type = inner.split("|")[0].strip().lower()
        chart_data = self.report_data.get("chart_data") or {}
        years_data = chart_data.get("years") or []
        if not years_data:
            return None
        output_dir = chart_data.get("output_dir") or os.path.join(os.getcwd(), "charts")
        os.makedirs(output_dir, exist_ok=True)
        try:
            if chart_type in ("trend", "pie"):
                from tools.energy_audit.indicators import YearlyEnergyData
                yd_objects = []
                for d in years_data:
                    yd_objects.append(YearlyEnergyData(
                        year=int(d.get("year", 0)),
                        electricity_kwh=float(d.get("electricity_kwh", 0) or 0),
                        water_m3=float(d.get("water_m3", 0) or 0),
                        natural_gas_m3=float(d.get("natural_gas_m3", 0) or 0),
                        heating_energy_heat=float(d.get("heating_energy_heat_gj", 0) or 0),
                        transportation_petrol_kg=float(d.get("petrol_kg", 0) or 0),
                        transportation_diesel_kg=float(d.get("diesel_kg", 0) or 0),
                        building_area=float(chart_data.get("building_area", 0) or 0),
                        people_count=float(chart_data.get("people_count", 0) or 0),
                    ))
                energy_types = chart_data.get("energy_types") or ["electricity_kwh", "water_m3", "natural_gas_m3"]
                if chart_type == "trend":
                    return self._generate_yearly_trend_chart(yd_objects, energy_types, output_dir)
                return self._generate_energy_pie_chart(yd_objects, energy_types, output_dir)
            if chart_type == "cost_pie":
                return self._generate_cost_pie_chart(years_data, output_dir)
            if chart_type.startswith("monthly_"):
                et_key = chart_type.replace("monthly_", "")
                monthly_attr = {"electricity_kwh": "monthly_electricity_kwh",
                                "water_m3": "monthly_water_m3",
                                "natural_gas_m3": "monthly_natural_gas_m3"}.get(et_key)
                if not monthly_attr:
                    return None
                name_cn = self._ENERGY_TYPE_CN.get(et_key, et_key)
                unit = {"electricity_kwh": "kWh", "water_m3": "m³", "natural_gas_m3": "m³"}.get(et_key, "")
                return _generate_monthly_bar_chart(years_data, et_key, monthly_attr, name_cn, unit, output_dir)
            if chart_type == "flow":
                from tools.energy_audit.energy_flow_chart import draw_energy_flow_diagram
                energy_types = chart_data.get("energy_types") or ["electricity_kwh", "water_m3", "natural_gas_m3"]
                # 时间戳文件名：避免与用户已打开的旧图（文件锁）冲突
                flow_path = os.path.join(output_dir, f"energy_flow_{int(__import__('time').time())}.png")
                return draw_energy_flow_diagram(
                    energy_types=energy_types,
                    equipment=chart_data.get("equipment"),
                    unit_name=chart_data.get("unit_name", ""),
                    output_path=flow_path,
                ) or (flow_path if os.path.exists(flow_path) else None)
        except Exception as e:
            self._chart_errors = getattr(self, "_chart_errors", []) + [f"{chart_type}: {e}"]
            return None
        return None

    def _append_chapter_images(self, chapter_key: str):
        mapping = {"第2章": "chapter2", "第3章": "chapter3"}
        data_key = mapping.get(chapter_key)
        if not data_key:
            return
        images = self.report_data.get(data_key, {}).get("images", []) or []
        prefix = "图2" if chapter_key == "第2章" else "图3"
        for idx, img_info in enumerate(images, 1):
            path = img_info.get("path", "") if isinstance(img_info, dict) else ""
            caption = (img_info.get("caption") if isinstance(img_info, dict) else None) or f"{prefix}-{idx}"
            if path and os.path.exists(path):
                self._add_image_with_caption(path, caption)

    def _build_imitated_chapters(self):
        builders = {
            "第1章": self.build_chapter1,
            "第2章": self.build_chapter2,
            "第3章": self.build_chapter3,
            "第4章": self.build_chapter4,
            "第5章": self.build_chapter5,
            "第6章": self.build_chapter6,
            "第7章": self.build_chapter7,
            "第8章": self.build_chapter8,
        }
        for ch_num, ch_title, _sub in self.chapters:
            text = self._imitated_chapter_text(ch_num)
            if text:
                self._add_heading_1(f"{ch_num}  {ch_title}")
                self._write_imitated_body(text)
                self._append_chapter_images(ch_num)
            else:
                builder = builders.get(ch_num)
                if builder:
                    builder()

    # ============================================================
    # 第1章模板
    # ============================================================

    # 1.6 审计依据 — 固定部分（国家法律 + 部委规章 + 国家标准）
    AUDIT_BASIS_NATIONAL = [
        "《中华人民共和国节约能源法》（2016年）；",
        "《公共机构节能条例》（2008年）；",
        "《公共机构能源审计管理暂行办法》（国家发展改革委、国管局令第32号）；",
        "《公共建筑能源审计技术导则》（2016年）；",
        "《公共机构能源审计技术导则》（GB/T 31342-2014）；",
        "《建筑节能与可再生能源利用通用规范》（GB 55015-2021）；",
        "《公共机构能源管理体系实施指南》（GB/T 32019-2015）；",
        "《民用建筑电气设计标准》（GB 51348-2019）；",
        "《空调通风系统运行管理规程》（GB 50365-2019）；",
        "《民用建筑供暖通风与空气调节设计规范》（GB 50736-2016）；",
        "《建筑给水排水设计标准》（GB 50015-2019）；",
        "《供配电系统设计规范》（GB 50052-2009）；",
        "《综合能耗计算通则》（GB/T 2589-2020）；",
        "《公共机构能源资源计量器具配备和管理要求》（GB/T29149-2012）；",
        "《建筑照明设计标准》（GB/T 50034-2024）；",
        "《热泵和冷水机组能效限定值及能效等级》（GB 19577-2024）；",
        "《室内空气质量标准》（GB/T 18883-2022）；",
        "《公共建筑节能改造技术规范》（JGJ 176-2009）；",
        "《城镇供热管网设计规范》（CJJ/T 34-2022）；",
    ]

    @staticmethod
    def _water_indicator_name(institution_category: str = '') -> str:
        """取水指标名按机构类型自适应（1.2 审计范围与 5.3.4 标题共用，2026-09-03 用户确认）。

        医疗 → 单位开放床日用水量；政务服务中心/场馆（含体育）→ 单位建筑面积年取水量；
        教育 → 人均用水量；党政机关（默认）→ 人均机关取水量。
        """
        ic = institution_category or ''
        if '医疗' in ic:
            return '单位开放床日用水量'
        if '政务' in ic or '场馆' in ic or '体育' in ic:
            return '单位建筑面积年取水量'
        if '教育' in ic:
            return '人均用水量'
        return '人均机关取水量'

    @staticmethod
    def _venue_sub_type(specific_type: str = '', institution_category: str = '') -> str:
        """场馆子类型识别（DB37/T 3780-2019 定额按子类型分档）。

        文化馆（宫）、美术馆等其他文化类场馆参照图书馆（标准注4）。
        """
        text = f"{specific_type or ''} {institution_category or ''}"
        for k in ('图书馆', '博物馆', '剧院', '体育馆', '科技馆'):
            if k in text:
                return k
        return '图书馆'

    def build_chapter1(self):
        """第1章 能源审计执行概要（模板填充）"""
        d = self.report_data
        ch1 = d.get('chapter1', {})
        unit = ch1.get('audited_unit', d.get('cover', {}).get('title', '【被审计单位】'))
        unit_short = ch1.get('audited_unit_short', unit.replace('能源审计报告', '').strip())
        auditor = ch1.get('audit_org', d.get('cover', {}).get('audit_organization', '【审计单位】'))
        province = ch1.get('province', '山东')

        self._add_heading_1("第1章  能源审计执行概要")

        # ---- 1.1 审计目的 ----
        self._add_heading_2("1.1  审计目的")
        self._add_body_text(
            "公共机构能源审计是指依据有关法律、法规、规章和标准，对公共机构的用能系统、设备的运行、管理"
            "及能源资源利用状况进行检验、核查和技术、经济分析评价，提出改进用能方式或提高用能效率建议和意见的行为。"
        )
        self._add_body_text(
            f"为了准确了解{unit_short}的用能情况，{province}省机关事务管理局"
            f"委托{auditor}对{unit_short}进行能源审计。"
        )

        # ---- 1.2 审计范围 ----
        self._add_heading_2("1.2  审计范围")
        addr = ch1.get('address', '【地址】')
        bldg = ch1.get('buildings', '【建筑描述】')
        # 年份：优先从 audit_period 解析，否则用 year_start/year_end
        cyc = ch1.get('audit_period', '')
        if cyc and '—' in cyc:
            parts = cyc.split('—')
            y1 = parts[0][:4] if len(parts[0]) >= 4 else ch1.get('audit_start', '')[:4]
            y2 = parts[1][:4] if len(parts) > 1 and len(parts[1]) >= 4 else ch1.get('audit_end', '')[:4]
        else:
            y1 = ch1.get('audit_start', '')[:4] or ch1.get('year_start', '【起始年】')
            y2 = ch1.get('audit_end', '')[:4] or ch1.get('year_end', '【结束年】')
        # 机构类别 → 取水指标名（与 5.3.4 动态一致）
        ic12 = d.get('tags', {}).get('institution_category', '')
        water_name12 = self._water_indicator_name(ic12)
        # 能源类型列表 → 中文
        et_raw = ch1.get('energy_types', [])
        if isinstance(et_raw, list):
            et_map = {'electricity_kwh':'电','water_m3':'水','natural_gas_m3':'天然气',
                      'petrol_kg':'汽油','diesel_kg':'柴油','heating_energy_heat_gj':'热'}
            energy_types = '、'.join(et_map.get(e, e) for e in et_raw if e in et_map) or '电、水、天然气'
        else:
            energy_types = str(et_raw) or '电、热、水、油'
        if not y1 or y1 == '【起始年】': y1 = ''
        if not y2 or y2 == '【结束年】': y2 = ''
        year_str = f"{y1}年-{y2}年" if y1 and y2 else f"{bldg}年度"
        self._add_body_text(
            f"本次能源审计范围包含位于{addr}的{unit_short}内{bldg}。"
        )
        self._add_body_text(
            f"本次能源审计工作基于{year_str}整年及每月的{energy_types}等能源账单及能耗统计数据，"
            f"同时结合现场勘察所收集的各建筑围护结构、各类用能设备资料、日常用能习惯等实际数据信息，"
            f"对{unit_short}年总能耗、单位建筑面积非供暖能耗、常规用能系统单位建筑面积电耗、"
            f"人均综合能耗、{water_name12}等进行分析计算，依据国家或地方能耗定额标准，"
            f"对建筑用能现状进行总体评价。"
        )

        # ---- 1.3 审计周期 ----
        self._add_heading_2("1.3  审计周期")
        audit_time = ch1.get('audit_time', '【YYYY年M月—YYYY年M月】')
        audit_period = ch1.get('audit_period', '【YYYY年M月-YYYY年M月】')
        base_period = ch1.get('base_period', '【YYYY年M月-YYYY年M月】')
        self._add_body_text("审计时间")
        self._add_body_text(audit_time)
        self._add_body_text("审计期")
        self._add_body_text(audit_period)
        self._add_body_text("基准期")
        self._add_body_text(base_period)

        # ---- 1.4 审计内容 ----
        self._add_heading_2("1.4  审计内容")
        self._add_body_text(
            "依据国家有关的节能法规和标准，对公共机构能源资源利用状况进行检验、核查和分析评价，主要包括以下内容："
        )
        for item in ["1.能源资源管理情况；", "2.能耗分析评价；", "3.节能潜力分析及建议。"]:
            self._add_body_text(item)

        # ---- 1.5 审计过程 ----
        self._add_heading_2("1.5  审计过程")
        self._add_body_text(
            f"审计前期，审计组对{unit_short}发送能源审计调研表，根据反馈调研表梳理分析能耗数据、用能设备情况等，"
            f"围绕单位用能水平、节能潜力点、能源管理等方面进行分析并形成初步评估意见。"
        )
        self._add_body_text(
            "审计中期，审计组开展项目现场调研，根据前期评估内容进行现场沟通核实，"
            "确认前期资料与现场数据的一致性。主要调研包括能耗数据异常情况确认、"
            "现场用能设备与系统确认、能源管理情况确认等。同时现场重点查看节能管理责任和日常节能措施落实情况、"
            "用能系统运维管理情况、能源资源计量器具配备情况等。"
        )
        self._add_body_text(
            "审计后期，审计组对现场调研数据及各类资料进行整理汇总，包含但不限于各单位能耗数据、用能系统、异常数据等，"
            "分析各单位用能结构和用能规律，进一步全面分析各单位能源应用情况，形成并提交能源审计报告。"
        )

        # ---- 1.6 审计依据 ----
        self._add_heading_2("1.6  审计依据")
        from tools.energy_audit.province_regulations import get_provincial_regulations
        # 确定机构类型
        tags = self.report_data.get('tags', {})
        inst_map = {'党政机关': 'government', '医疗': 'medical', '教育': 'education'}
        inst_type = inst_map.get(tags.get('institution_category', ''), 'government')
        provincial = ch1.get('provincial_regulations') or get_provincial_regulations(province, inst_type)
        for line in provincial:
            self._add_bullet_text(line)
        # 国家标准（固定）
        for line in self.AUDIT_BASIS_NATIONAL:
            self._add_bullet_text(line)
        self._add_bullet_text("调研资料及其他相关资料。")

    # ============================================================
    # 第2章：LLM生成 + 图片嵌入
    # ============================================================

    def build_chapter2(self):
        """第2章 公共机构概况（数据驱动：优先从 project_data 自动生成，用户文本兜底）"""
        ch2 = self.report_data.get('chapter2', {})
        ch1 = self.report_data.get('chapter1', {})
        tags = self.report_data.get('tags', {})
        self._add_heading_1("第2章  公共机构概况")

        unit = ch1.get('audited_unit_short', '') or ch2.get('unit_name', '被审计单位')
        address = ch1.get('address', '') or ch2.get('address', '')
        buildings = ch2.get('buildings', [])
        area = ch2.get('building_area', 0)
        people = ch2.get('people_count', 0)
        cat = tags.get('institution_category', '')
        spec = tags.get('specific_type', '')

        # ---- 2.1 ----
        self._add_heading_2("2.1  公共机构基本情况")
        text_21 = ch2.get('section_2_1', '')
        if not text_21:
            # 2.1 直取 basic_situation（load_from_project 已填 + PG 兜底）；这里仍为空才提示补全
            text_21 = (f"⚠️ 2.1 单位基本情况缺失：请补充 basic_situation 字段"
                       f"（数据收集阶段或 PG ts_customer_info.get_customer_info 均无值）"
                       f"或由 author agent 根据被审计单位信息（{buildings[0].get('name','') if buildings else ''}等"
                       f"{len(buildings)}栋建筑，面积{area}m²，人数{people}人）人工补全。")
        for p in text_21.split('\n'):
            if p.strip():
                self._add_body_text(p.strip())

        # 图片
        images = ch2.get('images', [])
        for idx, img_info in enumerate(images, 1):
            img_path = img_info.get('path', '')
            caption = img_info.get('caption', f'图2-{idx}')
            if img_path and os.path.exists(img_path):
                self._add_image_with_caption(img_path, caption)

        # ---- 2.2 ----
        self._add_heading_2("2.2  建筑物概况")
        text_22 = ch2.get('section_2_2', '')
        if not text_22 and buildings:
            # 段1: 总览
            bldg_names = '、'.join(b.get('name','') for b in buildings[:6])
            text_22 = f"{unit}院内主要建筑物包括{bldg_names}等{len(buildings)}栋建筑。"
            # 共性特征（扩展：含屋面保温 / 监测 / 遮阳 / 楼层计量）
            total = len(buildings)

            def _rate(field, want):
                """统计满足 want 的建筑数，返回 '全部N栋' / 'N栋（X%）' / None"""
                cnt = sum(1 for b in buildings if b.get(field) == want)
                if cnt == 0:
                    return None
                if cnt == total:
                    return f"全部{total}栋"
                return f"{cnt}栋（{cnt/total*100:.0f}%）"

            structures = set(b.get('structure','') for b in buildings if b.get('structure'))
            insulations = set(b.get('insulation','') for b in buildings if b.get('insulation'))
            windows = set(b.get('window_type','') for b in buildings if b.get('window_type'))
            sunshades = set(b.get('sunshade_type','') for b in buildings if b.get('sunshade_type'))
            common_parts = []
            if structures:
                s = '、'.join(structures)
                common_parts.append(f"均采用{s}" if '结构' in s else f"均采用{s}结构")
            if insulations and '有' in str(insulations): common_parts.append("设有外墙保温")
            if windows and '—' not in str(windows) and '无' not in str(windows):
                common_parts.append(f"外窗采用{'、'.join(windows)}")
            # ---- 新增维度 ----
            roof = _rate('roof_insulation', '有')
            if roof: common_parts.append(f"{roof}建筑设有屋面保温")
            mon = _rate('monitoring', '有')
            if mon: common_parts.append(f"{mon}建筑配备能耗在线监测系统")
            sm = _rate('storey_metrology', '是')
            if sm: common_parts.append(f"{sm}建筑实现楼层单独计量")
            if sunshades:
                common_parts.append(f"遮阳形式为{'、'.join(sunshades)}")
            if common_parts:
                text_22 += f"各建筑{'，'.join(common_parts)}。\n\n"

            # 面积汇总：供冷 / 供热 / 地下车库
            total_cool = sum(float(b.get('cooling_area') or 0) for b in buildings)
            total_heat = sum(float(b.get('heating_area') or 0) for b in buildings)
            total_garage = sum(float(b.get('garage_area') or 0) for b in buildings)
            agg = []
            if total_cool > 0: agg.append(f"供冷面积{total_cool:g}m²")
            if total_heat > 0: agg.append(f"供热面积{total_heat:g}m²")
            if total_garage > 0: agg.append(f"地下车库面积{total_garage:g}m²")
            if agg:
                text_22 += f"\n\n全院合计：{'，'.join(agg)}。\n"

            # 段2: 逐栋详情
            text_22 += "各建筑物详细情况如下：\n"
            for b in buildings:
                parts = [f"{b.get('name','')}"]
                if b.get('year'): parts.append(f"{b.get('year')}年竣工")
                if b.get('floors'): parts.append(b.get('floors'))
                if b.get('area'): parts.append(f"建筑面积{b.get('area')}m²")
                if b.get('structure'):
                    st = b.get('structure')
                    parts.append(st if '结构' in st else f'{st}结构')
                # ---- 新增字段 (可选，避免空值噪音) ----
                if b.get('orientation'): parts.append(f"朝{b.get('orientation')}")
                cooling_area = b.get('cooling_area') or 0
                heating_area = b.get('heating_area') or 0
                if cooling_area and float(cooling_area) > 0:
                    parts.append(f"供冷面积{float(cooling_area):g}m²")
                if heating_area and float(heating_area) > 0:
                    parts.append(f"供热面积{float(heating_area):g}m²")
                if b.get('roof_insulation') == '有':
                    mat = b.get('roof_insulation_material', '')
                    parts.append(f"屋面保温（{mat}）" if mat else "设有屋面保温")
                if b.get('sunshade_type'):
                    parts.append(f"采用{b.get('sunshade_type')}")
                if b.get('run_time'):
                    parts.append(f"运行时间为{b.get('run_time')}")
                if b.get('monitoring') == '有':
                    parts.append("配备能耗在线监测系统")
                text_22 += '，'.join(parts) + "。\n"

            # 收口
            text_22 += f"各建筑详细参数见表2-1至表2-{len(buildings)}。"
        for p in text_22.split('\n'):
            if p.strip():
                self._add_body_text(p.strip())

        # 建筑参数表（自动）
        for table_num, bldg in enumerate(buildings, 1):
            self._add_building_param_table(bldg, table_num)

        # ---- 2.3 ----
        self._add_heading_2("2.3  能源资源利用情况")
        text_23 = ch2.get('section_2_3', '')
        if not text_23:
            et = ch1.get('energy_types', [])
            et_map = {'electricity_kwh':'电','water_m3':'水','natural_gas_m3':'天然气','petrol_kg':'汽油','diesel_kg':'柴油'}
            eq = self.report_data.get('chapter6', {}).get('_equipment', [])
            parts = []

            # 电系统
            if 'electricity_kwh' in et:
                cooling = [e for e in eq if e.get('category') == '空调']
                lighting = [e for e in eq if e.get('category') == '照明']
                office = [e for e in eq if e.get('category') == '办公']
                elec_text = f"{unit}用电系统主要包括"
                elec_items = []
                if cooling:
                    qty = sum(c.get('quantity',0) for c in cooling)
                    elec_items.append(f"{'、'.join(c.get('name','') for c in cooling[:3])}等共{qty}台")
                if lighting:
                    qty = sum(l.get('quantity',0) for l in lighting)
                    elec_items.append(f"{'、'.join(l.get('name','') for l in lighting[:3])}共{qty}套")
                if office:
                    ofc_parts = [f"{o.get('name','')}{o.get('quantity','')}台" if o.get('quantity') else o.get('name','') for o in office[:3]]
                    elec_items.append(f"{'、'.join(ofc_parts)}")
                elec_text += '、'.join(elec_items) + "。" if elec_items else "照明、空调、办公设备等。"
                parts.append(elec_text)

            # 水系统
            if 'water_m3' in et:
                parts.append(f"{unit}用水系统主要为生活用水、卫生清洁用水等，由市政自来水供水。")

            # 天然气
            if 'natural_gas_m3' in et:
                kitchen = [e for e in eq if e.get('category') == '厨房']
                if kitchen:
                    k_parts = [f"{k.get('name','')}" for k in kitchen[:3]]
                    parts.append(f"{unit}用气系统主要为厨房设备({', '.join(k_parts)})。")
                else:
                    parts.append(f"{unit}用气系统主要为厨房炊事用气。")

            # 汽/柴油
            if 'petrol_kg' in et or 'diesel_kg' in et:
                parts.append(f"{unit}用油系统主要为公务用车燃油消耗。")

            # 供暖（从能源数据推断）
            if any(ey.get('heating_energy_heat_gj', 0) or ey.get('heating_cost_wan', 0) for ey in (proj.energy_yearly if hasattr(self, 'proj') else [])):
                parts.append(f"{unit}供暖采用市政集中供热，按面积缴费。")

            text_23 = '\n\n'.join(parts) if parts else f"{unit}主要用能类型包括{'、'.join(et_map.get(e, e) for e in et if e in et_map)}。"
        for p in text_23.split('\n'):
            if p.strip():
                self._add_body_text(p.strip())

    # ============================================================
    # 第3章：LLM生成 + 图片嵌入
    # ============================================================

    def build_chapter3(self):
        """第3章 能源资源管理状况（从 management 数据自动生成）"""
        ch3 = self.report_data.get('chapter3', {})
        self._add_heading_1("第3章  能源资源管理状况")
        unit = (self.report_data.get('chapter1', {}).get('audited_unit_short', '被审计单位'))

        # 3.1
        self._add_heading_2("3.1  能源资源管理机构职责")
        t31 = ch3.get('section_3_1', '')
        if not t31:
            t31 = f"{unit}按照《公共机构节能条例》要求，设立了能源管理岗位和责任人，明确了节能管理职责，确保能源资源管理工作有人抓、有人管。"
        for p in t31.split('\n'): p=p.strip(); p and self._add_body_text(p)

        # 3.2
        self._add_heading_2("3.2  能源资源管理目标和方针")
        t32 = ch3.get('section_3_2', '')
        if not t32:
            t32 = f"{unit}坚持'节约优先、高效利用'的能源管理方针，将节能管理纳入日常运营，通过制度建设、全员参与、定期监督等方式实现能耗合理控制。"
        for p in t32.split('\n'): p=p.strip(); p and self._add_body_text(p)

        # 3.3
        self._add_heading_2("3.3  能源资源管理成效与问题")
        t33 = ch3.get('section_3_3', '')
        if not t33:
            t33 = f"{unit}在能源管理方面虽已建立基本的制度框架，但仍存在进一步改进的空间。"
        for p in t33.split('\n'): p=p.strip(); p and self._add_body_text(p)

        # 3.4 节能改造与管理措施（有节能管理信息记录时渲染，否则跳过）
        t34 = ch3.get('section_3_4', '')
        if t34:
            self._add_heading_2("3.4  节能改造与管理措施")
            for p in t34.split('\n'): p=p.strip(); p and self._add_body_text(p)

        # 图片
        for idx, img_info in enumerate(ch3.get('images', []), 1):
            p = img_info.get('path', '')
            if p and os.path.exists(p):
                self._add_image_with_caption(p, img_info.get('caption', f'图3-{idx}'))

    # ============================================================
    # 第4章：计量及统计状况
    # ============================================================

    _ENERGY_TYPE_CN = {
        'electricity_kwh': '电', 'water_m3': '水', 'natural_gas_m3': '天然气',
        'petrol_kg': '汽油', 'diesel_kg': '柴油', 'heating_energy_heat_gj': '热',
    }

    def build_chapter4(self):
        """第4章 能源资源计量及统计状况（从 metering 数据自动生成）"""
        ch4 = self.report_data.get('chapter4', {})
        self._add_heading_1("第4章  能源资源计量及统计状况")
        unit = (self.report_data.get('chapter1', {}).get('audited_unit_short', '被审计单位'))

        for key, title in [('section_4_1','4.1  能源资源计量体系'), ('section_4_2','4.2  计量器具配备及管理'),
                           ('section_4_3','4.3  能源资源统计情况'), ('section_4_4','4.4  能源资源统计成效及问题')]:
            self._add_heading_2(title)
            t = ch4.get(key, '')
            if not t:
                if '4.1' in key:
                    t = f"{unit}能源资源计量体系按GB/T29149-2012要求划分为一级（总表）、二级（分区）、三级（设备）三级计量体系。"
                elif '4.2' in key:
                    t = f"{unit}计量器具主要包括电表、水表、天然气表，由供电公司、自来水公司、燃气公司分别计量。"
                elif '4.3' in key:
                    t = f"{unit}能耗数据采集主要依靠人工抄表和账单汇总，按月统计核算。"
                else:
                    t = f"{unit}能源资源统计工作总体运行正常，但仍存在精细化管理的提升空间。"
            for p in t.split('\n'): p=p.strip(); p and self._add_body_text(p)

    # ============================================================
    # 第5章：能耗指标分析
    # ============================================================

    def build_chapter5(self):
        """第5章 能源资源消费/消耗指标分析（结构化表格 + 指标计算）"""
        ch5 = self.report_data.get('chapter5', {})
        self._add_heading_1("第5章  能源资源消费/消耗指标分析")

        # 从 chapter5 或 chapter1 获取必要参数
        unit_name = ch5.get('unit_name', self.report_data.get('chapter1', {}).get('audited_unit_short', '被审计单位'))
        years_raw = ch5.get('years', [])
        energy_data_list = ch5.get('energy_data', [])  # List[dict] with keys: year, electricity_kwh, water_m3, ...
        cost_data = ch5.get('cost_data', {})
        area = ch5.get('building_area', self.report_data.get('chapter2', {}).get('building_area', 0))
        people = ch5.get('people_count', self.report_data.get('chapter2', {}).get('people_count', 0))
        institution_type = ch5.get('institution_type', 'government')
        _tags5 = self.report_data.get('tags', {})
        venue_sub = self._venue_sub_type(_tags5.get('specific_type', ''), _tags5.get('institution_category', ''))

        if not years_raw:
            years_raw = sorted(set(str(d.get('year', '')) for d in energy_data_list if d.get('year')))
        if not energy_data_list:
            self.doc.add_page_break()
            para = self.doc.add_paragraph()
            para.alignment = 1  # CENTER
            run = para.add_run("⚠️ 第5章数据待补充")
            self._set_font(run, "宋体", "宋体", 14, bold=True)
            tip = self.doc.add_paragraph()
            tip.alignment = 1
            r2 = tip.add_run(
                "需要被审计单位的逐月/逐年电耗、水耗、天然气消耗等能源账单数据。\n"
                "数据来源：ts_institution_energy_main/data 表查询 → Excel/CSV 文件 → 用户手动输入。\n"
                "请提供上述数据后重新生成报告。"
            )
            self._set_font(r2, FMT.body_font_cn, FMT.body_font_en, 12)
            return

        # 用 indicators.py 计算指标
        from tools.energy_audit.indicators import (
            YearlyEnergyData, calc_unit_area_non_heating_energy, calc_unit_area_electricity,
            calc_per_capita_energy, calc_per_capita_water, resolve_benchmark,
        )
        yd_objects = []
        for d in energy_data_list:
            yd = YearlyEnergyData(
                year=int(d.get('year', 0)),
                electricity_kwh=float(d.get('electricity_kwh', 0) or 0),
                water_m3=float(d.get('water_m3', 0) or 0),
                natural_gas_m3=float(d.get('natural_gas_m3', 0) or 0),
                heating_energy_heat=float(d.get('heating_energy_heat_gj', 0) or 0),
                transportation_petrol_kg=float(d.get('petrol_kg', 0) or 0),
                transportation_diesel_kg=float(d.get('diesel_kg', 0) or 0),
                building_area=area, people_count=people,
                electricity_cost_wan=float(d.get('electricity_cost_wan', 0) or 0),
                water_cost_wan=float(d.get('water_cost_wan', 0) or 0),
                natural_gas_cost_wan=float(d.get('natural_gas_cost_wan', 0) or 0),
                heating_cost_wan=float(d.get('heating_cost_wan', 0) or 0),
                petrol_cost_wan=float(d.get('petrol_cost_wan', 0) or 0),
            )
            yd_objects.append(yd)

        # 优先使用 project.indicators 中预计算的指标
        indicators = self.report_data.get('chapter5', {}).get('indicators', {})
        indicator_yearly = {
            item['year']: item
            for item in indicators.get('yearly', [])
            if isinstance(item, dict) and 'year' in item
        }

        # 指标单元格安全取值：指标 dict 带 error（基础数据不足）或缺失时返回 default，
        # 不抛 KeyError，并在 5.3 末尾汇总友好提示。
        indicator_warnings = []

        def _cell(r, key, fmt='{:.2f}', divisor=None, default='—'):
            """从指标 dict 取单元格值；error/缺失/非数字时返回 default。"""
            if isinstance(r, dict) and not r.get('error'):
                v = r.get(key)
                if v is not None:
                    try:
                        if divisor:
                            v = v / divisor
                        return fmt.format(v)
                    except (TypeError, ValueError):
                        pass
            return default

        def _eval_cell(r, bm, val_key='kgce_per_m2'):
            """评价结果单元格：优先指标自带评价，其次按约束/基准/引导值判断。"""
            if isinstance(r, dict) and not r.get('error'):
                bm_r = r.get('benchmark')
                if isinstance(bm_r, dict) and bm_r.get('评价结果'):
                    return bm_r['评价结果']
                val = r.get(val_key)
                if val is not None:
                    if val <= bm['引导值']:
                        return '低于引导值'
                    if val <= bm['基准值']:
                        return '低于基准值'
                    if val <= bm['约束值']:
                        return '低于约束值'
                    return '高于约束值'
            return '—'

        # 收集能源类型
        all_types = set()
        for d in energy_data_list:
            for et in ['electricity_kwh','water_m3','natural_gas_m3','petrol_kg','diesel_kg']:
                if d.get(et, 0) and float(d.get(et, 0) or 0) > 0:
                    all_types.add(et)
        energy_types = sorted(all_types)

        # ===== 总述 =====
        years_str = '、'.join([str(yd.year) for yd in yd_objects])
        self._add_body_text(
            f"为全面准确分析{unit_name}用能情况和用能规律，此次能源审计工作选取"
            f"{yd_objects[0].year}年-{yd_objects[-1].year}年完整年周期内"
            f"{'、'.join([self._ENERGY_TYPE_CN.get(et, et) for et in energy_types])}"
            f"等用能数据，并根据近三年总用能及各项用能数据进行计算分析。"
        )

        # ===== 5.1 概况 =====
        if yd_objects:
            latest = yd_objects[-1]
            self._add_heading_2("5.1  能源资源消费/消耗概况")
            et_cn = [self._ENERGY_TYPE_CN.get(et, et) for et in energy_types]
            self._add_body_text(f"{unit_name}主要用能类型包括{'、'.join(et_cn)}。")

            # 能源流向图
            chart_dir = os.path.join(os.getcwd(), 'charts')
            flow_path = os.path.join(chart_dir, 'energy_flow.png')
            from tools.energy_audit.energy_flow_chart import draw_energy_flow_diagram
            draw_energy_flow_diagram(
                energy_types=energy_types,
                equipment=self.report_data.get('chapter6', {}).get('_equipment', []),
                unit_name=unit_name,
                output_path=flow_path,
            )
            if os.path.exists(flow_path):
                self._add_image_with_caption(flow_path, '图5.1  能源流向图')

        # ===== 5.2 能源资源消耗/消费数据（按用能类型动态生成） =====
        if yd_objects:
            self._add_heading_2("5.2  能源资源消耗/消费数据")
            section_num = 1

            # 每种能源类型一个 H3 小节
            energy_type_sections = [
                ('electricity_kwh', '用电情况分析', 'kWh', '万kWh'),
                ('water_m3', '用水情况分析', 'm³', 'm³'),
                ('natural_gas_m3', '用气情况分析', 'm³', 'm³'),
                ('heating_energy_heat_gj', '用暖情况分析', 'GJ', 'GJ'),
                ('petrol_kg', '用油情况分析（汽油）', 'kg', 'kg'),
                ('diesel_kg', '用油情况分析（柴油）', 'kg', 'kg'),
            ]

            for et_key, et_title, unit, _ in energy_type_sections:
                # et_key → YearlyEnergyData 属性名映射
                _yd_attr_map = {
                    'electricity_kwh': 'electricity_kwh', 'water_m3': 'water_m3',
                    'natural_gas_m3': 'natural_gas_m3', 'heating_energy_heat_gj': 'heating_energy_heat',
                    'petrol_kg': 'transportation_petrol_kg', 'diesel_kg': 'transportation_diesel_kg',
                }
                yd_attr = _yd_attr_map.get(et_key, et_key)
                # 检查该能源类型是否有实际数据
                has_data = any(getattr(yd, yd_attr, 0) and float(getattr(yd, yd_attr, 0) or 0) > 0 for yd in yd_objects)
                if not has_data:
                    continue

                self._add_heading_3(f"5.2.{section_num}  {et_title}")
                vals = [f'{float(getattr(yd, yd_attr, 0) or 0):,.0f}' for yd in yd_objects]
                name_map = {'electricity_kwh':'用电量','water_m3':'用水量','natural_gas_m3':'用气量',
                            'heating_energy_heat_gj':'用热量','petrol_kg':'汽油用量','diesel_kg':'柴油用量'}
                name_cn = name_map.get(et_key, et_title)

                # 段1: 设备描述（按能源类型区分） + 各年具体数值
                eq = self.report_data.get('chapter6', {}).get('_equipment', [])
                if et_key == 'electricity_kwh':
                    elec_cats = ('空调','照明','办公','动力','专用','热水器')
                    eq_names = [e.get('name','') for e in eq if e.get('category') in elec_cats]
                    eq_str = f"主要用电设备为{'、'.join(eq_names)}等" if eq_names else "主要用电设备为照明、空调、办公设备等"
                elif et_key == 'water_m3':
                    eq_str = "用水系统主要为生活用水、卫生清洁用水等"
                elif et_key == 'natural_gas_m3':
                    kitchen = [e.get('name','') for e in eq if e.get('category') == '厨房'][:3]
                    eq_str = f"用气系统主要为{'、'.join(kitchen) if kitchen else '厨房炊事'}用气"
                elif et_key == 'heating_energy_heat_gj':
                    eq_str = "供暖采用市政集中供热，末端为散热器采暖"
                elif et_key == 'petrol_kg':
                    eq_str = "汽油主要用于公务车辆"
                elif et_key == 'diesel_kg':
                    eq_str = "柴油主要用于公务车辆或备用发电机"
                else:
                    eq_str = f"主要消耗{'、'.join(e.get('name','') for e in eq[:3])}等" if eq else ""

                year_parts = []
                for j, v in enumerate(vals):
                    vn = float(v.replace(',',''))
                    year_parts.append(f"{yd_objects[j].year}年全年{name_cn}为{vn:,.0f}{unit}")
                year_str = "，".join(year_parts)

                self._add_body_text(
                    f"{unit_name}{eq_str}。"
                    f"根据能耗账单及缴费发票统计分析，{unit_name}{year_str}。"
                    f"{yd_objects[0].year}年-{yd_objects[-1].year}年总{name_cn}变化趋势如图5.{section_num}所示。"
                )

                # 逐年柱状图
                chart_path = _generate_single_energy_chart(yd_objects, yd_attr, name_cn, unit, chart_dir)
                if chart_path and os.path.exists(chart_path):
                    self._add_image_with_caption(chart_path,
                        f'图5.{section_num}  {yd_objects[0].year}年-{yd_objects[-1].year}年总{name_cn}（单位：{unit}）')

                # 段2: 由图分析 —— 整体评价 + 逐年增减量率 + 异常标注
                if len(vals) >= 2:
                    raw_vals = [float(v.replace(',','')) for v in vals]
                    # 计算各年同比变化率
                    yoy_changes = []
                    yoy_rates = []
                    for j in range(1, len(raw_vals)):
                        diff = raw_vals[j] - raw_vals[j-1]
                        rate = diff / raw_vals[j-1] * 100 if raw_vals[j-1] else 0
                        yoy_changes.append(diff)
                        yoy_rates.append(rate)

                    # --- 整体评价 ---
                    max_rate = max(abs(r) for r in yoy_rates) if yoy_rates else 0
                    all_positive = all(r > 0 for r in yoy_rates)
                    all_negative = all(r < 0 for r in yoy_rates)

                    if all_positive:
                        trend_desc = "整体呈上升趋势" if max_rate < 5 else ("逐年增长" if max_rate < 15 else "逐年持续增长")
                    elif all_negative:
                        trend_desc = "整体呈下降趋势" if max_rate < 5 else ("逐年递减" if max_rate < 15 else "逐年持续下降")
                    elif len(yoy_rates) >= 2 and yoy_rates[0] > 0 and yoy_rates[-1] < 0:
                        trend_desc = "先增加后减少，在一定范围内波动"
                    elif len(yoy_rates) >= 2 and yoy_rates[0] < 0 and yoy_rates[-1] > 0:
                        trend_desc = "先减少后增加，在一定范围内波动"
                    elif max_rate < 5:
                        trend_desc = "整体波动不大"
                    elif max_rate < 15:
                        trend_desc = "整体呈小幅波动"
                    else:
                        trend_desc = "整体波动较大"

                    self._add_body_text(
                        f"由图5.{section_num}分析，{unit_name}{yd_objects[0].year}年-"
                        f"{yd_objects[-1].year}年总{name_cn}{trend_desc}。"
                    )

                    # --- 逐年增减量 + 增减率 ---
                    year_lines = []
                    anomaly_years = []
                    for j in range(1, len(raw_vals)):
                        diff = yoy_changes[j-1]
                        rate = yoy_rates[j-1]
                        abs_diff = abs(diff)
                        abs_rate = abs(rate)
                        prev_val = raw_vals[j-1]
                        curr_val = raw_vals[j]
                        if prev_val == 0 and curr_val > 0:
                            year_lines.append(
                                f"{yd_objects[j].year}年新增{name_cn}{curr_val:,.0f}{unit}"
                            )
                        elif prev_val == 0 and curr_val == 0:
                            year_lines.append(
                                f"{yd_objects[j].year}年与{yd_objects[j-1].year}年均无{name_cn}"
                            )
                        elif abs_rate < 0.01:
                            year_lines.append(
                                f"{yd_objects[j].year}年{name_cn}较{yd_objects[j-1].year}年基本持平"
                            )
                        else:
                            direction = "增加" if diff > 0 else "降低"
                            year_lines.append(
                                f"{yd_objects[j].year}年{name_cn}较{yd_objects[j-1].year}年"
                                f"{direction}{abs_diff:,.0f}{unit}，{direction}率为{abs_rate:.2f}%"
                            )
                        if abs_rate > 15:
                            anomaly_years.append(yd_objects[j].year)

                    if year_lines:
                        self._add_body_text("，".join(year_lines) + "。")

                    # --- 异常标注 ---
                    if anomaly_years:
                        anomaly_text = (
                            f"其中{'、'.join(str(y) for y in anomaly_years)}年"
                            f"{name_cn}变化幅度超过15%，属于异常波动，"
                            f"建议进一步核实该年度用能异常原因。"
                        )
                        self._add_body_text(anomaly_text)

                # 段3: 逐月趋势引语 + 逐月柱状图 + 逐月分析
                monthly_attr = {'electricity_kwh':'monthly_electricity_kwh','water_m3':'monthly_water_m3',
                                'natural_gas_m3':'monthly_natural_gas_m3'}.get(et_key)
                monthly_path = _generate_monthly_bar_chart(energy_data_list, et_key, monthly_attr, name_cn, unit, chart_dir)
                if monthly_path and os.path.exists(monthly_path):
                    self._add_body_text(
                        f"结合{unit_name}审计期内逐月{name_cn}统计数据，"
                        f"分析近三年逐月{name_cn}情况，变化趋势如图5.{section_num}a所示。"
                    )
                    self._add_image_with_caption(monthly_path,
                        f'图5.{section_num}a  {yd_objects[0].year}年-{yd_objects[-1].year}年逐月{name_cn}趋势（单位：{unit}）')
                    # 逐月分析（段4）
                    monthly_data_for_analysis = [d.get(monthly_attr) for d in energy_data_list if d.get(monthly_attr) and len(d.get(monthly_attr, []) or []) == 12]
                    if monthly_data_for_analysis:
                        # 取各年月均（跨年取均值，避免单年偏差）
                        month_avgs = [sum(m[i] for m in monthly_data_for_analysis) / len(monthly_data_for_analysis) for i in range(12)]
                        mc = ['1月','2月','3月','4月','5月','6月','7月','8月','9月','10月','11月','12月']
                        avg_all = sum(month_avgs) / 12
                        mx_val = max(month_avgs); mx_idx = month_avgs.index(mx_val)
                        mn_val = min(month_avgs); mn_idx = month_avgs.index(mn_val)
                        # 找出高于均值的月份（高峰月）
                        high_months = [i for i, v in enumerate(month_avgs) if v > avg_all * 1.05]

                        # 获取建筑冷/热源信息
                        buildings = self.report_data.get('chapter2', {}).get('buildings', [])
                        cold_sources = list(set(b.get('cooling_source','') for b in buildings if b.get('cooling_source')))
                        heat_sources = list(set(b.get('heating_source','') for b in buildings if b.get('heating_source')))
                        cooling_terminals = list(set(b.get('cooling_terminal','') for b in buildings if b.get('cooling_terminal')))

                        if et_key == 'electricity_kwh' and (cold_sources or heat_sources):
                            # === 用电逐月分析（结合冷/热源因果链） ===
                            lines = []
                            # 1. 识别高峰月份段
                            high_month_names = [mc[i] for i in high_months]
                            if high_month_names:
                                lines.append(
                                    f"由图5.{section_num}a分析，{unit_name}用电量较大月份为"
                                    f"{'、'.join(high_month_names)}。"
                                )

                            # 2. 制冷季分析：基于冷源类型解释
                            cooling_months = [i for i in high_months if i in (5,6,7,8)]  # 6-9月
                            transition_low = [i for i in (3,4,9,10) if month_avgs[i] < avg_all]
                            if cooling_months:
                                cold_str = '、'.join(cold_sources) if cold_sources else '制冷机组'
                                # 判断冷源是否电驱动
                                elec_driven_keywords = ['水冷','风冷','冷水机组','多联机','分体','VRV','变频']
                                is_elec_cooling = any(kw in str(cold_sources) for kw in elec_driven_keywords)
                                if is_elec_cooling:
                                    lines.append(
                                        f"结合制冷空调形式分析，夏季空调冷源为{cold_str}，"
                                        f"为电驱动式制冷机组，制冷动力消耗为电，受夏季制冷因素的影响，"
                                        f"{'、'.join(mc[i] for i in cooling_months)}单月用电量明显高于"
                                        f"{'、'.join(mc[i] for i in transition_low) if transition_low else '过渡季'}"
                                        f"无制冷需求或需求较少的月份，单月用电量变化受季节性影响因素较大。"
                                    )
                                else:
                                    lines.append(
                                        f"夏季制冷由{cold_str}提供，制冷季（{'、'.join(mc[i] for i in cooling_months)}）"
                                        f"用电量明显高于过渡季。"
                                    )

                            # 3. 供暖季分析
                            heating_months = [i for i in high_months if i in (11,0,1)]  # 12-2月
                            if heating_months:
                                heat_str = '、'.join(heat_sources) if heat_sources else '市政供暖'
                                is_central_heat = any(kw in str(heat_sources) for kw in ['市政','集中','蒸汽','热力'])
                                if is_central_heat:
                                    lines.append(
                                        f"冬季以{heat_str}为主，热源主要为供暖热水，热源形式虽不是电力直接驱动，"
                                        f"但换热站循环泵24h开启，同时部分房间使用分体空调辅助供暖，"
                                        f"因此供暖季（{'、'.join(mc[i] for i in heating_months)}）用电量同样高于"
                                        f"{'、'.join(mc[i] for i in transition_low) if transition_low else '过渡季'}无供暖需求或需求较少的月份。"
                                    )
                                else:
                                    lines.append(
                                        f"冬季供暖由{heat_str}提供，供暖季（{'、'.join(mc[i] for i in heating_months)}）"
                                        f"用电量同样高于过渡季。"
                                    )

                            # 4. 收口总结
                            lines.append(
                                f"{unit_name}供冷、供暖系统年用电量占总用电量的比重较大，"
                                f"逐月用电量变化符合单位用能变化以及季节性用能变化规律。"
                            )

                            for line in lines:
                                self._add_body_text(line)
                        else:
                            # 非用电类型：简化分析
                            season_peaks = [i for i in high_months if i in (5,6,7,8)]
                            winter_peaks = [i for i in high_months if i in (11,0,1)]
                            if season_peaks and len(season_peaks) >= 3:
                                s = "夏季月份用量明显偏高，受季节性因素影响较大"
                            elif winter_peaks and len(winter_peaks) >= 3:
                                s = "冬季月份用量明显偏高，受季节性因素影响较大"
                            elif mx_val > mn_val * 1.5:
                                s = f"各月用量存在一定波动，最高月份为{mc[mx_idx]}"
                            else:
                                s = "各月用量分布较为均匀"
                            self._add_body_text(
                                f"由图5.{section_num}a分析，{unit_name}逐月{name_cn}"
                                f"{s}，最高月份为{mc[mx_idx]}（{mx_val:,.0f}{unit}），"
                                f"最低月份为{mc[mn_idx]}（{mn_val:,.0f}{unit}）。"
                            )

                    # 月份同比异常检测（跨年同期对比）
                    if monthly_data_for_analysis and len(monthly_data_for_analysis) >= 2:
                        yoy_anomalies = []  # [(year, month_name, rate, current_val, prev_val)]
                        for m_idx in range(12):
                            for y_idx in range(1, len(monthly_data_for_analysis)):
                                prev_val = monthly_data_for_analysis[y_idx-1][m_idx]
                                curr_val = monthly_data_for_analysis[y_idx][m_idx]
                                if prev_val and prev_val > 0:
                                    yoy_rate = (curr_val - prev_val) / prev_val * 100
                                    if abs(yoy_rate) > 30:
                                        year = yd_objects[y_idx].year
                                        yoy_anomalies.append((year, mc[m_idx], yoy_rate, curr_val, prev_val))

                        if yoy_anomalies:
                            # 分组：按年份汇总异常月份
                            from collections import defaultdict
                            by_year = defaultdict(list)
                            for yr, mn, rate, cv, pv in yoy_anomalies:
                                direction = "增长" if rate > 0 else "降低"
                                by_year[yr].append(
                                    f"{mn}（较上年同期{direction}{abs(rate):.1f}%，"
                                    f"由{pv:,.0f}{unit}变为{cv:,.0f}{unit}）"
                                )

                            anomaly_lines = []
                            for yr in sorted(by_year.keys()):
                                months_str = "、".join(by_year[yr])
                                anomaly_lines.append(f"{yr}年：{months_str}")

                            # 判断是否可解释（如学校寒暑假）：当前仅对医院/机关等无长假的机构做标记
                            inst_cat = self.report_data.get('tags', {}).get('institution_category', '')
                            is_school = '教育' in str(inst_cat)
                            if is_school:
                                self._add_body_text(
                                    f"经月份同比分析发现，审计期内部分月份{name_cn}存在较大年际波动："
                                    f"{'；'.join(anomaly_lines)}。上述波动可能与学校寒暑假期间用能规律有关，"
                                    f"属正常周期性变化。"
                                )
                            else:
                                self._add_body_text(
                                    f"经月份同比分析发现，审计期内部分月份{name_cn}存在较大年际波动"
                                    f"（同比变化率超过±30%）：{'；'.join(anomaly_lines)}。"
                                    f"建议进一步核实以上月份{name_cn}异常波动原因。"
                                )
                section_num += 1

            # 5.2.N+1 能源资源费用分析
            cost_sections = [
                ('electricity_cost_wan', '电费', '万元'),
                ('natural_gas_cost_wan', '燃气费', '万元'),
                ('water_cost_wan', '水费', '万元'),
                ('heating_cost_wan', '供暖费', '万元'),
                ('petrol_cost_wan', '油费', '万元'),
            ]
            has_cost = any(
                any(getattr(yd, k, 0) and float(getattr(yd, k, 0) or 0) > 0 for yd in yd_objects)
                for k, _, _ in cost_sections
            )
            if has_cost:
                self._add_heading_3(f"5.2.{section_num}  能源资源费用分析")

                # 统计有数据的费用类型
                active_costs = [(k, n, u) for k, n, u in cost_sections
                                if any(getattr(yd, k, 0) and float(getattr(yd, k, 0) or 0) > 0 for yd in yd_objects)]
                cost_names = [n for _, n, _ in active_costs]

                # 段1：开头介绍 + 引出表格
                cost_name_str = '、'.join(cost_names)
                self._add_body_text(
                    f"根据所提供的能耗账单及缴费发票，{unit_name}能源费用主要包括"
                    f"{cost_name_str}等。"
                    f"{yd_objects[0].year}年-{yd_objects[-1].year}年各项能源费用统计"
                    f"见表5.{section_num}。"
                )

                # 费用统计表
                cost_headers = ['费用类型'] + [f'{yd.year}年（万元）' for yd in yd_objects]
                cost_rows = []
                for cost_key, cost_name, _ in active_costs:
                    row = [cost_name]
                    for yd in yd_objects:
                        val = float(getattr(yd, cost_key, 0) or 0)
                        row.append(f'{val:,.2f}')
                    cost_rows.append(row)
                # 合计行
                total_row = ['合计']
                for yd in yd_objects:
                    total = sum(float(getattr(yd, k, 0) or 0) for k, _, _ in active_costs)
                    total_row.append(f'{total:,.2f}')
                cost_rows.append(total_row)

                self._add_table(
                    cost_headers, cost_rows,
                    title=f'表5.{section_num}  {yd_objects[0].year}年-{yd_objects[-1].year}年能源费用统计'
                )

                # 各年能源费用占比饼状图
                for yd in yd_objects:
                    cost_vals = [float(getattr(yd, k, 0) or 0) for k, _, _ in active_costs]
                    if all(v == 0 for v in cost_vals):
                        continue
                    pie_path = _generate_cost_pie_chart(yd.year, cost_names, cost_vals, chart_dir)
                    if pie_path and os.path.exists(pie_path):
                        self._add_image_with_caption(pie_path,
                            f'图5.{section_num}  {yd.year}年能源费用占比')

                # 费用分析文字（三层：整体→逐年→占比）
                # 计算各年总费用
                total_costs = []
                for yd in yd_objects:
                    total = sum(float(getattr(yd, k, 0) or 0) for k, _, _ in active_costs)
                    total_costs.append(total)

                # 1. 整体变化评价
                if len(total_costs) >= 2:
                    total_rates = []
                    for j in range(1, len(total_costs)):
                        prev = total_costs[j-1]
                        if prev > 0:
                            total_rates.append((total_costs[j] - prev) / prev * 100)
                    max_total_rate = max(abs(r) for r in total_rates) if total_rates else 0
                    if max_total_rate < 5:
                        overall = "整体变化不大"
                    elif max_total_rate < 15:
                        overall = "整体呈小幅波动"
                    else:
                        overall = "整体波动较大"
                    self._add_body_text(
                        f"{unit_name}{yd_objects[0].year}年-{yd_objects[-1].year}年能源费用{overall}。"
                    )

                # 2. 逐年费用增减
                if len(total_costs) >= 2:
                    year_cost_lines = []
                    for j in range(1, len(total_costs)):
                        diff = total_costs[j] - total_costs[j-1]
                        rate = diff / total_costs[j-1] * 100 if total_costs[j-1] else 0
                        abs_diff = abs(diff); abs_rate = abs(rate)
                        if total_costs[j-1] == 0 and total_costs[j] > 0:
                            year_cost_lines.append(
                                f"{yd_objects[j].year}年新增能源费用{total_costs[j]:,.2f}万元"
                            )
                        elif abs_rate < 0.01:
                            year_cost_lines.append(
                                f"{yd_objects[j].year}年能源费用与{yd_objects[j-1].year}年基本持平"
                            )
                        else:
                            direction = "增加" if diff > 0 else "降低"
                            year_cost_lines.append(
                                f"{yd_objects[j].year}年能源费用较{yd_objects[j-1].year}年"
                                f"{direction}{abs_diff:,.2f}万元，{direction}率为{abs_rate:.2f}%"
                            )
                    if year_cost_lines:
                        self._add_body_text("，".join(year_cost_lines) + "。")

                # 3. 占比最高分析
                latest_yd = yd_objects[-1]
                latest_costs = [(cost_name, float(getattr(latest_yd, cost_key, 0) or 0)) 
                                for cost_key, cost_name, _ in active_costs]
                latest_total = sum(v for _, v in latest_costs)
                if latest_total > 0:
                    # 找占比最高的
                    top_name, top_val = max(latest_costs, key=lambda x: x[1])
                    top_pct = top_val / latest_total * 100
                    top_sorted = sorted(latest_costs, key=lambda x: x[1], reverse=True)
                    if top_pct >= 50:
                        self._add_body_text(
                            f"能源费用占比最高的为{top_name}，"
                            f"占该单位能源消费总额的{top_pct:.0f}%以上，"
                            f"影响{unit_name}能源费用的主要因素是{top_name}支出。"
                        )
                    elif len(top_sorted) >= 2:
                        n1, v1 = top_sorted[0]; n2, v2 = top_sorted[1]
                        p1 = v1 / latest_total * 100; p2 = v2 / latest_total * 100
                        self._add_body_text(
                            f"能源费用占比最高的为{n1}（{p1:.1f}%）和{n2}（{p2:.1f}%），"
                            f"两项合计占能源消费总额的{p1+p2:.0f}%以上，"
                            f"是影响{unit_name}能源费用的主要支出类型。"
                        )
            # 计算 5.3 的表号起点
            next_table = section_num + 1
        else:
            next_table = 2

        # ===== 5.3 指标 =====
        self._add_heading_2("5.3  能耗资源消耗/消费指标")

        # 5.3.1 非供暖能耗
        self._add_heading_3("5.3.1  单位建筑面积非供暖能耗")

        # 机构类型名称映射（定义动词 + 特殊用能注）
        inst_cat = self.report_data.get('tags', {}).get('institution_category', '')
        ic = str(inst_cat)
        if '医疗' in ic:
            org_type = '医疗机构'
            action_desc = '从事疾病诊断、治疗活动'
            special_note = '大型医疗设备、数据中心、厨房炊事、洗衣房'
        elif '政务' in ic or '服务' in ic:
            org_type = '政务服务中心'
            action_desc = '日常办公'
            special_note = '数据中心、厨房炊事'
        elif '场馆' in ic or '体育' in ic or '文化' in ic:
            org_type = '场馆机构'
            action_desc = '运行'
            special_note = '数据中心、厨房炊事、专业设备'
        elif '教育' in ic:
            org_type = '教育机构'
            action_desc = '运行'
            special_note = '数据中心、实验室、厨房炊事'
        elif '党政' in ic:
            org_type = '党政机关'
            action_desc = '日常办公'
            special_note = '数据中心、厨房炊事、专业用途设备'
        else:
            org_type = '公共机构'
            action_desc = '日常运行'
            special_note = '数据中心、厨房炊事'

        # 段1: 定义
        self._add_body_text(
            f"{org_type}单位建筑面积非供暖能耗：在统计报告期内，{org_type}{action_desc}过程中，"
            f"除供暖能耗和交通能耗之外消耗的各种能源实物量，包括机构内供冷、通风、照明、"
            f"生活热水、电梯、办公设备以及机构内供暖系统的热水循环泵电耗、供暖用的风机电耗等"
            f"所使用的能耗，按照规定的计算方法和单位折算为标准煤后的总和与建筑面积的比值。"
            f"单位为千克标准煤每平方米年，kgce/（m²·a）。计算公式如下："
        )

        # 段2: OMML公式 + "式中：" + 符号说明（公式变量斜体）
        self._add_formula("E_jrcn = (E - E_gn - E_jt) / M")
        self._add_body_text("式中：")
        self._add_formula_symbol("E_jrcn——", "单位建筑面积非供暖能耗，单位为千克标准煤每平方米年，kgce/(m²·a)；")
        self._add_formula_symbol("E——", "综合能耗，单位为千克标准煤每年，kgce/a；")
        self._add_formula_symbol("E_gn——", "供暖能耗，单位为千克标准煤每年，kgce/a；")
        self._add_formula_symbol("E_jt——", "交通能耗，单位为千克标准煤每年，kgce/a；")
        self._add_formula_symbol("M——", "建筑面积，单位为平方米，m²。")

        # 段3: 注
        self._add_body_text(
            f"注：{org_type}内的{special_note}等特定功能的用能不计入{org_type}非供暖能耗，"
            f"计算单位建筑面积非供暖能耗时，应同时剔除特殊用能系统对应的建筑面积。"
        )

        # 段4: 特殊说明（信息中心未单独计量，不剔除）
        self._add_body_text(
            f"{unit_name}设置信息中心，但信息中心用电未进行单独计量，"
            f"计算单位建筑面积非供暖能耗时未进行剔除，"
            f"{yd_objects[0].year}年-{yd_objects[-1].year}年单位建筑面积非供暖能耗指标如表5.{next_table}所示："
        )

        # 数据表
        headers = ['统计周期'] + [f'{yd.year}年' for yd in yd_objects]
        rows = [['非供暖能耗(tce)'], ['建筑面积(m²)'], ['单位面积非供暖能耗(kgce/m²)'], ['约束值'], ['基准值'], ['引导值'], ['评价结果']]
        bm = resolve_benchmark(institution_type, 'unit_area_non_heating')
        for yd in yd_objects:
            r = indicator_yearly.get(yd.year, {}).get('unit_area_non_heating')
            if not r:
                r = calc_unit_area_non_heating_energy(yd)
            if isinstance(r, dict) and r.get('error'):
                indicator_warnings.append(f"{yd.year}年单位建筑面积非供暖能耗：{r['error']}")
            rows[0].append(_cell(r, 'non_heating_kgce', divisor=1000))
            rows[1].append(f"{area:.2f}")
            rows[2].append(_cell(r, 'kgce_per_m2'))
            rows[3].append(str(bm['约束值']))
            rows[4].append(str(bm['基准值']))
            rows[5].append(str(bm['引导值']))
            rows[6].append(_eval_cell(r, bm))
        self._add_table(headers, rows, f"表5.{next_table}  单位建筑面积非供暖能耗")

        # 5.3.2 常规用能系统单位建筑面积电耗
        self._add_heading_3("5.3.2  常规用能系统单位建筑面积电耗")

        # 段1: 定义（统一使用 action_desc，无特殊动词区分）
        self._add_body_text(
            f"{org_type}常规用能系统单位建筑面积电耗：在统计报告期内，{org_type}"
            f"{action_desc}过程中，由照明插座、空调、动力等用能系统消耗的电量总和"
            f"与建筑面积的比值。单位为千瓦时每平方米年，kWh/（m²·a）。计算公式如下："
        )

        # 段2: OMML公式 E_ja = E_D / M
        self._add_simple_fraction_formula("E_ja", "E_D", "M")
        self._add_body_text("式中：")
        self._add_formula_symbol("E_ja——", "常规用能系统单位建筑面积电耗，单位为千瓦时每平方米年，kWh/(m²·a)；")
        self._add_formula_symbol("E_D——", "电量总和，单位为千瓦时每年，kWh/a；")
        self._add_formula_symbol("M——", "建筑面积，单位为平方米，m²。")

        # 段3: 引出表格
        self._add_body_text(
            f"{unit_name}{yd_objects[0].year}年-{yd_objects[-1].year}年"
            f"常规用能系统单位建筑面积电耗指标如表5.{next_table+1}所示："
        )

        headers = ['统计周期'] + [f'{yd.year}年' for yd in yd_objects]
        rows = [['用电量(kWh)'], ['单位面积电耗(kWh/m²)'], ['约束值'], ['基准值'], ['引导值'], ['评价结果']]
        bm2 = resolve_benchmark(institution_type, 'unit_area_elec')
        for yd in yd_objects:
            r = indicator_yearly.get(yd.year, {}).get('unit_area_electricity')
            if not r:
                r = calc_unit_area_electricity(yd, institution_type=institution_type, sub_type=venue_sub)
            if isinstance(r, dict) and r.get('error'):
                indicator_warnings.append(f"{yd.year}年常规用能系统单位建筑面积电耗：{r['error']}")
            rows[0].append(_cell(r, 'total_electricity_kwh', '{:,.0f}'))
            rows[1].append(_cell(r, 'kwh_per_m2'))
            rows[2].append(str(bm2['约束值']))
            rows[3].append(str(bm2['基准值']))
            rows[4].append(str(bm2['引导值']))
            rows[5].append(_eval_cell(r, bm2, val_key='kwh_per_m2'))
        self._add_table(headers, rows, f"表5.{next_table+1}  常规用能系统单位建筑面积电耗")

        # 5.3.3 人均综合能耗
        self._add_heading_3("5.3.3  人均综合能耗")

        # P的人员构成按机构类型
        if '医疗' in ic:
            p_desc = "在岗在编的人员、各类编外工作人员、门诊人数、床位数"
        elif '教育' in ic:
            p_desc = "在岗在编的教职员工、各类编外工作人员及学生人数"
        else:
            p_desc = "在岗在编的人员、各类编外工作人员"

        # 段1: 定义
        self._add_body_text(
            f"{org_type}人均综合能耗：在统计报告期内，{org_type}综合能耗与用能人数的比值。"
            f"单位为千克标准煤每人年，kgce/（p·a）。计算公式如下："
        )

        # 段2: 公式 + 式中（E_r = E / P）
        self._add_simple_fraction_formula("E_r", "E", "P")
        self._add_body_text("式中：")
        self._add_formula_symbol("E_r——", "人均综合能耗，单位为千克标准煤每人年，kgce/（p·a）；")
        self._add_formula_symbol("E——",
            f"综合能耗，在统计报告期内，{org_type}{action_desc}过程中，"
            f"实际消耗的各种能源实物量，按照规定的计算方法和单位分别折算为标准煤后的总和。"
            f"单位为千克标准煤每年，kgce/a；")
        self._add_formula_symbol("P——",
            f"用能人数，单位为人，p。在统计报告期内，{org_type}用能人数为{org_type}"
            f"消耗能源的日平均人员数量，包括{p_desc}。"
            f"经统计，{unit_name}{yd_objects[0].year}年-{yd_objects[-1].year}年"
            f"用能人数为{people}人。")

        # 段3: 引出表
        self._add_body_text(
            f"{unit_name}{yd_objects[0].year}年-{yd_objects[-1].year}年"
            f"人均综合能耗指标如表5.{next_table+2}所示："
        )

        headers = ['统计周期'] + [f'{yd.year}年' for yd in yd_objects]
        rows = [['综合能耗(kgce)'], ['用能人数'], ['人均综合能耗(kgce/人)'], ['约束值'], ['基准值'], ['引导值'], ['评价结果']]
        bm3 = resolve_benchmark(institution_type, 'per_capita_energy')
        for yd in yd_objects:
            r = indicator_yearly.get(yd.year, {}).get('per_capita_energy')
            if not r:
                r = calc_per_capita_energy(yd, institution_type=institution_type, sub_type=venue_sub)
            if isinstance(r, dict) and r.get('error'):
                indicator_warnings.append(f"{yd.year}年人均综合能耗：{r['error']}")
            rows[0].append(_cell(r, 'total_kgce'))
            rows[1].append(str(people))
            rows[2].append(_cell(r, 'kgce_per_person'))
            rows[3].append(str(bm3['约束值']))
            rows[4].append(str(bm3['基准值']))
            rows[5].append(str(bm3['引导值']))
            rows[6].append(_eval_cell(r, bm3, val_key='kgce_per_person'))
        self._add_table(headers, rows, f"表5.{next_table+2}  人均综合能耗")

        # 5.3.4 用水指标（按机构类型分三种形态）
        if '医疗' in ic:
            # === 医疗：卫生业单位用水量（单位开放床日用水量） ===
            self._add_heading_3("5.3.4  单位开放床日用水量")
            self._add_body_text(
                f"单位时间内，三级、二级、一级综合医院住院部的"
                f"单位开放床日用水量计算公式如下："
            )
            self._add_simple_fraction_formula("V_Z", "W_Z", "N_i")
            self._add_body_text("式中：")
            self._add_formula_symbol("V_Z——",
                "三级、二级、一级综合医院住院部单位开放床日用水量，单位为L/（床·d）；")
            self._add_formula_symbol("W_Z——",
                f"三级、二级、一级综合医院住院部年用水总量（包括住院部、医技部、教学科研、"
                f"后勤、行政管理等用水量，不包括洗衣、制药、试验用水量和家属区、宿舍、"
                f"幼儿园、招待所等用水量和外供水量），单位为m³/a；")
            self._add_formula_symbol("N_i——",
                "三级、二级、一级综合医院第i日的实际开放床日数，单位为床·d。")
            self._add_body_text(
                f"{unit_name}{yd_objects[0].year}年-{yd_objects[-1].year}年"
                f"卫生业单位用水量指标如表5.{next_table+3}所示："
            )
            headers = ['统计周期'] + [f'{yd.year}年' for yd in yd_objects]
            rows = [['取水量(m³)'], ['床位数'], ['单位开放床日用水量(L/床·d)'], ['通用值'], ['先进值'], ['评价结果']]
            beds = self.report_data.get('chapter2', {}).get('beds_count', 0) or 0
            for yd in yd_objects:
                r = indicator_yearly.get(yd.year, {}).get('per_capita_water')
                if not r:
                    r = calc_per_capita_water(yd, institution_type=institution_type, bed_count=beds)
                if isinstance(r, dict) and r.get('error'):
                    indicator_warnings.append(f"{yd.year}年卫生业单位用水量：{r['error']}")
                rows[0].append(_cell(r, 'total_water_m3', '{:,.2f}'))
                rows[1].append(str(beds))
                rows[2].append(_cell(r, 'L_per_bed_day', '{:,.2f}'))
                bm_w = (r.get('benchmark') or {}) if isinstance(r, dict) else {}
                rows[3].append(str(bm_w.get('约束值', bm_w.get('通用值', ''))))
                rows[4].append(str(bm_w.get('引导值', bm_w.get('先进值', ''))))
                rows[5].append(bm_w.get('评价结果', '—'))
            self._add_table(headers, rows, f"表5.{next_table+3}  卫生业单位用水量")

        elif '教育' in ic:
            # === 教育：单位标准人数用水量 ===
            self._add_heading_3("5.3.4  单位标准人数用水量")
            self._add_body_text(f"单位时间内，单位标准人数用水量计算公式如下：")
            self._add_simple_fraction_formula("V_rs", "V_k", "N_p")
            self._add_body_text("式中：")
            self._add_formula_symbol("V_rs——",
                "单位标准人数用水量，单位为立方米每人年（m³/p·a）；")
            self._add_formula_symbol("V_k——",
                "年用水量，单位为立方米每年（m³/a）；")
            self._add_formula_symbol("N_p——",
                "标准人数，单位为p。")
            self._add_body_text(
                f"{unit_name}{yd_objects[0].year}年-{yd_objects[-1].year}年"
                f"单位标准人数用水量指标如表5.{next_table+3}所示："
            )
            headers = ['统计周期'] + [f'{yd.year}年' for yd in yd_objects]
            rows = [['取水量(m³)'], ['标准人数'], ['单位标准人数用水量(m³/p·a)'], ['通用值'], ['先进值'], ['评价结果']]
            for yd in yd_objects:
                r = indicator_yearly.get(yd.year, {}).get('per_capita_water')
                if not r:
                    r = calc_per_capita_water(yd, institution_type=institution_type)
                if isinstance(r, dict) and r.get('error'):
                    indicator_warnings.append(f"{yd.year}年单位标准人数用水量：{r['error']}")
                rows[0].append(_cell(r, 'total_water_m3', '{:,.2f}'))
                rows[1].append(str(people))
                rows[2].append(_cell(r, 'm3_per_person', '{:,.2f}'))
                bm_w = (r.get('benchmark') or {}) if isinstance(r, dict) else {}
                rows[3].append(str(bm_w.get('约束值', bm_w.get('通用值', ''))))
                rows[4].append(str(bm_w.get('引导值', bm_w.get('先进值', ''))))
                rows[5].append(bm_w.get('评价结果', '—'))
            self._add_table(headers, rows, f"表5.{next_table+3}  单位标准人数用水量")

        else:
            # === 党政/政务/场馆/教育：取水指标名按机构类型动态 ===
            water_name = self._water_indicator_name(ic)
            area_based = '政务' in ic or '场馆' in ic or '体育' in ic
            water_unit = '立方米每平方米年（m³/(m²·a)）' if area_based else '立方米每人年（m³/(p·a)）'
            self._add_heading_3(f"5.3.4  {water_name}")
            self._add_body_text(f"单位时间内，{water_name}计算公式如下：")
            self._add_simple_fraction_formula("V_rc", "V_k", "N_p")
            self._add_body_text("式中：")
            self._add_formula_symbol("V_rc——",
                f"{water_name}，单位为{water_unit}；")
            self._add_formula_symbol("V_k——",
                "年取水量，单位为立方米每年（m³/a）；")
            self._add_formula_symbol("N_p——",
                ("建筑面积，单位为平方米（m²）。" if area_based
                 else "用能人数（机关人数/在校人数），单位为p。"))
            self._add_body_text(
                f"{unit_name}{yd_objects[0].year}年-{yd_objects[-1].year}年"
                f"{water_name}指标如表5.{next_table+3}所示："
            )
            headers = ['统计周期'] + [f'{yd.year}年' for yd in yd_objects]
            if area_based:
                rows = [['取水量(m³)'], ['建筑面积(m²)'], ['单位建筑面积年取水量(m³/(m²·a))'], ['通用值'], ['先进值'], ['评价结果']]
            else:
                rows = [['取水量(m³)'], ['用能人数'], ['人均取水量(m³/人)'], ['通用值'], ['先进值'], ['评价结果']]
            for yd in yd_objects:
                r = indicator_yearly.get(yd.year, {}).get('per_capita_water')
                if not r:
                    r = calc_per_capita_water(yd, institution_type=institution_type, building_area=area)
                if isinstance(r, dict) and r.get('error'):
                    indicator_warnings.append(f"{yd.year}年{water_name}：{r['error']}")
                rows[0].append(_cell(r, 'total_water_m3', '{:,.2f}'))
                if area_based:
                    if not area or area <= 0:
                        indicator_warnings.append(f"{yd.year}年{water_name}：建筑面积缺失，无法计算")
                    rows[1].append(str(area) if area else '—')
                    rows[2].append(_cell(r, 'm3_per_area', '{:,.4f}'))
                else:
                    rows[1].append(str(people))
                    rows[2].append(_cell(r, 'm3_per_person', '{:,.2f}'))
                bm_w = (r.get('benchmark') or {}) if isinstance(r, dict) else {}
                rows[3].append(str(bm_w.get('约束值', bm_w.get('通用值', ''))))
                rows[4].append(str(bm_w.get('引导值', bm_w.get('先进值', ''))))
                rows[5].append(bm_w.get('评价结果', '—'))
            self._add_table(headers, rows, f"表5.{next_table+3}  {water_name}")

        # 5.3 指标计算不足提示（缺数据时友好提示而非崩溃）
        if indicator_warnings:
            self._add_body_text(
                "注：因基础数据不足，上表部分指标未能计算，以'—'标注。"
                "具体原因：" + '；'.join(sorted(set(indicator_warnings))) + "。"
                "请补充建筑面积、用能人数、床位数等基础参数及逐年能耗数据后重新生成报告。"
            )

        # ===== 5.4 能耗基准 =====
        self._add_heading_2("5.4  建筑能耗基准")
        from tools.energy_audit.indicators import calc_baseline
        # 优先使用 project.indicators 中预计算的基准
        bl = indicators.get('baseline', {})
        if not bl:
            # 只对有数据的能源类型计算基准
            bl_yd = [yd for yd in yd_objects if yd.electricity_kwh > 0]
            if bl_yd:
                bl = calc_baseline(bl_yd)
            # 用量基准表
            hdr_use = ['能源类型','基准值','单位','计算方法']
            rows_use = []
            for label, info in bl.get('usage', {}).items():
                rows_use.append([label, str(info['基准值']), info.get('单位',''), info.get('方法','')])
            if rows_use:
                self._add_table(hdr_use, rows_use, f"表5.{next_table+4}  能源资源用量基准")
            # 费用基准表
            hdr_cost = ['费用类型','基准值(万元)','计算方法']
            rows_cost = []
            for label, info in bl.get('cost', {}).items():
                rows_cost.append([label, str(info['基准值']), info.get('方法','')])
            if rows_cost:
                self._add_table(hdr_cost, rows_cost, f"表5.{next_table+5}  能源资源费用基准")

    # ============================================================
    # 第6章：系统分析（按实际设备动态生成）

    def build_chapter6(self):
        """第6章 按参考报告结构：用电(8子节) → 用水 → 其他用能 → 室内环境"""
        ch6 = self.report_data.get('chapter6', {})
        unit = self.report_data.get('chapter1', {}).get('audited_unit_short', '被审计单位')
        eq_list = ch6.get('_equipment', [])
        buildings = self.report_data.get('chapter2', {}).get('buildings', [])
        equip_images = self.report_data.get('images_equipment', [])
        institution_cat = self.report_data.get('tags', {}).get('institution_category', '')

        self._add_heading_1("第6章  主要能源资源利用系统分析")

        # ========== 6.1 用电系统 ==========
        self._add_heading_2("6.1  用电系统运行分析")

        # 总述：列出所有用电设备
        all_elec_cat = ['空调','照明','办公','热水器','厨房']
        if '医疗' in institution_cat:
            all_elec_cat.append('医疗')
        all_elec = [e for e in eq_list if e.get('category') in all_elec_cat]
        if all_elec:
            cat_summary = {}
            for e in all_elec:
                cat = e.get('category','')
                cat_summary[cat] = cat_summary.get(cat, []) + [e.get('name','')]
            parts = []
            cat_names = {'空调':'空调系统','照明':'照明系统','办公':'办公系统','热水器':'电热水器','厨房':'厨房设备','医疗':'大型医疗设备'}
            for cat in ['空调','照明','办公','医疗','热水器','厨房']:
                if cat in cat_summary:
                    names = list(dict.fromkeys(cat_summary[cat]))[:3]
                    parts.append(f"{cat_names.get(cat,cat)}，主要为{'、'.join(names)}等")
            self._add_body_text(f"{unit}用电系统主要包括{'；'.join(parts)}。")

        # 6.1.1 空调与供暖系统
        cooling_eq = [e for e in all_elec if e.get('category') == '空调']
        if cooling_eq:
            self._add_heading_3("空调与供暖系统")
            # 冷源
            b0 = buildings[0] if buildings else {}
            cold_src = b0.get('cooling_source', '') or '分体式空调'
            cold_term = b0.get('cooling_terminal', '') or '—'
            heat_src = b0.get('heating_source', '') or '市政供暖'
            heat_term = b0.get('heating_terminal', '') or '散热器'
            self._add_body_text(
                f"{unit}夏季供冷冷源主要为{cold_src}，末端形式为{cold_term}。"
                f"主要设备有{'、'.join(e.get('name','') for e in cooling_eq)}等，"
                f"共{sum(e.get('quantity',0) for e in cooling_eq)}台。"
            )
            # 热源
            self._add_body_text(
                f"冬季供暖热源为{heat_src}，末端形式为{heat_term}。"
                f"供冷时间为每年6月～9月，供暖时间为每年11月至次年3月，"
                f"共计约120天，办公时间运行。"
            )
            # 照片 + 设备表
            for img in equip_images[:2]:
                if os.path.exists(img):
                    self._add_image(img, width_cm=7)
            eq_table = [{'name':e.get('name',''),'spec':e.get('spec',''),'qty':e.get('quantity',1),'remark':e.get('remark',''),'independent_metering':e.get('independent_metering',''),'independent_metering_desc':e.get('independent_metering_desc','')} for e in cooling_eq]
            self._add_equipment_table(eq_table, "表6.1  空调与供暖系统设备清单")

        # 6.1.2 照明系统
        lighting_eq = [e for e in all_elec if e.get('category') == '照明']
        if lighting_eq:
            self._add_heading_3("照明系统")
            light_names = '、'.join(e.get('name','') for e in lighting_eq)
            light_qty = sum(e.get('quantity',0) for e in lighting_eq)
            self._add_body_text(
                f"{unit}照明灯具主要为{light_names}等，共{light_qty}套。"
                f"办公区域照明运行时间约为每天8:00-18:00，公共区域采用定时或声控开关控制。"
            )
            for img in equip_images[2:4]:
                if os.path.exists(img):
                    self._add_image(img, width_cm=7)
            eq_table = [{'name':e.get('name',''),'spec':e.get('spec',''),'qty':e.get('quantity',1),'remark':e.get('remark',''),'independent_metering':e.get('independent_metering',''),'independent_metering_desc':e.get('independent_metering_desc','')} for e in lighting_eq]
            self._add_equipment_table(eq_table, "表6.2  照明系统设备清单")

        # 6.1.3 办公设备
        office_eq = [e for e in all_elec if e.get('category') == '办公']
        if office_eq:
            self._add_heading_3("办公设备系统")
            self._add_body_text(f"{unit}办公设备主要包括{'、'.join(e.get('name','') for e in office_eq)}等。")
            eq_table = [{'name':e.get('name',''),'spec':e.get('spec',''),'qty':e.get('quantity',1),'remark':e.get('remark',''),'independent_metering':e.get('independent_metering',''),'independent_metering_desc':e.get('independent_metering_desc','')} for e in office_eq]
            self._add_equipment_table(eq_table, "表6.3  办公设备清单")

        # 6.1.4 大型医疗设备 / 其他用电设备（根据机构类型）
        medical_eq = [e for e in all_elec if e.get('category') == '医疗']
        hot_water_eq = [e for e in all_elec if e.get('category') == '热水器']
        if medical_eq:
            self._add_heading_3("大型医疗设备系统")
            self._add_body_text(f"{unit}大型医疗设备主要包括{'、'.join(e.get('name','') for e in medical_eq)}等。")
            eq_table = [{'name':e.get('name',''),'spec':e.get('spec',''),'qty':e.get('quantity',1),'remark':e.get('remark',''),'independent_metering':e.get('independent_metering',''),'independent_metering_desc':e.get('independent_metering_desc','')} for e in medical_eq]
            self._add_equipment_table(eq_table, "表6.4  大型医疗设备清单")
        elif hot_water_eq:
            self._add_heading_3("其他用电设备")
            self._add_body_text(
                f"{unit}其他用电设备主要包括{'、'.join(e.get('name','') for e in hot_water_eq)}等，"
                f"共{sum(e.get('quantity',0) for e in hot_water_eq)}台，"
                f"用于职工日常饮用水供应，全天候运行。"
            )
            eq_table = [{'name':e.get('name',''),'spec':e.get('spec',''),'qty':e.get('quantity',1),'remark':e.get('remark',''),'independent_metering':e.get('independent_metering',''),'independent_metering_desc':e.get('independent_metering_desc','')} for e in hot_water_eq]
            self._add_equipment_table(eq_table, "表6.4  其他用电设备清单")

        # 6.1.5 厨房设备
        kitchen_eq = [e for e in all_elec if e.get('category') == '厨房']
        if kitchen_eq:
            self._add_heading_3("厨房设备")
            self._add_body_text(f"{unit}厨房设备主要包括{'、'.join(e.get('name','') for e in kitchen_eq)}等。")
            eq_table = [{'name':e.get('name',''),'spec':e.get('spec',''),'qty':e.get('quantity',1),'remark':e.get('remark',''),'independent_metering':e.get('independent_metering',''),'independent_metering_desc':e.get('independent_metering_desc','')} for e in kitchen_eq]
            self._add_equipment_table(eq_table, "表6.5  厨房设备清单")

        # 6.1.6 信息机房（提示补充）
        self._add_heading_3("信息机房系统")
        self._add_body_text(
            f"⚠️ 请提供{unit}信息机房系统信息：机房面积、机柜数量、服务器数量、UPS配置、"
            f"精密空调情况；信息机房用电是否纳入本单位用能统计、是否实现单独计量管理。"
        )

        # 6.1.7 变配电系统（提示补充）
        self._add_heading_3("变配电系统")
        self._add_body_text(
            f"⚠️ 请提供{unit}变配电系统信息：变压器型号、容量、数量、安装位置；"
            f"配电柜配置情况；是否有节能型变压器。"
        )

        # ========== 6.2 用水系统 ==========
        self._add_heading_2("6.2  用水系统运行分析")
        self._add_body_text(
            f"{unit}用水主要为生活用水、卫生间用水及空调系统补水。"
            f"生活用水由市政自来水直供，主要为清洗用水、饮用水。"
        )

        # ========== 6.3 其他用能系统 ==========
        has_gas = any(e.get('category') == '厨房' for e in eq_list)
        if has_gas:
            self._add_heading_2("6.3  其他用能系统运行分析")
            self._add_body_text(
                f"{unit}设有餐厅，用气设备为厨房炊具，使用天然气。"
            )

        # ========== 6.4 室内环境检测 ==========
        indoor = ch6.get('indoor_env')
        if indoor and isinstance(indoor, dict):
            self._add_heading_2("6.4  室内环境检测")
            rooms = indoor.get('rooms', [])
            test_date = indoor.get('test_date', '')
            test_cond = indoor.get('test_conditions', '')
            if test_date or test_cond:
                self._add_body_text(
                    f"本次能源审计对{unit}建筑室内环境进行了现场检测，"
                    f"检测时间为{test_date or '—'}，{test_cond or '—'}。"
                )
            if rooms:
                # 检测结果表
                self._add_table(
                    ['房间','温度(℃)','湿度(%)','CO2(ppm)','照度(lx)','风速(m/s)','VOC(mg/m³)','PM2.5(μg/m³)'],
                    [[r.get('room',''), str(r.get('temp','')), str(r.get('humidity','')),
                      str(r.get('co2','')), str(r.get('illumination','')),
                      str(r.get('wind_speed','')), str(r.get('voc','')), str(r.get('pm25',''))]
                     for r in rooms],
                    "表6.N  室内环境检测结果"
                )
                # 标准参考表
                self._add_table(
                    ['参数','标准值','标准来源'],
                    [['温度','22~28℃(夏季)/16~24℃(冬季)','GB/T 18883-2022'],
                     ['湿度','40%~80%(夏季)/30%~60%(冬季)','GB/T 18883-2022'],
                     ['CO2','≤1000ppm','GB/T 18883-2022'],
                     ['照度','300~500lx(办公)','GB 50034-2024'],
                     ['风速','≤0.3m/s(夏季)/≤0.2m/s(冬季)','GB/T 18883-2022'],
                     ['PM2.5','≤75μg/m³','GB/T 18883-2022']],
                    "表6.N+1  室内环境标准参考值"
                )
                self._add_body_text(
                    f"参考《建筑照明设计标准》(GB 50034-2024)及《室内空气质量标准》"
                    f"(GB/T 18883-2022)，{unit}室内温湿度满足舒适性标准，"
                    f"CO2浓度正常，室内环境质量较好。"
                )
        else:
            self._add_body_text(f"⚠️ 请提供{unit}室内环境检测数据。")
        
    def _add_equipment_table(self, equipment: list, title: str, headers=None):
        """添加设备清单表。headers 可自定义，默认：序号|设备名称|规格/功率|数量|备注"""
        if not equipment:
            return
        
        from docx.shared import Cm as _Cm
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        if headers is None:
            has_metering = any(str(eq.get('independent_metering') or '').strip() for eq in equipment)
            if has_metering:
                headers = ['序号', '设备名称', '规格/功率', '数量', '独立计量', '备注']
            else:
                headers = ['序号', '设备名称', '规格/功率', '数量', '备注']
        rows = []
        show_metering = '独立计量' in headers
        for i, eq in enumerate(equipment, 1):
            remark = eq.get('remark', '') or ''
            desc = (eq.get('independent_metering_desc') or '').strip()
            if desc and desc not in remark:
                remark = f"{remark}；{desc}".strip('；') if remark else desc
            row = [str(i), eq.get('name', ''), eq.get('spec', ''), str(eq.get('qty', ''))]
            if show_metering:
                row.append(eq.get('independent_metering', '') or '')
            row.append(remark)
            rows.append(row)

        # 标题
        para = self.doc.add_paragraph()
        para.alignment = self.WD_ALIGN.CENTER
        run = para.add_run(title)
        self._set_font(run, FMT.body_font_cn, FMT.body_font_en, FMT.table_header_size, bold=True)

        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')

        # 表头
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.clear()
            p.alignment = self.WD_ALIGN.CENTER
            r = p.add_run(h)
            self._set_font(r, FMT.body_font_cn, FMT.body_font_en, FMT.table_content_size, bold=True)

        # 数据行
        for ri, row_data in enumerate(rows):
            row = table.rows[ri + 1]
            row.height = _Cm(1.0)
            for ci, val in enumerate(row_data):
                cell = row.cells[ci]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.clear()
                p.alignment = self.WD_ALIGN.CENTER
                r = p.add_run(val)
                self._set_font(r, FMT.body_font_cn, FMT.body_font_en, FMT.table_content_size)

        self.doc.add_paragraph()

    # ============================================================
    # 第7章：节能效果与节能潜力分析（动态H3）
    # ============================================================

    def build_chapter7(self):
        """第7章 节能效果与节能潜力分析（问题从数据推断 + 用户补充兜底）"""
        ch7 = self.report_data.get('chapter7', {})
        ch4 = self.report_data.get('chapter4', {})
        ch6 = self.report_data.get('chapter6', {})
        tags = self.report_data.get('tags', {})
        unit = (self.report_data.get('chapter1', {}).get('audited_unit_short', '') or
                self.report_data.get('chapter2', {}).get('unit_name', '被审计单位'))
        self._add_heading_1("第7章  节能效果与节能潜力分析")

        problems = ch7.get('problems', [])

        # 如果用户没提供问题，尝试从 metering/equipment 数据推断
        if not problems:
            inferred = []

            # 从 chapter4 计量信息推断
            metering = self.report_data.get('metering', {})
            has_monitoring = metering.get('has_monitoring_system', False) if metering else False
            has_separate = metering.get('has_separate_metering', False) if metering else False
            has_household = metering.get('has_household_metering', False) if metering else False

            if not has_separate:
                inferred.append({
                    'title': '分项计量体系不完善',
                    'text': '目前仅实现了建筑总用电量的计量，尚未对空调、照明、办公设备等主要用能系统进行分项计量，无法精确掌握各系统的能耗分布情况。建议按照GB/T29149-2012要求，完善二级和三级计量体系，实现用能分项、分区精准计量。'
                })

            # 从 chapter6 设备推断
            eq_list = ch6.get('_equipment', [])
            lighting_eq = [e for e in eq_list if e.get('category') == '照明']
            cooling_eq = [e for e in eq_list if e.get('category') == '空调']
            all_specs = ' '.join(e.get('spec','') for e in eq_list)

            # 空调能效
            if cooling_eq and '3级' in all_specs:
                inferred.append({
                    'title': '部分空调设备能效偏低',
                    'text': f'现有空调设备中部分为3级能效产品，制冷效率偏低。建议制定分年度更换计划，优先更换使用年限长、能效低的设备，逐步替换为1级能效产品。'
                })

            # 水泵/电机设备
            pump_eq = [e for e in eq_list if e.get('category') in ('水泵','电机','动力')]
            if not pump_eq:
                # 没有水泵数据 → 不写水泵类问题
                pass

            # 热水壶/开水器
            hot_water_eq = [e for e in eq_list if e.get('category') == '热水器']
            if hot_water_eq:
                inferred.append({
                    'title': '饮用水设备节能优化',
                    'text': f'现有饮用水热水器{sum(e.get("quantity",0) for e in hot_water_eq)}台，全天候运行。建议加装定时控制器，在非工作时间（18:00-次日8:00及周末）自动断电，预计可减少约30%的待机能耗。'
                })

            # 从建筑数据推断
            buildings = self.report_data.get('chapter2', {}).get('buildings', [])
            old_buildings = [b for b in buildings if b.get('year', 9999) < 2000]
            if old_buildings:
                inferred.append({
                    'title': '老旧建筑围护结构节能改造',
                    'text': f'{",".join(b.get("name","") for b in old_buildings)}建成于{old_buildings[0].get("year","")}年，建议对建筑外窗、外墙保温等围护结构进行节能改造，降低建筑冷热负荷。'
                })

            # 建议项（按实际情况判断）
            if not has_household:
                inferred.append({
                    'title': '完善能源管理制度',
                    'text': '建议建立健全能源管理体系，制定年度节能目标并分解至各科室，将节能工作纳入绩效考核；加强用能行为管理，减少不必要的能源浪费。'
                })

            if inferred:
                problems = inferred
                print(f"[自动推断] 第7章问题 ({len(problems)}项)")

        # 7.1
        if problems:
            self._add_heading_2("7.1  用能系统现状分析")
            for i, p in enumerate(problems, 1):
                self._add_heading_3(f"7.1.{i}  {p['title']}")
                for para_text in p['text'].split('\n'):
                    if para_text.strip():
                        self._add_body_text(para_text.strip())

        solutions = ch7.get('solutions', [])
        if not solutions and problems:
            solutions = []
            sol_map = {
                '计量': {'title': '完善分项计量体系', 'text': '按照GB/T29149-2012要求，增设二级和三级计量器具，实现空调、照明、办公设备等各用能系统的独立分项计量。为安装分项计量表计的位置（配电室、空调机房等）配置数据采集器，实现能耗数据的自动采集和实时上传。'},
                '空调': {'title': '空调系统能效提升', 'text': '制定分年度更换计划，将3级能效空调设备逐步更换为1级能效产品。同时优化空调运行策略：夏季制冷温度不低于26℃，冬季供暖温度不高于20℃；非工作时间关闭无人区域空调。预计可节省空调用电15%～20%。'},
                '饮用水': {'title': '饮用水设备节能改造', 'text': '为饮用水热水器加装定时控制器，设定工作日8:00-18:00运行，周末及节假日自动断电。年节电量约为设备全年用电量的30%，预计每年可节约用电约1,000～2,000kWh。'},
                '围护结构': {'title': '建筑围护结构节能改造', 'text': '对老旧建筑的外窗和外墙保温系统进行节能改造，采用中空Low-E玻璃替换普通玻璃窗，增加外墙保温层厚度。改造后可降低建筑冬季采暖负荷约15%，夏季空调冷负荷约10%。'},
                '能源管理': {'title': '健全能源管理体系', 'text': '建立由单位主要负责人牵头的能源管理领导小组，制定年度节能目标并分解到各科室。建立健全能耗统计台账，定期开展能耗分析。加强节能宣传培训，提高全体职工的节能意识。'},
                '监测': {'title': '建设能耗监测系统', 'text': '按GB/T29149-2012要求建设能耗监测系统，实现用电分项计量和用水分区计量。'},
                '可再生': {'title': '推进可再生能源利用', 'text': '利用屋顶建设光伏发电或太阳能热水系统，提高可再生能源利用率。'},
            }
            for p in problems:
                matched = None
                for kw, sol in sol_map.items():
                    if kw in p['title']:
                        matched = sol.copy()
                        break
                if matched:
                    solutions.append(matched)
            if solutions:
                solutions.append({'title': '综合节能效果评估', 'text': f'上述措施实施后，预计可实现年节约用电20,000~50,000kWh，综合节能率约为5%~10%。'})

        if solutions:
            self._add_heading_2("7.2  节能潜力分析及建议")
            for i, s in enumerate(solutions, 1):
                self._add_heading_3(f"7.2.{i}  {s['title']}")
                for para_text in s['text'].split('\n'):
                    if para_text.strip():
                        self._add_body_text(para_text.strip())

        # 汇总表
        summary = ch7.get('summary', {})
        if summary:
            hdrs = summary.get('headers', ['序号','改造项目','预估投资(万元)','年节能量(tce)','节能率','回收期(年)'])
            rows = summary.get('rows', [])
            if rows:
                self._add_table(title="表7.1  节能改造措施汇总", headers=hdrs, rows=rows)

    # ============================================================
    # 第8章：审计结论（LLM综合生成）
    # ============================================================

    def build_chapter8(self):
        """第8章 审计结论（自动聚合前7章指标+问题+建议）"""
        ch8 = self.report_data.get('chapter8', {})
        ch5 = self.report_data.get('chapter5', {})
        ch7 = self.report_data.get('chapter7', {})
        ch1 = self.report_data.get('chapter1', {})
        tags = self.report_data.get('tags', {})
        unit = ch1.get('audited_unit_short', '被审计单位')
        inst = 'government'; cat = tags.get('institution_category', '')
        if '医疗' in cat: inst = 'medical'

        self._add_heading_1("第8章  审计结论")

        # 如果用户提供了文本，直接用
        text = ch8.get('text', '')
        if text:
            for p in text.split('\n'):
                if p.strip(): self._add_body_text(p.strip())
            return

        # 自动生成
        from tools.energy_audit.indicators import resolve_benchmark
        bm = resolve_benchmark(inst, 'unit_area_non_heating')
        std_name = bm.get('标准', '')
        # 修正：按机构类型取正确的定额标准名
        if not std_name or '医疗' in std_name:
            std_map = {'medical': 'DB37/T 2673-2019《医疗机构能源消耗定额标准》',
                       'government': 'DB37/T 2672-2019《党政机关能源消耗定额标准》',
                       'education': 'DB37/T 2674-2019《教育机构能源消耗定额标准》'}
            std_name = std_map.get(inst, std_name or 'DB37/T 2672-2019')
        bm2 = resolve_benchmark(inst, 'unit_area_elec')
        bm3 = resolve_benchmark(inst, 'per_capita_energy')

        self._add_body_text(
            f"此次对{unit}的能源审计工作主要采用了资料收集、现场勘察与问询、数据计算与分析等多种审计方法，"
            f"力求全面准确地对单位整体用能情况与用能系统作出分析评价。"
        )

        # 指标汇总（优先使用 chapter5 indicators 预计算结果）
        indicators = ch5.get('indicators', {})
        if indicators.get('status') == 'ok' and indicators.get('yearly'):
            self._add_body_text(f"参照{std_name}及DB37/T 4452-2021，各项能耗指标评价结果如下：")
            beds = ch5.get('beds_count', 0) or 0
            for item in indicators['yearly']:
                yr = item.get('year', '')
                r_nh = item.get('unit_area_non_heating', {})
                r_ed = item.get('unit_area_electricity', {})
                r_pe = item.get('per_capita_energy', {})
                r_pw = item.get('per_capita_water', {})

                nh = r_nh.get('kgce_per_m2', 0)
                ed = r_ed.get('kwh_per_m2', 0)
                pe = r_pe.get('kgce_per_person', 0)
                ev1 = _indicator_eval(r_nh)
                ev2 = _indicator_eval(r_ed)
                ev3 = _indicator_eval(r_pe)

                if beds:
                    pw = r_pw.get('L_per_bed_day', 0)
                    ev4 = _indicator_eval(r_pw)
                    pw_text = f"单位开放床日用水量{pw:.2f}L/(床·d)（{ev4}）"
                else:
                    pw = r_pw.get('m3_per_person', 0)
                    ev4 = _indicator_eval(r_pw)
                    pw_text = f"人均取水量{pw:.2f}m³/(p·a)（{ev4}）"

                self._add_body_text(
                    f"{yr}年：单位建筑面积非供暖能耗{nh:.2f}kgce/(m²·a)（{ev1}），"
                    f"单位建筑面积电耗{ed:.2f}kWh/(m²·a)（{ev2}），"
                    f"人均综合能耗{pe:.2f}kgce/(p·a)（{ev3}），"
                    f"{pw_text}。"
                )
        elif ch5.get('energy_data'):
            # 兜底：使用 indicators 现场计算
            self._add_body_text(f"参照{std_name}及DB37/T 4452-2021，各项能耗指标评价结果如下：")
            from tools.energy_audit.indicators import (
                YearlyEnergyData, calc_unit_area_non_heating_energy,
                calc_unit_area_electricity, calc_per_capita_energy, calc_per_capita_water,
            )
            area = ch5.get('building_area', 0)
            people = ch5.get('people_count', 0)
            beds = ch5.get('beds_count', 0) or 0
            for d in ch5['energy_data']:
                yr = d.get('year', '')
                yd = YearlyEnergyData(
                    year=int(yr),
                    electricity_kwh=float(d.get('electricity_kwh', 0) or 0),
                    water_m3=float(d.get('water_m3', 0) or 0),
                    natural_gas_m3=float(d.get('natural_gas_m3', 0) or 0),
                    heating_energy_heat=float(d.get('heating_energy_heat_gj', 0) or 0),
                    transportation_petrol_kg=float(d.get('petrol_kg', 0) or 0),
                    transportation_diesel_kg=float(d.get('diesel_kg', 0) or 0),
                    building_area=area, people_count=people,
                )
                r_nh = calc_unit_area_non_heating_energy(yd)
                r_ed = calc_unit_area_electricity(yd, institution_type=inst, sub_type=venue_sub)
                r_pe = calc_per_capita_energy(yd, institution_type=inst, sub_type=venue_sub)
                r_pw = calc_per_capita_water(yd, institution_type=inst, bed_count=beds)

                nh = r_nh.get('kgce_per_m2', 0)
                ed = r_ed.get('kwh_per_m2', 0)
                pe = r_pe.get('kgce_per_person', 0)
                ev1 = _indicator_eval(r_nh)
                ev2 = _indicator_eval(r_ed)
                ev3 = _indicator_eval(r_pe)

                if beds:
                    pw = r_pw.get('L_per_bed_day', 0)
                    ev4 = _indicator_eval(r_pw)
                    pw_text = f"单位开放床日用水量{pw:.2f}L/(床·d)（{ev4}）"
                else:
                    pw = r_pw.get('m3_per_person', 0)
                    ev4 = _indicator_eval(r_pw)
                    pw_text = f"人均取水量{pw:.2f}m³/(p·a)（{ev4}）"

                self._add_body_text(
                    f"{yr}年：单位建筑面积非供暖能耗{nh:.2f}kgce/(m²·a)（{ev1}），"
                    f"单位建筑面积电耗{ed:.2f}kWh/(m²·a)（{ev2}），"
                    f"人均综合能耗{pe:.2f}kgce/(p·a)（{ev3}），"
                    f"{pw_text}。"
                )

        # 问题汇总
        problems = ch7.get('problems', [])
        if problems:
            titles = [p['title'] for p in problems]
            self._add_body_text(f"经能源审计发现，{unit}目前存在{'、'.join(titles)}等问题。")

        # 建议条数
        solutions = ch7.get('solutions', [])
        if solutions:
            sol_titles = [s['title'] for s in solutions[:6]]
            self._add_body_text(f"针对上述问题，提出可行节能改造建议共计{len(solutions)}条，主要包括：{'、'.join(sol_titles)}等。")

    def _generate_cost_pie_chart(self, years_data, output_dir: str = './charts'):
        """最新年能源费用占比饼图（电费/水费/天然气费/热力费，单位：万元）"""
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            setup_chart_font(plt)
        except ImportError:
            return None

        os.makedirs(output_dir, exist_ok=True)
        latest = years_data[-1]
        cost_fields = [
            ("electricity_cost_wan", "电费"),
            ("water_cost_wan", "水费"),
            ("natural_gas_cost_wan", "天然气费"),
            ("heating_cost_wan", "热力费"),
        ]
        labels = []
        values = []
        for key, label in cost_fields:
            v = float(latest.get(key, 0) or 0)
            if v > 0:
                labels.append(label)
                values.append(v)
        if not values:
            return None
        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90,
               colors=['#4CAF50', '#2196F3', '#FF9800', '#F44336'])
        ax.set_title(chart_text(f'{latest.get("year", "")}年能源费用占比'))
        path = os.path.join(output_dir, 'chart_cost_structure.png')
        fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        return path

    def _generate_energy_pie_chart(self, yd_objects, energy_types, output_dir: str = './charts'):
        """生成最新年能源结构的饼图"""
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            setup_chart_font(plt)
        except ImportError:
            return None

        os.makedirs(output_dir, exist_ok=True)
        latest = yd_objects[-1]
        labels = [self._ENERGY_TYPE_CN.get(et, et) for et in energy_types]
        coeff_map = {'electricity_kwh':0.1229,'water_m3':0.2571,'natural_gas_m3':1.33,'petrol_kg':1.4714,'diesel_kg':1.4571}
        values = [getattr(latest, et, 0) * coeff_map.get(et, 1) / 1000 for et in energy_types]
        if not any(v > 0 for v in values):
            return None
        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90)
        ax.set_title(chart_text(f'{latest.year}年能源消费结构'))
        path = os.path.join(output_dir, 'chart_structure.png')
        fig.savefig(path, dpi=150, bbox_inches='tight')
        plt.close(fig)
        return path

    def _generate_yearly_trend_chart(self, yd_objects, energy_types, output_dir: str = './charts'):
        """逐年能耗趋势柱状图"""
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            setup_chart_font(plt)
        except ImportError:
            return None

        os.makedirs(output_dir, exist_ok=True)
        years = [str(yd.year) for yd in yd_objects]
        # 只显示 top 3 能源类型
        coeff_map = {'electricity_kwh':0.1229,'water_m3':0.2571,'natural_gas_m3':1.33,'petrol_kg':1.4714,'diesel_kg':1.4571}
        et_values = {}
        for et in energy_types:
            vals = [getattr(yd, et, 0) * coeff_map.get(et, 1) / 1000 for yd in yd_objects]
            if any(v > 0 for v in vals):
                et_values[et] = vals
        if not et_values:
            return None

        fig, ax = plt.subplots(figsize=(8, 5))
        x = range(len(years))
        width = 0.8 / len(et_values)
        for i, (et, vals) in enumerate(et_values.items()):
            ax.bar([xi + i*width for xi in x], vals, width, label=self._ENERGY_TYPE_CN.get(et, et))
        ax.set_xticks([xi + width*(len(et_values)-1)/2 for xi in x])
        ax.set_xticklabels(years)
        ax.set_ylabel('tce')
        ax.set_title(chart_text('逐年能耗趋势'))
        ax.legend()
        ax.grid(True, alpha=0.3, axis='y')
        path = os.path.join(output_dir, 'chart_trend.png')
        fig.savefig(path, dpi=150, bbox_inches='tight')
        plt.close(fig)
        return path

    def _add_image_with_caption(self, image_path: str, caption: str):
        """嵌入图片（宽度12cm，居中），下方居中添加图注（12pt 宋体，1.5倍行距，居中）"""
        from docx.shared import Cm as _Cm
        para = self.doc.add_paragraph()
        para.alignment = self.WD_ALIGN.CENTER
        try:
            run = para.add_run()
            run.add_picture(image_path, width=_Cm(12))
        except Exception as e:
            run = para.add_run(f"[图片无法嵌入: {e}]")
            self._set_font(run, FMT.body_font_cn, FMT.body_font_en, 10)

        # 图注（12pt 宋体，1.5倍行距，居中）
        cap_para = self.doc.add_paragraph()
        cap_para.alignment = self.WD_ALIGN.CENTER
        cap_para.paragraph_format.line_spacing = FMT.body_line_spacing
        cap_run = cap_para.add_run(caption)
        self._set_font(cap_run, FMT.body_font_cn, FMT.body_font_en, FMT.body_size)

    def _add_building_param_table(self, bldg: dict, table_num: int):
        """生成单个建筑的参数表（4列键值对），标题在表格上方"""
        name = bldg.get('name', '')
        title = f"表2-{table_num}  {name}基本信息"
        # 标题单独一段
        self._add_table_title(title)

        rows = [
            ["建筑物名称", name, "建筑地址", bldg.get('address', '')],
            ["建造年代", f"{bldg.get('year', '')}年" if bldg.get('year') else '', "建筑功能", bldg.get('function', '')],
            ["建筑层数", bldg.get('floors', ''), "建筑面积", str(bldg.get('area', '')) + ('m²' if bldg.get('area') else '')],
            ["建筑结构形式", bldg.get('structure', ''), "建筑外窗类型", bldg.get('window_type', '')],
            ["建筑外墙保温", bldg.get('insulation', ''), "建筑朝向", bldg.get('orientation', '')],
            ["建筑功能分区", bldg.get('function_zoning', ''), "建筑高度", bldg.get('height', '')],
            ["夏季空调冷源", bldg.get('cooling_source', ''), "冬季供暖热源", bldg.get('heating_source', '')],
            ["夏季空调末端", bldg.get('cooling_terminal', ''), "冬季供暖末端", bldg.get('heating_terminal', '')],
            ["建筑给水系统", bldg.get('water_system', ''), "建筑消防给水系统", bldg.get('fire_system', '')],
            ["生活热水系统", bldg.get('hot_water', ''), "能耗在线监测系统", bldg.get('monitoring', '')],
            # ============ Part B: 扩展信息 (新增字段) ============
            ["使用面积", str(bldg.get('use_area', '')) + ('m²' if bldg.get('use_area') else ''),
             "供冷面积", str(bldg.get('cooling_area', '')) + ('m²' if bldg.get('cooling_area') else '')],
            ["供热面积", str(bldg.get('heating_area', '')) + ('m²' if bldg.get('heating_area') else ''),
             "外墙主体材料", bldg.get('wall_body_material', '')],
            ["屋面保温", bldg.get('roof_insulation', ''),
             "屋面保温材料", bldg.get('roof_insulation_material', '')],
            ["遮阳形式", bldg.get('sunshade_type', ''),
             "遮阳材料", bldg.get('sunshade_material', '')],
            ["建筑运行时间", bldg.get('run_time', ''),
             "楼层单独计量", bldg.get('storey_metrology', '')],
            ["地下车库", bldg.get('garage', ''),
             "地下车库面积", str(bldg.get('garage_area', '')) + ('m²' if bldg.get('garage_area') else '')],
        ]

        from docx.shared import Cm as _Cm
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        table = self.doc.add_table(rows=len(rows), cols=4, style='Table Grid')
        table.autofit = True

        for ri, row_data in enumerate(rows):
            row = table.rows[ri]
            row.height = _Cm(1.0)
            for ci, val in enumerate(row_data):
                cell = row.cells[ci]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.clear()
                if ci in (0, 2):  # label columns
                    p.alignment = self.WD_ALIGN.CENTER
                    r = p.add_run(str(val) if val else "")
                    self._set_font(r, FMT.body_font_cn, FMT.body_font_en, FMT.table_content_size, bold=True)
                else:  # value columns
                    p.alignment = self.WD_ALIGN.CENTER
                    r = p.add_run(str(val) if val else "")
                    self._set_font(r, FMT.body_font_cn, FMT.body_font_en, FMT.table_content_size)

        self.doc.add_paragraph()

    def get_chapter_reference(self, chapter_key: str, context: str = "") -> str:
        """
        检索同类报告的对应章节作为 LLM 参考（RAG 库函数包装）。

        注意：build_chapter1/3 等章节生成器 **不调用** 本方法。
        第3章现行路径是 PG 节能管理表 + 制度文件 LLM 提炼 + 模板兜底。
        search_for_chapter() 仍可用，但未接入 rest_generate 管线。

        用法:
          ref = builder.get_chapter_reference('第2章', '公共机构基本情况')
          # → 返回 Markdown 格式的参考文本；不会自动写入报告
        """
        tags = self.report_data.get('tags', {})
        if not tags:
            return ""
        try:
            from rag.rag_search import search_for_chapter
            return search_for_chapter(chapter_key, tags, context)
        except Exception as e:
            return f"[RAG参考检索失败: {e}]"

    def generate(self, output_path: str):
        """生成 Word 报告（含目录自动更新与水印）"""
        self.doc = self.Document()

        # 页面设置
        from docx.shared import Cm
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

        self.build_cover()
        self.build_audit_info_tables()
        self.build_toc()
        self.build_all_chapters()

        self.doc.save(output_path)

        # 目录域自动更新：Word 打开时刷新 TOC/页码
        self._set_update_fields_on_open(output_path)
        # 页脚页码：第 X 页 共 Y 页
        self._add_page_numbers(output_path)
        return output_path

    def _add_page_numbers(self, output_path: str):
        """向页脚写入居中页码域：第 X 页 共 Y 页（Word/WPS 打开时随 updateFields 自动刷新）。"""
        from docx import Document as _Doc
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = _Doc(output_path)
        sec = doc.sections[0]
        footer = sec.footer
        # 复用/清空默认页脚段落
        if footer.paragraphs:
            p = footer.paragraphs[0]
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
        else:
            p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _set_font_run(run, size=10.5, bold=False):
            run.font.name = "Times New Roman"
            run.font.size = self.Pt(size)
            run.font.bold = bold
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(_qn("w:rFonts"))
            if rFonts is None:
                from lxml import etree as _etree
                rFonts = _etree.SubElement(rPr, _qn("w:rFonts"))
            rFonts.set(_qn("w:eastAsia"), "宋体")

        def _add_field(fld_type: str):
            """插入 PAGE / NUMPAGES 域"""
            run = p.add_run()
            _set_font_run(run)
            r = run._r
            b = OxmlElement("w:fldChar"); b.set(_qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText"); instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            instr.text = f" {fld_type} "
            s = OxmlElement("w:fldChar"); s.set(_qn("w:fldCharType"), "separate")
            t = OxmlElement("w:t"); t.text = "1"
            e = OxmlElement("w:fldChar"); e.set(_qn("w:fldCharType"), "end")
            r.append(b); r.append(instr); r.append(s); r.append(t); r.append(e)

        r1 = p.add_run("第 "); _set_font_run(r1)
        _add_field("PAGE")
        r2 = p.add_run(" 页 共 "); _set_font_run(r2)
        _add_field("NUMPAGES")
        r3 = p.add_run(" 页"); _set_font_run(r3)

        doc.save(output_path)

    def _set_update_fields_on_open(self, output_path: str):
        """在 word/settings.xml 写入 <w:updateFields w:val="true"/>，使 Word 打开时自动刷新目录域。"""
        import tempfile
        import zipfile
        import shutil
        from pathlib import Path
        from lxml import etree

        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        tmp = Path(tempfile.mkdtemp(prefix="ea_updatefields_"))
        try:
            with zipfile.ZipFile(output_path, "r") as z:
                z.extractall(tmp)
            settings = tmp / "word" / "settings.xml"
            settings.parent.mkdir(parents=True, exist_ok=True)
            if settings.exists():
                tree = etree.parse(str(settings))
                root = tree.getroot()
            else:
                root = etree.Element(f"{{{W_NS}}}settings")
                tree = etree.ElementTree(root)
            uf = root.find(f"{{{W_NS}}}updateFields")
            if uf is None:
                uf = etree.Element(f"{{{W_NS}}}updateFields")
                root.insert(0, uf)
            uf.set(f"{{{W_NS}}}val", "true")
            tree.write(str(settings), xml_declaration=True, encoding="UTF-8", standalone="yes")
            with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
                for f in tmp.rglob("*"):
                    if f.is_dir():
                        continue
                    z.write(f, f.relative_to(tmp).as_posix())
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class MarkdownReportBuilder:
    """生成 Markdown 格式报告（完整内容，结构对齐 Word 版）"""

    def __init__(self, audit_type: str):
        self.audit_type = audit_type
        self.chapters = CHAPTER_STRUCTURES.get(audit_type, CHAPTER_STRUCTURES["公共机构"])
        self.report_data: Dict = {}

    def set_data(self, data: Dict):
        self.report_data = data

    def generate(self) -> str:
        d = self.report_data
        lines = []

        # 封面
        cover = d.get('cover', {})
        lines.append(f"# {cover.get('title', '能源审计报告')}")
        if cover.get('audit_organization'):
            lines.append(f"**{cover['audit_organization']}**")
        if cover.get('report_date'):
            lines.append(f"*{cover['report_date']}*")
        lines.append("")

        # 三张表
        tabs = d.get('audit_info_tables', {})
        inst = tabs.get('institution', {})
        if inst.get('name'):
            lines.append("## 能源审计机构信息表\n")
            lines.append("| 项目 | 内容 |")
            lines.append("|------|------|")
            for k, v in [('机构名称', inst.get('name','')), ('地址', inst.get('address','')),
                         ('负责人', inst.get('contact','')), ('联系方式', inst.get('phone',''))]:
                lines.append(f"| {k} | {v} |")
            lines.append("")

        members = tabs.get('team_members', [])
        if members:
            lines.append("## 能源审计组人员名单\n")
            lines.append("| 职务 | 姓名 | 学历 | 资质 | 专业 |")
            lines.append("|------|------|------|------|------|")
            for m in members[:10]:
                lines.append(f"| {m.get('role','')} | {m.get('name','')} | {m.get('education','')} | {m.get('certification','')} | {m.get('major','')} |")
            lines.append("")

        # 各章
        ch2 = d.get('chapter2', {})
        ch5 = d.get('chapter5', {})
        energy_data = ch5.get('energy_data', [])

        for ch_num, ch_title, sub_sections in self.chapters:
            lines.append(f"## {ch_num}  {ch_title}\n")

            if '第1章' in ch_num:
                ch1 = d.get('chapter1', {})
                unit = ch1.get('audited_unit_short', '') or cover.get('title', '').replace('能源审计报告', '')
                if unit: lines.append(f"**被审计单位**: {unit}")
                if ch1.get('address'): lines.append(f"**地址**: {ch1['address']}")
                if ch1.get('audit_time'): lines.append(f"**审计时间**: {ch1['audit_time']}")
                if ch1.get('audit_period'): lines.append(f"**审计期**: {ch1['audit_period']}")
                if ch1.get('base_period'): lines.append(f"**基准期**: {ch1['base_period']}")
                lines.append("")

            elif '第2章' in ch_num:
                buildings = ch2.get('buildings', [])
                area = ch2.get('building_area', 0)
                people = ch2.get('people_count', 0)
                unit = ch2.get('unit_name', '')
                lines.append(f"**单位**: {unit}")
                lines.append(f"**建筑面积**: {area} m²")
                lines.append(f"**用能人数**: {people}")
                if buildings:
                    lines.append(f"\n### 建筑列表 ({len(buildings)}栋)\n")
                    lines.append("| 名称 | 年代 | 层数 | 结构 |")
                    lines.append("|------|------|------|------|")
                    for b in buildings:
                        lines.append(f"| {b.get('name','')} | {b.get('year','')} | {b.get('floors','')} | {b.get('structure','')} |")
                lines.append("")

            elif '第5章' in ch_num:
                # 优先使用 project.indicators 预计算结果
                indicators = ch5.get('indicators', {})
                if indicators.get('status') == 'ok' and indicators.get('yearly'):
                    lines.append(f"### 逐年能耗数据\n")
                    lines.append("| 年份 | 综合能耗(tce) |")
                    lines.append("|------|-------------|")
                    for item in indicators['yearly']:
                        lines.append(f"| {item.get('year','')} | {item.get('total_energy_tce', 0):.2f} |")
                    lines.append("")
                    lines.append(f"### 指标汇总\n")
                    lines.append("| 年份 | 非供暖(kgce/m²) | 电耗(kWh/m²) | 人均能耗(kgce/人) |")
                    lines.append("|------|----------------|-------------|-----------------|")
                    for item in indicators['yearly']:
                        nh = item.get('unit_area_non_heating', {}).get('kgce_per_m2', 0)
                        ed = item.get('unit_area_electricity', {}).get('kwh_per_m2', 0)
                        pe = item.get('per_capita_energy', {}).get('kgce_per_person', 0)
                        lines.append(f"| {item.get('year','')} | {nh:.2f} | {ed:.2f} | {pe:.2f} |")
                    lines.append("")
                elif energy_data:
                    # 兜底：使用 indicators 现场计算
                    from tools.energy_audit.indicators import (
                        YearlyEnergyData, calc_unit_area_non_heating_energy,
                        calc_unit_area_electricity, calc_per_capita_energy,
                    )
                    area = ch5.get('building_area', 0)
                    people = ch5.get('people_count', 0)
                    inst = ch5.get('institution_type', 'government')
                    lines.append(f"### 逐年能耗数据\n")
                    lines.append("| 年份 | 综合能耗(tce) |")
                    lines.append("|------|-------------|")
                    yd_cache = {}
                    for d in energy_data:
                        yd = YearlyEnergyData(
                            year=int(d.get('year', 0)),
                            electricity_kwh=float(d.get('electricity_kwh', 0) or 0),
                            water_m3=float(d.get('water_m3', 0) or 0),
                            natural_gas_m3=float(d.get('natural_gas_m3', 0) or 0),
                            heating_energy_heat=float(d.get('heating_energy_heat_gj', 0) or 0),
                            transportation_petrol_kg=float(d.get('petrol_kg', 0) or 0),
                            transportation_diesel_kg=float(d.get('diesel_kg', 0) or 0),
                            building_area=area, people_count=people,
                        )
                        yd_cache[d.get('year','')] = yd
                        lines.append(f"| {d.get('year','')} | {yd.total_energy_tce:.2f} |")
                    lines.append("")
                    if area and people:
                        lines.append(f"### 指标汇总\n")
                        lines.append("| 年份 | 非供暖(kgce/m²) | 电耗(kWh/m²) | 人均能耗(kgce/人) |")
                        lines.append("|------|----------------|-------------|-----------------|")
                        for yr, yd in yd_cache.items():
                            r_nh = calc_unit_area_non_heating_energy(yd)
                            r_ed = calc_unit_area_electricity(yd, institution_type=inst)
                            r_pe = calc_per_capita_energy(yd, institution_type=inst)
                            lines.append(
                                f"| {yr} | {r_nh['kgce_per_m2']:.2f} | "
                                f"{r_ed['kwh_per_m2']:.2f} | {r_pe['kgce_per_person']:.2f} |"
                            )
                        lines.append("")

            elif '第6章' in ch_num:
                ch6 = d.get('chapter6', {})
                eq = ch6.get('_equipment', [])
                cooling = [e for e in eq if e.get('category') == '空调']
                if cooling:
                    lines.append("### 主要用电设备\n")
                    lines.append("| 名称 | 规格 | 数量 |")
                    lines.append("|------|------|------|")
                    for e in cooling[:10]:
                        lines.append(f"| {e.get('name','')} | {e.get('spec','')} | {e.get('quantity','')} |")
                    lines.append("")

            elif '第7章' in ch_num:
                problems = d.get('chapter7', {}).get('problems', [])
                if problems:
                    lines.append("### 发现的问题\n")
                    for p in problems[:10]:
                        lines.append(f"- **{p.get('title','')}**: {p.get('text','')}")
                    lines.append("")

            elif '第8章' in ch_num:
                if energy_data:
                    lines.append("### 审计结论\n")
                    lines.append(f"经对 {unit} 进行能源审计，主要结论如下：")
                    lines.append("")
                    lines.append("")

            # 子节
            for sec_num, sec_title in sub_sections:
                content = d.get('sections', {}).get(sec_num, '')
                if content:
                    lines.append(f"### {sec_num}  {sec_title}\n")
                    lines.append(f"{content}\n")

        return '\n'.join(lines)


# ============================================================
# 统一接口
# ============================================================

class ReportGenerator:
    """能源审计报告生成器（支持 Word 和 Markdown）"""

    def __init__(self, audit_type: str):
        self.audit_type = audit_type
        self.report_data: Dict = {}
        self.word_builder = WordReportBuilder(audit_type)
        self.md_builder = MarkdownReportBuilder(audit_type)

    def set_report_data(self, data: Dict):
        self.report_data = data
        self.word_builder.set_data(data)
        self.md_builder.set_data(data)

    def _query_basic_situation_fallback(self, unit_name: str) -> str:
        """2.1 兜底：项目数据无 basic_situation 时，从 PG ts_customer_info 反查。

        链路: find_project_by_name(unit_name) → customer_id →
              get_customer_info(customer_id) → basic_situation
        任何一步失败/无数据都返回 ''（不抛异常，避免生成报告被 DB 故障阻断）。
        """
        if not unit_name:
            return ''
        try:
            from tools.energy_audit import PgDataQuery
            pg = PgDataQuery()
            pg.connect()
            try:
                proj = pg.find_project_by_name(unit_name)
                customer_id = proj.get('customer_id') if proj else None
                if not customer_id:
                    return ''
                cust = pg.get_customer_info(customer_id=customer_id)
                if cust and cust[0].get('basic_situation'):
                    return str(cust[0]['basic_situation'])
                return ''
            finally:
                pg.disconnect()
        except Exception:
            return ''

    def load_from_project(self, project) -> Dict:
        """从 AuditProject 自动构建 report_data（字段溯源: project_data → 用户 → 默认）"""
        from tools.energy_audit.project_data import AuditProject
        if not isinstance(project, AuditProject):
            raise TypeError("需要 AuditProject 对象")

        b = project.base

        # ---- 照片按分类路由（对齐 photo_manager.PHOTO_REQUIREMENTS）----
        # 未分类照片（category=''）兜底进第2章；分类照片按章节/系统分发。
        img_by_cat: Dict[str, List] = {}
        for _img in project.images:
            _cat = getattr(_img, 'category', '') or ''
            img_by_cat.setdefault(_cat, []).append(_img)

        def _chapter_imgs(*cats: str, prefix: str = '图2') -> List[Dict]:
            """取指定分类的照片 → [{path, caption}]；caption 未填时按章节自动编号"""
            items = []
            for c in cats:
                items.extend(img_by_cat.get(c, []))
            return [{'path': _i.path, 'caption': _i.caption or f'{prefix}-{idx+1}'}
                    for idx, _i in enumerate(items)]

        # 第3章：管理信息 + 节能管理信息（ts_institution_energy_saving）合并，
        # 节能管理信息有记录时用真实数据生成，无记录时保留既有兜底文案。
        chapter3 = {
            'section_3_1': project.management.management_org or '',
            # 3.2 目标与方针：优先用 LLM 从制度文件提炼的正文，否则留空由兜底文案补
            'section_3_2': project.management.management_policy or '',
            'section_3_3': project.management.honors or '',
        }
        es_list = sorted((es for es in project.energy_saving if es is not None),
                         key=lambda es: es.statistical_year or 0, reverse=True)
        _ch3_imgs = []
        if es_list:
            es_secs = _energy_saving_chapter3_sections(es_list[0], b.unit_short or b.unit_name)
            for k, v in es_secs.items():
                # 3.2 若已从制度文件提炼到正文（management_policy），不再叠加兜底文案
                if k == 'section_3_2' and chapter3.get('section_3_2'):
                    continue
                chapter3[k] = '\n\n'.join(filter(None, [chapter3.get(k, ''), v]))
            # 管理文件 / 获奖证书附件图片（file_resolver 已下载到本地）
            for _p in (es_list[0].management_file_images + es_list[0].award_certificate_images):
                if _p and os.path.exists(_p):
                    _ch3_imgs.append({'path': _p, 'caption': f'图3-{len(_ch3_imgs)+1}'})
        # 第3章照片：合并数据模型分类图片（管理文件/荣誉）
        _ch3_imgs += _chapter_imgs('管理文件/荣誉', prefix='图3')
        if _ch3_imgs:
            chapter3['images'] = _ch3_imgs

        rd = {
            'tags': {
                'audit_type': b.unit_type or '公共机构',
                'institution_category': b.institution_category or '',
                'specific_type': b.specific_type or '',
            },
            'cover': {
                'title': f"{b.unit_name}能源审计报告",
                'audit_organization': b.auditor or '同方德诚（山东）科技股份公司',
                'report_date': b.report_date or datetime.now().strftime('%Y年%m月'),
            },
            'audit_info_tables': {
                # 能源审计机构信息表：机构名称/详细地址来自 ts_register_info（采集进 base），
                # 负责人/联系方式由用户提问提供（base.audit_org_contact/audit_org_phone）；
                # 表内 contact/mobile 仅作预填参考。缺失时保留空串，由 V1/V3 校验拦截提示。
                'institution': {
                    'name': b.audit_org_name or b.auditor,
                    'address': b.audit_org_address,
                    'contact': b.audit_org_contact,
                    'phone': b.audit_org_phone,
                },
                'team_members': [
                    {'role': m.role, 'name': m.name, 'education': m.education,
                     'certification': m.certification, 'major': m.major}
                    for m in project.audit_team
                ],
                'cooperation': [
                    {'role': c.role, 'dept': c.dept, 'name': c.name,
                     'gender': c.gender, 'position': c.position}
                    for c in project.cooperation
                ],
            },
            'chapter1': {
                'audited_unit_short': b.unit_short or b.unit_name,
                'address': b.address,
                'buildings': f"{len(project.buildings)}栋建筑",
                'audit_time': f"{b.audit_start}—{b.audit_end}" if b.audit_start and b.audit_end else '',
                'audit_period': b.audit_period or '',
                'base_period': b.base_period or '',
                'energy_types': [],
                'province': b.province or '山东',
            },
            'chapter2': {
                'unit_name': b.unit_name,
                # 2.1 直接取项目数据 basic_situation（数据收集阶段已从 ts_customer_info 解析，
                # 溯源 PG → Excel → Config）；为空则兜底查询 PG get_customer_info
                'section_2_1': b.basic_situation or self._query_basic_situation_fallback(b.unit_name),
                'building_area': b.building_area,
                'people_count': b.people_count,
                'beds_count': b.beds_count,
                'buildings': [{'name':bg.name or f'建筑{i+1}', 'address':bg.address or b.address,
                               'year':bg.year, 'function':bg.function, 'floors':bg.floors,
                               'area':bg.area, 'use_area':bg.use_area, 'structure':bg.structure,
                               'wall_body_material':bg.wall_body_material,
                               'window_type':bg.window_type, 'insulation':bg.insulation,
                               'roof_insulation':bg.roof_insulation,
                               'roof_insulation_material':bg.roof_insulation_material,
                               'sunshade_type':bg.sunshade_type, 'sunshade_material':bg.sunshade_material,
                               'orientation':bg.orientation, 'function_zoning':bg.function_zoning,
                               'height':bg.height,
                               'cooling_source':bg.cooling_source, 'heating_source':bg.heating_source,
                               'cooling_area':bg.cooling_area, 'heating_area':bg.heating_area,
                               'cooling_terminal':bg.cooling_terminal, 'heating_terminal':bg.heating_terminal,
                               'water_system':bg.water_system, 'fire_system':bg.fire_system,
                               'hot_water':bg.hot_water, 'monitoring':bg.monitoring,
                               'run_time':bg.run_time, 'storey_metrology':bg.storey_metrology,
                               'garage':bg.garage, 'garage_area':bg.garage_area}
                              for i, bg in enumerate(project.buildings)],
                # 建筑照片：分类（建筑外观/各建筑外观）+ 未分类兜底
                'images': _chapter_imgs('建筑外观', '各建筑外观', '', prefix='图2'),
            },
            'chapter3': chapter3,
            'chapter4': {
                'images': _chapter_imgs('计量器具', prefix='图4'),
                'section_4_1': f"计量体系按GB/T29149-2012划分为三级计量。⚠️ 请说明具体计量分级和监测系统建设情况。" if not project.metering.has_monitoring_system else f"计量体系按GB/T29149-2012划分为三级计量。",
                'section_4_2': f"电表{project.metering.electric_meters}块、水表{project.metering.water_meters}块。⚠️ 请确认具体数量和型号。" if not project.metering.electric_meters else f"电表{project.metering.electric_meters}块、水表{project.metering.water_meters}块、气表{project.metering.gas_meters}块。",
                'section_4_3': '⚠️ 请提供数据采集方式和统计制度说明。来源：被审计单位调研表。',
                'section_4_4': '⚠️ 请提供统计工作的成效和存在的问题。来源：现场调研和用户访谈。',
            },
            'chapter5': {
                'text': '',
            },
            'chapter6': {},
            'chapter7': {'images': _chapter_imgs('节能改造示意', prefix='图7')},
            'chapter8': {},
            'sections': {},
        }

        # 能耗数据 → chapter5（结构化，走 indicators 自动计算）
        if hasattr(project, 'energy_yearly') and project.energy_yearly:
            energy_data_list = []
            for ey in project.energy_yearly:
                energy_data_list.append({
                    'year': ey.year,
                    'electricity_kwh': ey.electricity_kwh,
                    'water_m3': ey.water_m3,
                    'natural_gas_m3': ey.natural_gas_m3,
                    'heating_energy_heat_gj': ey.heating_energy_heat_gj,
                    'petrol_kg': ey.petrol_kg,
                    'diesel_kg': ey.diesel_kg,
                    'electricity_cost_wan': ey.electricity_cost_wan,
                    'water_cost_wan': ey.water_cost_wan,
                    'heating_cost_wan': ey.heating_cost_wan,
                    'monthly_electricity_kwh': ey.monthly_electricity_kwh,
                    'monthly_water_m3': ey.monthly_water_m3,
                    'monthly_natural_gas_m3': ey.monthly_natural_gas_m3,
                })
            institution_type = 'medical' if '医疗' in (b.institution_category or '') else ('government' if '党政' in (b.institution_category or '') else 'education')
            # 统一使用 indicators 的机构类型映射
            from tools.energy_audit.indicators import institution_category_to_type
            institution_type = institution_category_to_type(b.institution_category)
            rd['chapter5'] = {
                'energy_data': energy_data_list,
                'unit_name': b.unit_short or b.unit_name,
                'building_area': b.building_area,
                'people_count': b.people_count,
                'beds_count': b.beds_count,
                'institution_type': institution_type,
            }
            # 优先使用 project.indicators（数据提取后已预计算）
            if getattr(project, 'indicators', None) and project.indicators.get('status') == 'ok':
                rd['chapter5']['indicators'] = project.indicators
            else:
                # 兜底：现场计算一次
                try:
                    from tools.energy_audit.indicators import compute_project_indicators
                    rd['chapter5']['indicators'] = compute_project_indicators(project)
                except Exception as e:
                    import logging as _logging
                    _logging.getLogger(__name__).warning(f"load_from_project 兜底计算指标失败: {e}")
            # 自动识别的能源类型同步到 chapter1
            counts = {}
            for d in energy_data_list:
                for k in ['electricity_kwh','water_m3','natural_gas_m3','petrol_kg','diesel_kg']:
                    if d.get(k, 0) and float(d.get(k, 0) or 0) > 0:
                        counts[k] = counts.get(k, 0) + 1
            energy_types = [k for k, v in counts.items() if v > 0]
            rd['chapter1']['energy_types'] = energy_types  # 供 chapter1 使用

        # 第5章账单照片（分类：能耗账单）
        rd['chapter5']['images'] = _chapter_imgs('能耗账单', prefix='图5')

        # 设备 → chapter6
        if project.equipment:
            eq_list = [{'name':eq.name,'category':eq.category,'spec':eq.spec,
                        'quantity':eq.quantity,'remark':eq.remark,
                        'independent_metering':eq.independent_metering,
                        'independent_metering_desc':eq.independent_metering_desc}
                       for eq in project.equipment]
            rd['chapter6'] = ch6 = rd.get('chapter6', {}) or {}
            ch6['_equipment'] = eq_list
            # 第6章设备照片：按系统分组；兼容旧版顶层 images_equipment 消费方（顺序: 制冷→照明→变配电→水泵→厨房）
            _eq_imgs = _chapter_imgs('制冷设备', '照明设备', '变压器/配电', '水泵/水箱', '厨房设备', prefix='图6')
            if _eq_imgs:
                rd['images_equipment'] = [_i['path'] for _i in _eq_imgs]
                ch6['images'] = {
                    'cooling': _chapter_imgs('制冷设备', prefix='图6'),
                    'lighting': _chapter_imgs('照明设备', prefix='图6'),
                    'transformer': _chapter_imgs('变压器/配电', prefix='图6'),
                    'water': _chapter_imgs('水泵/水箱', prefix='图6'),
                    'other_energy': _chapter_imgs('厨房设备', prefix='图6'),
                }

        self.set_report_data(rd)
        return rd

    def generate_word(self, output_path: str) -> str:
        """生成符合格式规范的 Word (.docx) 报告"""
        return self.word_builder.generate(output_path)

    def generate_markdown(self, output_path: str = None) -> str:
        """生成 Markdown 报告（快速预览）"""
        content = self.md_builder.generate()
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
        return content

    def generate(self, output_path: str) -> str:
        """自动根据扩展名选择格式"""
        if output_path.endswith('.docx'):
            return self.generate_word(output_path)
        else:
            return self.generate_markdown(output_path)


# ============================================================
# 使用示例
# ============================================================

def _indicator_eval(r) -> str:
    """从指标 dict 安全取评价结果；dict 缺失/含 error 或 benchmark 为 None 时返回 '—'。"""
    if isinstance(r, dict) and not r.get('error'):
        bm = r.get('benchmark')
        if isinstance(bm, dict) and bm.get('评价结果'):
            return bm['评价结果']
    return '—'


def _energy_saving_chapter3_sections(es, unit: str) -> dict:
    """从一条节能管理信息（ts_institution_energy_saving）生成第3章各节真实文案。

    返回 {section_3_2 / section_3_3 / section_3_4: 文本}，无对应数据的键不返回；
    load_from_project 会与 management 文案合并，无记录时保留既有兜底文案。
    """
    sections = {}

    # 3.2 能源管理制度
    if es.energy_management == 1:
        sections['section_3_2'] = (f"{unit}已建立能源管理制度，将节能管理纳入日常运营，"
                                   "通过制度建设、定期监督等方式落实节能责任。")
    elif es.energy_management == 0:
        sections['section_3_2'] = f"{unit}目前尚未建立完善的能源管理制度，节能管理仍有提升空间。"

    # 3.3 管理成效与问题
    parts_33 = []
    if es.has_awards == 1 and es.award_name:
        parts_33.append(f"{unit}节能工作取得成效，{es.award_name}。")
    if es.energy_pain_points:
        parts_33.append(f"目前能源利用方面存在的主要痛点：{es.energy_pain_points}。")
    if parts_33:
        sections['section_3_3'] = '\n'.join(parts_33)

    # 3.4 节能改造与管理措施
    parts_34 = []
    replaced = []
    if es.lighting_replacement == 1:
        replaced.append('照明灯具')
    if es.ac_replacement == 1:
        replaced.append('空调设备')
    if es.water_saving_fixture_replacement == 1:
        replaced.append('节水型卫生器具')
    if replaced:
        parts_34.append(f"{unit}已实施{'、'.join(replaced)}更换等节能改造措施。")
    if es.central_ac_control == 1:
        parts_34.append("中央空调系统已增加集中控制，以提升运行能效。")
    if es.other_measures:
        parts_34.append(f"其他节能改造措施：{es.other_measures}。")
    if es.third_party_system:
        parts_34.append(f"能源系统已由第三方托管运营：{es.third_party_system}。")
    if es.charging_pile == 1:
        cp = "单位已配置充电桩"
        if es.charging_settlement:
            cp += f"，结算方式为{es.charging_settlement}"
        if es.charging_installation:
            cp += f"，安装方式为{es.charging_installation}"
        parts_34.append(cp + "。")
    if es.third_party_outsource == 1:
        out = "用能系统已由第三方外包管理"
        if es.outsource_content:
            out += f"，内容包括{es.outsource_content}"
        if es.outsource_settlement:
            out += f"，结算方式为{es.outsource_settlement}"
        parts_34.append(out + "。")
    if parts_34:
        sections['section_3_4'] = '\n'.join(parts_34)

    return sections


def _iso_to_cn(date_str: str) -> str:
    """2022-01-01 → 2022年1月1日"""
    if not date_str or '年' in date_str:
        return date_str
    parts = date_str.split('-')
    if len(parts) >= 3:
        y, m, d = parts[0], str(int(parts[1])), str(int(parts[2]))
        return f"{y}年{m}月{d}日"
    return date_str


def _generate_single_energy_chart(yd_objects, et_key, label, unit, output_dir='./charts'):
    """生成单种能源的逐年柱状图"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        setup_chart_font(plt)
        years = [yd.year for yd in yd_objects]
        values = [float(getattr(yd, et_key, 0) or 0) for yd in yd_objects]
        if not any(v > 0 for v in values):
            return None
        fig, ax = plt.subplots(figsize=(5, 3))
        bars = ax.bar([str(y) for y in years], values, color='#4CAF50', width=0.4)
        for bar, val in zip(bars, values):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(values)*0.02,
                    f'{val:,.0f}', ha='center', va='bottom', fontsize=8)
        ax.set_ylabel(chart_text(f'{label}({unit})'), fontsize=9)
        ax.set_title(chart_text(f'逐年{label}趋势'), fontsize=11)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()
        os.makedirs(output_dir, exist_ok=True)
        path = os.path.join(output_dir, f'chart_{et_key}.png')
        fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        return path
    except Exception:
        return None

def _generate_monthly_bar_chart(energy_data_list, et_key, monthly_attr, label, unit, output_dir='./charts'):
    """生成逐月柱状图（年度对比）"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        setup_chart_font(plt)

        months = ['1月','2月','3月','4月','5月','6月','7月','8月','9月','10月','11月','12月']
        fig, ax = plt.subplots(figsize=(10, 5))
        has_data = False
        years_data = []

        for d in energy_data_list:
            monthly = d.get(monthly_attr)
            if monthly and isinstance(monthly, list) and len(monthly) == 12:
                values = [float(v or 0) for v in monthly]
                years_data.append((d.get('year', '?'), values))
                has_data = True

        if not has_data:
            return None

        n_years = len(years_data)
        bar_width = 0.25
        x = range(len(months))
        colors = ['#4CAF50', '#2196F3', '#FF9800']

        for i, (year_label, values) in enumerate(years_data):
            offset = (i - (n_years - 1) / 2) * bar_width
            bars = ax.bar([j + offset for j in x], values, bar_width,
                          label=f'{year_label}年', color=colors[i % 3], edgecolor='white')

        ax.set_xticks(x)
        ax.set_xticklabels(months, fontsize=8)
        ax.set_ylabel(chart_text(f'{label}({unit})'), fontsize=9)
        ax.set_title(chart_text(f'逐月{label}趋势（年度对比）'), fontsize=12, fontweight='bold')
        ax.legend(fontsize=9)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()
        os.makedirs(output_dir, exist_ok=True)
        path = os.path.join(output_dir, f'chart_{et_key}_monthly.png')
        fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        return path
    except Exception:
        return None

def _generate_cost_pie_chart(year: int, labels: list, values: list, output_dir='./charts'):
    """生成单年能源费用占比饼状图"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        setup_chart_font(plt)

        # 过滤0值
        filtered = [(l, v) for l, v in zip(labels, values) if v > 0]
        if not filtered:
            return None
        lbls, vals = zip(*filtered)

        fig, ax = plt.subplots(figsize=(5, 4))
        colors = ['#4CAF50','#2196F3','#FF9800','#F44336','#9C27B0']
        wedges, texts, autotexts = ax.pie(
            vals, labels=lbls, autopct='%1.1f%%',
            colors=colors[:len(vals)], startangle=90,
            textprops={'fontsize': 9}
        )
        ax.set_title(chart_text(f'{year}年能源费用占比'), fontsize=12)

        plt.tight_layout()
        os.makedirs(output_dir, exist_ok=True)
        path = os.path.join(output_dir, f'cost_pie_{year}.png')
        fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        return path
    except Exception:
        return None


if __name__ == "__main__":
    sample_data = {
        'cover': {
            'title': '山东省某政府机关能源审计报告',
            'audit_organization': '同方德诚科技有限公司',
            'report_date': '2026年6月',
        },
        'audit_info_tables': {
            'institution': {'name': '测试单位', 'address': '测试路1号', 'contact': '张三', 'phone': '13800000000'},
            'team_members': [{'role':'组长','name':'李四','education':'硕士','certification':'工程师','major':'暖通'}],
            'cooperation': [{'role':'配合','dept':'办公室','name':'王五','gender':'男','position':'主任'}],
        },
        'chapter1': {'audited_unit_short': '测试单位', 'audit_time': '2025年6月—2025年7月', 'audit_period': '2022年1月-2024年12月', 'base_period': '2021年1月-2023年12月', 'energy_types': ['electricity_kwh','water_m3'], 'province': '山东'},
        'chapter2': {'unit_name': '测试单位', 'building_area': 5000, 'people_count': 200, 'buildings': [{'name':'主楼','year':2020,'floors':'5层','structure':'框架'}]},
        'chapter5': {'energy_data': [{'year':2022,'electricity_kwh':500000,'water_m3':5000,'natural_gas_m3':3000},{'year':2023,'electricity_kwh':520000,'water_m3':5500,'natural_gas_m3':3200}], 'building_area': 5000, 'people_count': 200},
        'chapter6': {'_equipment': [{'name':'中央空调','category':'空调','spec':'30kW','quantity':3}]},
        'sections': {
            '1.1': '为贯彻落实《公共机构节能条例》，评估该单位能源使用状况，发现节能潜力。',
        }
    }

    gen = ReportGenerator("公共机构")
    gen.set_report_data(sample_data)

    # 生成 Word
    gen.generate_word("能源审计报告.docx")
    print("✓ Word报告已生成: 能源审计报告.docx")

    # 生成 Markdown
    md = gen.generate_markdown("能源审计报告.md")
    print(f"✓ Markdown报告已生成: 能源审计报告.md ({len(md)} 字符)")
