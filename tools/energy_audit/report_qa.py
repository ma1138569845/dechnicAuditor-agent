"""
报告自动质检

检查项:
  1. 章节完整性（1-8章顺序且不为空）
  2. 表格数量合理性
  3. 指标数值范围
  4. 残留占位符
  5. 章节编号连续性
"""

import re
from docx import Document


def check_report(doc_path: str) -> dict:
    """对生成的 Word 报告执行自动质检。返回 {ok, issues, warnings}"""
    doc = Document(doc_path)
    issues = []
    warnings = []

    paragraphs = [p.text.strip() for p in doc.paragraphs]
    tables = doc.tables

    # 1. 章节完整性
    expected_chapters = list(range(1, 9))
    found_chapters = []
    for i, p in enumerate(paragraphs):
        m = re.match(r'第(\d+)章', p)
        if m:
            found_chapters.append(int(m.group(1)))
    if found_chapters != expected_chapters:
        missing = set(expected_chapters) - set(found_chapters)
        extra = set(found_chapters) - set(expected_chapters)
        if missing: issues.append(f"缺少章节: 第{'、'.join(str(c) for c in sorted(missing))}章")
        if extra: issues.append(f"多余章节: 第{'、'.join(str(c) for c in sorted(extra))}章")
        if found_chapters != sorted(found_chapters):
            issues.append(f"章节顺序异常: {found_chapters}")
    else:
        warnings.append(f"✅ 章节顺序正确 (1→8)")

    # 2. 表格数量
    if len(tables) < 5:
        issues.append(f"表格过少 ({len(tables)}张)，应至少有5张（表1~表3 + 第2章建筑表 + 第5章指标表）")
    elif len(tables) >= 5:
        warnings.append(f"✅ 表格数量正常 ({len(tables)}张)")

    # 3. 残留占位符（段落 + 表格单元格都扫；表内【待补充】同样算残留）
    full_text = '\n'.join(paragraphs)
    try:
        _cell_texts = [cell.text for t in doc.tables for row in t.rows for cell in row.cells]
        if _cell_texts:
            full_text += '\n' + '\n'.join(_cell_texts)
    except Exception:
        pass
    placeholders = []
    for p in [
        ('【待补充】', '待补充'),
        ('【XX', '占位符'),
        ('YYYY年M月', '日期占位符'),
        ('待LLM生成', 'LLM占位符'),
    ]:
        if p[0] in full_text:
            count = full_text.count(p[0])
            placeholders.append(f"{p[1]} ({count}处)")
    if placeholders:
        issues.append(f"残留占位符: {'; '.join(placeholders)}")

    # 4. 指标数值检查
    tce_patterns = re.findall(r'(\d+\.?\d*)\s*tce', full_text)
    for val_str in tce_patterns:
        val = float(val_str)
        if val < 0: issues.append(f"负值能耗: {val} tce")
        if val > 100000: warnings.append(f"异常高能耗: {val} tce")

    kgce_patterns = re.findall(r'(\d+\.?\d*)\s*kgce', full_text)
    for val_str in kgce_patterns:
        val = float(val_str)
        if val < 1: issues.append(f"异常低单位面积能耗: {val} kgce/m²")
        if val > 1000: issues.append(f"异常高单位面积能耗: {val} kgce/m²")

    # 5. 覆盖/表1/表2/表3检查
    for tid in ['能源审计机构信息表', '能源审计组人员名单', '能源审计配合人员名单']:
        if tid not in full_text:
            issues.append(f"缺少必备表: {tid}")

    # 判定
    ok = len(issues) == 0

    return {
        'ok': ok,
        'issues': issues,
        'warnings': warnings,
        'chapters': found_chapters,
        'tables': len(tables),
        'paragraphs': len(paragraphs),
    }
