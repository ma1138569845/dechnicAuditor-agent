# -*- coding: utf-8 -*-
"""能源审计报告 MD -> DOCX 转换脚本（用户格式规范硬编码）
格式规范：H1宋体15pt居中加粗 / H2宋体14pt加粗 / H3宋体12pt加粗 /
正文宋体+Times New Roman 12pt两端对齐、首行缩进2字符、1.5倍行距 /
表格12pt居中、行高1.01cm(AT_LEAST)、垂直居中、表头加粗、Table Grid /
表注(>引用)灰色12pt无缩进 / 表题(**表X-X**)居中加粗 /
页脚居中页码域 / 目录TOC域(1-3级, 打开Word后Ctrl+A→F9更新)

用法:
  python md_to_docx_energy_audit.py <input.md> <output.docx> [单位名称] [审计期]

说明: md 开头结构约定 = 首个 "# " 封面标题 -> "> " 编制说明引用块 ->
首个表格为报告信息表(放封面) -> "# 目 录" -> "# 0.摘要" 起正文。
本机运行必须用 anaconda python: D:/develop/anaconda3/python.exe
"""
import re
import sys

from docx import Document
from docx.enum.table import (WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE,
                             WD_TABLE_ALIGNMENT)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SONG = '宋体'
TNR = 'Times New Roman'
GRAY = (0x59, 0x59, 0x59)
DEFAULT_UNIT = '××市第一中学'
DEFAULT_PERIOD = '2025年度'


def set_run_font(run, size=12, bold=False, color=None):
    run.font.name = TNR
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), TNR)
    rFonts.set(qn('w:hAnsi'), TNR)
    rFonts.set(qn('w:eastAsia'), SONG)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_rich(para, text, size=12, bold=False, color=None):
    """支持 **加粗** 内联标记"""
    for part in re.split(r'(\*\*.+?\*\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            r = para.add_run(part[2:-2])
            set_run_font(r, size, True, color)
        else:
            r = para.add_run(part)
            set_run_font(r, size, bold, color)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); run._r.append(f1)
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'; run._r.append(instr)
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate'); run._r.append(f2)
    t = OxmlElement('w:t'); t.text = '（打开文档后在Word中按 Ctrl+A → F9 更新目录）'; run._r.append(t)
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end'); run._r.append(f3)
    set_run_font(run, 12)


def add_page_field(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); run._r.append(f1)
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'; run._r.append(instr)
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end'); run._r.append(f2)
    set_run_font(run, 10.5)


def add_heading(doc, text, level):
    sizes = {1: 15, 2: 14, 3: 12}
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    add_rich(p, text, size=sizes[level], bold=True)
    pPr = p._element.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), str(level - 1))
    pPr.append(ol)
    return p


def add_body(doc, text, indent=True, color=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Pt(24)  # 12pt x 2字符
    add_rich(p, text, color=color)
    return p


def add_table(doc, rows):
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        row_obj = table.rows[ri]
        row_obj.height = Cm(1.01)
        row_obj.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for ci in range(n_cols):
            cell = row_obj.cells[ci]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[ci] if ci < len(row) else ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = 1.0
            add_rich(para, text, bold=(ri == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def parse_table(lines, start):
    rows = []
    j = start
    while j < len(lines):
        line = lines[j].strip()
        if not line.startswith('|'):
            break
        cells = [c.strip() for c in line.strip('|').split('|')]
        if all(re.fullmatch(r'[\s:\-]+', c) for c in cells):
            j += 1
            continue
        rows.append(cells)
        j += 1
    return rows, j


def main():
    args = sys.argv[1:]
    if len(args) < 2:
        print(__doc__)
        sys.exit(1)
    md_path, out_path = args[0], args[1]
    unit = args[2] if len(args) > 2 else DEFAULT_UNIT
    period = args[3] if len(args) > 3 else DEFAULT_PERIOD
    # 正式模式: 传入第5个参数(审计机构名)即视为正式报告, 封面不显示"模板示例"
    real_mode = len(args) > 4
    audit_org = args[4] if real_mode else '同方德诚科技有限公司'
    report_date = args[5] if len(args) > 5 else '二〇二六年×月'

    with open(md_path, encoding='utf-8') as f:
        lines = [l.rstrip('\n') for l in f.readlines()]

    # 阶段1: 封面标题 / 编制说明 / 封面信息表
    cover_title = None
    usage = []
    cover_table = None
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('# ') and cover_title is None:
            cover_title = line[2:].strip()
            i += 1
            continue
        if line.startswith('> '):
            usage.append(line[2:].strip())
            i += 1
            continue
        if line.startswith('|'):
            rows, i = parse_table(lines, i)
            cover_table = rows
            break  # 封面信息表之后不再收集编制说明（避免误收正文表注）
        i += 1

    # 阶段2: 定位"# 目 录"及目录前内容（封面信息页）与正文起点
    toc_idx = -1
    body_start = 0
    for idx in range(len(lines)):
        if lines[idx].startswith('# ') and lines[idx][2:].strip() == '目 录':
            toc_idx = idx
            j = idx + 1
            while j < len(lines) and not lines[j].startswith('# '):
                j += 1
            body_start = j
            break
    # 目录前信息页内容 = cover_table 解析结束位置 到 "# 目 录" 之间
    pre_toc = lines[i:toc_idx] if toc_idx > i else []

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.0)
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer_p)

    # 封面
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(unit); set_run_font(r, 22, True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('能源审计报告'); set_run_font(r, 26, True)
    if not real_mode:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('（模板示例，数据均为虚拟编制）'); set_run_font(r, 14, False, GRAY)
    for _ in range(6):
        doc.add_paragraph()
    if real_mode:
        # 正式报告封面: 机构名 + 报告日期（与参考报告一致，无标签前缀）
        for label in (audit_org, report_date):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(label); set_run_font(r, 16)
    else:
        for label in ('审计机构：' + audit_org,
                      '审计期：' + period,
                      '报告日期：' + report_date):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(label); set_run_font(r, 16)
    if cover_table:
        doc.add_paragraph()
        add_table(doc, cover_table)
    doc.add_page_break()

    # 目录前信息页（如"能源审计机构信息表"）：渲染 # 标题、表格、**加粗**行、普通行
    if pre_toc:
        pt = 0
        while pt < len(pre_toc):
            s = pre_toc[pt].strip()
            if not s:
                pt += 1
                continue
            if s.startswith('# '):
                add_heading(doc, s[2:].strip(), 1)
                pt += 1
            elif s.startswith('|'):
                rows, nj = parse_table(pre_toc, pt)
                add_table(doc, rows)
                pt = nj
            elif s.startswith('**') and s.endswith('**'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_rich(p, s, bold=True)
                pt += 1
            else:
                add_body(doc, s)
                pt += 1
        doc.add_page_break()

    # 目录页
    add_heading(doc, '目 录', 1)
    add_toc_field(doc.add_paragraph())
    doc.add_page_break()

    # 编制说明页（仅当 md 开头提供了 > 引用块时输出）
    if usage:
        add_heading(doc, '编制说明', 1)
        for u in usage:
            add_body(doc, u, indent=False)
        doc.add_page_break()

    # 正文
    j = body_start
    while j < len(lines):
        line = lines[j].strip()
        if not line:
            j += 1
            continue
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1)
            j += 1
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
            j += 1
        elif line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
            j += 1
        elif line.startswith('|'):
            rows, nj = parse_table(lines, j)
            add_table(doc, rows)
            j = nj
        elif line.startswith('> '):
            add_body(doc, line[2:].strip(), indent=False, color=GRAY)
            j += 1
        elif re.match(r'^---+\s*$', line):
            j += 1
        elif re.match(r'^\*\*表[\d\-]+[^*]*\*\*$', line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            add_rich(p, line, bold=True)
            j += 1
        else:
            add_body(doc, line)
            j += 1

    doc.save(out_path)
    print('OK ->', out_path)


if __name__ == '__main__':
    main()
