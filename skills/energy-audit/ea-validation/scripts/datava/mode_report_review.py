"""V3 REPORT_REVIEW — 报告生成后审查。

三条主线：
    跨章数据一致性  第2章面积 vs 第5章面积、第4章能耗 vs 第5章能耗、与 data.json 对齐
    章节完整性      1~8章齐备、1.6省级规章≥3条、第6章动态H3、第8章指标汇总表
    格式规范        字体/字号/行距/缩进/对齐/表格行高（对齐 report_generator.FormatSpec）

产出：report_review.json + report_review.txt
存在 P0 时进程以退出码 2 结束，上游据此 kanban_block。
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from .common import (
    SEV_P0,
    SEV_P1,
    SEV_P2,
    Finding,
    ReviewResult,
    build_result,
    fmt_num,
    project_data_path,
    project_dir,
    read_json,
    safe_float,
    with_artifacts,
    write_json,
    write_text,
)

MODE = "REPORT_REVIEW"

EXPECTED_CHAPTERS: Tuple[int, ...] = tuple(range(1, 9))
MIN_PROVINCE_RULES = 3
MIN_CHAPTER6_H3 = 3
MIN_TABLES = 5
AREA_TOLERANCE_PCT = 1.0
ENERGY_TOLERANCE_PCT = 1.0
FORMAT_SAMPLE_LIMIT = 3
BODY_MIN_CHARS = 24  # 短于此长度的段落视为表题/图注，不参与正文格式判定

REQUIRED_TABLES: Tuple[str, ...] = (
    "能源审计机构信息表",
    "能源审计组人员名单",
    "能源审计配合人员名单",
)

PLACEHOLDERS: Tuple[Tuple[str, str], ...] = (
    ("【待补充】", "待补充占位符"),
    ("【XX", "XX 占位符"),
    ("YYYY年M月", "日期占位符"),
    ("待LLM生成", "LLM 占位符"),
    ("TODO", "TODO 标记"),
)

# 对齐 tools/energy_audit/report_generator.py 的 FormatSpec
FMT_BODY = {"font": "宋体", "size": 12.0, "line_spacing": 1.5, "indent_pt": 24.0, "align": 3}
FMT_H1 = {"font": "宋体", "size": 15.0, "bold": True, "align": 1}
FMT_H2 = {"font": "宋体", "size": 14.0, "bold": True, "align": 0}
FMT_H3 = {"font": "宋体", "size": 12.0, "bold": True, "align": 0}
TABLE_ROW_HEIGHT_EMU = 363600  # Cm(1.01)
TABLE_ROW_HEIGHT_TOL = 12000

_CN_DIGITS = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}

RE_CHAPTER = re.compile(r"^第\s*([0-9]+|[一二三四五六七八九])\s*章")
RE_H2 = re.compile(r"^(\d+)\.(\d+)(?![.\d])")
RE_H3 = re.compile(r"^(\d+)\.(\d+)\.(\d+)")
RE_AREA = re.compile(
    r"(?:总?建筑面积|建筑总面积)[^0-9]{0,16}([0-9][0-9,]*(?:\.[0-9]+)?)\s*(万)?\s*(?:m2|m²|平方米|㎡)"
)
RE_TCE = re.compile(r"([0-9][0-9,]*(?:\.[0-9]+)?)\s*(?:tce|吨标(?:准)?煤)")
RE_LIST_ITEM = re.compile(r"^(?:[（(]?\d+[)）\.、]|[·•\-—])")


# ================================================================
#  docx 结构解析
# ================================================================

@dataclass(frozen=True)
class Block:
    kind: str  # 'p' | 'tbl'
    index: int  # 文档内块序号
    text: str = ""
    obj: Any = None
    chapter: int = 0


def _chapter_number(text: str) -> int:
    match = RE_CHAPTER.match(text)
    if not match:
        return 0
    token = match.group(1)
    return int(token) if token.isdigit() else _CN_DIGITS.get(token, 0)


def parse_blocks(document: Any) -> List[Block]:
    """按文档顺序遍历段落与表格，并标注所属章节。"""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    paragraph_map = {p._element: p for p in document.paragraphs}
    table_map = {t._element: t for t in document.tables}

    blocks: List[Block] = []
    chapter = 0
    for index, child in enumerate(body.iterchildren()):
        if child.tag == qn("w:p"):
            paragraph = paragraph_map.get(child) or Paragraph(child, document)
            text = paragraph.text.strip()
            found = _chapter_number(text)
            if found:
                chapter = found
            blocks.append(Block("p", index, text, paragraph, chapter))
        elif child.tag == qn("w:tbl"):
            table = table_map.get(child) or Table(child, document)
            blocks.append(Block("tbl", index, "", table, chapter))
    return blocks


def chapter_text(blocks: Sequence[Block], chapter: int) -> str:
    return "\n".join(b.text for b in blocks if b.kind == "p" and b.chapter == chapter and b.text)


def chapter_tables(blocks: Sequence[Block], chapter: int) -> int:
    return sum(1 for b in blocks if b.kind == "tbl" and b.chapter == chapter)


# ================================================================
#  跨章数据一致性
# ================================================================

def _normalize_area(value: str, wan: Optional[str]) -> float:
    number = safe_float(value.replace(",", ""))
    return number * 10000.0 if wan else number


def extract_areas(blocks: Sequence[Block]) -> Dict[int, List[float]]:
    """按章节提取建筑面积陈述值（万 m² 已换算为 m²）。"""
    found: Dict[int, List[float]] = {}
    for block in blocks:
        if block.kind != "p" or not block.text:
            continue
        for value, wan in RE_AREA.findall(block.text):
            found.setdefault(block.chapter, []).append(_normalize_area(value, wan))
    return found


def check_area_consistency(
    areas: Dict[int, List[float]], truth_area: float
) -> List[Finding]:
    findings: List[Finding] = []
    flat = [(ch, v) for ch, values in areas.items() for v in values if v > 0]
    if not flat:
        return [
            Finding(
                code="V3.CROSS.AREA_MISSING",
                category="跨章一致性",
                severity=SEV_P1,
                title="全文未检出建筑面积陈述",
                location="第2章 / 第5章",
                suggestion="确认第2章建筑概况与第5章指标计算说明中写明总建筑面积",
            )
        ]

    values = [v for _, v in flat]
    low, high = min(values), max(values)
    if low > 0 and (high - low) / low * 100.0 > AREA_TOLERANCE_PCT:
        detail = "; ".join(f"第{ch}章 {fmt_num(v)} m²" for ch, v in sorted(flat))
        findings.append(
            Finding(
                code="V3.CROSS.AREA_MISMATCH",
                category="跨章一致性",
                severity=SEV_P0,
                title=f"各章建筑面积不一致（极差 {(high - low) / low * 100.0:.2f}%）",
                detail=detail,
                location="跨章",
                expected="全文口径一致",
                actual=f"{fmt_num(low)} ~ {fmt_num(high)} m²",
                suggestion="统一以 data.json base.building_area 为准重新生成受影响章节",
            )
        )

    if truth_area > 0:
        for chapter, value in sorted(flat):
            deviation = abs(value - truth_area) / truth_area * 100.0
            if deviation > AREA_TOLERANCE_PCT:
                findings.append(
                    Finding(
                        code="V3.CROSS.AREA_VS_SOURCE",
                        category="跨章一致性",
                        severity=SEV_P0,
                        title=f"第{chapter}章建筑面积与项目数据不符",
                        location=f"第{chapter}章",
                        expected=f"{fmt_num(truth_area)} m²（data.json）",
                        actual=f"{fmt_num(value)} m²",
                        suggestion="修正报告数据源绑定，勿在文本中硬写面积",
                    )
                )
                break  # 同一偏差不重复报，定位一处即可
    return findings


def check_energy_consistency(blocks: Sequence[Block]) -> List[Finding]:
    """第4章综合能耗与第5章综合能耗的最大陈述值比对。"""
    def max_tce(chapter: int) -> Optional[float]:
        values = [
            safe_float(v.replace(",", ""))
            for v in RE_TCE.findall(chapter_text(blocks, chapter))
        ]
        positives = [v for v in values if v > 0]
        return max(positives) if positives else None

    ch4, ch5 = max_tce(4), max_tce(5)
    if ch4 is None or ch5 is None:
        return []
    deviation = abs(ch4 - ch5) / max(ch4, ch5) * 100.0
    if deviation <= ENERGY_TOLERANCE_PCT:
        return []
    return [
        Finding(
            code="V3.CROSS.ENERGY_MISMATCH",
            category="跨章一致性",
            severity=SEV_P1,
            title=f"第4章与第5章综合能耗最大值偏差 {deviation:.2f}%",
            detail=f"第4章 {fmt_num(ch4)} tce，第5章 {fmt_num(ch5)} tce",
            location="第4章 vs 第5章",
            expected=f"偏差 ≤ {ENERGY_TOLERANCE_PCT:.0f}%",
            actual=f"{deviation:.2f}%",
            suggestion="核对两章是否同口径（是否含交通能耗/供暖能耗），统一后重新生成",
        )
    ]


def check_year_coverage(blocks: Sequence[Block], years: Sequence[int]) -> List[Finding]:
    text = chapter_text(blocks, 5) + "\n" + chapter_text(blocks, 4)
    missing = [year for year in years if str(year) not in text]
    if not missing:
        return []
    return [
        Finding(
            code="V3.CROSS.YEAR_MISSING",
            category="跨章一致性",
            severity=SEV_P1,
            title=f"审计年份未在第4/5章出现：{missing}",
            location="第4章 / 第5章",
            expected=f"覆盖 {list(years)}",
            actual="缺 " + ", ".join(str(y) for y in missing),
            suggestion="核对能耗表与图表是否遗漏该年数据",
        )
    ]


# ================================================================
#  章节完整性
# ================================================================

def check_chapters(blocks: Sequence[Block]) -> List[Finding]:
    found: List[int] = []
    for block in blocks:
        if block.kind != "p":
            continue
        number = _chapter_number(block.text)
        if number and number not in found:
            found.append(number)

    findings: List[Finding] = []
    missing = [c for c in EXPECTED_CHAPTERS if c not in found]
    if missing:
        findings.append(
            Finding(
                code="V3.STRUCT.CHAPTER_MISSING",
                category="章节完整性",
                severity=SEV_P0,
                title=f"缺少章节：第{'、'.join(str(c) for c in missing)}章",
                expected="第1~8章齐备",
                actual=f"实检出 {found}",
                suggestion="补齐缺失章节后重新生成报告",
            )
        )
    elif found != sorted(found):
        findings.append(
            Finding(
                code="V3.STRUCT.CHAPTER_ORDER",
                category="章节完整性",
                severity=SEV_P0,
                title="章节顺序异常",
                expected="1→8 递增",
                actual=str(found),
                suggestion="检查生成器章节装配顺序",
            )
        )

    empty = [
        c for c in EXPECTED_CHAPTERS
        if c in found and len(chapter_text(blocks, c)) < 80 and chapter_tables(blocks, c) == 0
    ]
    if empty:
        findings.append(
            Finding(
                code="V3.STRUCT.CHAPTER_EMPTY",
                category="章节完整性",
                severity=SEV_P0,
                title=f"章节内容近乎为空：第{'、'.join(str(c) for c in empty)}章",
                expected="正文 ≥ 80 字或含表格",
                actual="仅有标题",
                suggestion="补齐章节正文",
            )
        )
    return findings


def check_province_rules(blocks: Sequence[Block]) -> List[Finding]:
    """1.6 省级规章条目数（≥3 条）。"""
    section: List[str] = []
    collecting = False
    for block in blocks:
        if block.kind != "p" or not block.text:
            continue
        if RE_H2.match(block.text):
            head = RE_H2.match(block.text)
            collecting = head.group(1) == "1" and head.group(2) == "6"
            continue
        if block.chapter >= 2:
            collecting = False
        if collecting:
            section.append(block.text)

    count = sum(1 for line in section if "《" in line or RE_LIST_ITEM.match(line))
    if count >= MIN_PROVINCE_RULES:
        return []
    return [
        Finding(
            code="V3.STRUCT.PROVINCE_RULES",
            category="章节完整性",
            severity=SEV_P1,
            title=f"1.6 省级规章条目不足（{count} 条）",
            location="第1章 1.6",
            expected=f"≥ {MIN_PROVINCE_RULES} 条",
            actual=f"{count} 条",
            suggestion="用 web_search 核实省级节能规章后补录，禁止字符串替换套用他省规章",
        )
    ]


def check_chapter6_h3(blocks: Sequence[Block]) -> List[Finding]:
    count = sum(
        1
        for block in blocks
        if block.kind == "p" and block.chapter == 6 and RE_H3.match(block.text)
    )
    if count >= MIN_CHAPTER6_H3:
        return []
    return [
        Finding(
            code="V3.STRUCT.CHAPTER6_H3",
            category="章节完整性",
            severity=SEV_P1,
            title=f"第6章动态三级标题不足（{count} 个）",
            detail="第6章应按实际用能系统（空调/照明/电梯/热水/变配电…）动态展开 H3",
            location="第6章",
            expected=f"≥ {MIN_CHAPTER6_H3} 个 6.x.y 小节",
            actual=f"{count} 个",
            suggestion="按设备清单 category 动态生成系统小节",
        )
    ]


def check_chapter8_summary(blocks: Sequence[Block]) -> List[Finding]:
    if chapter_tables(blocks, 8) >= 1:
        return []
    return [
        Finding(
            code="V3.STRUCT.CHAPTER8_TABLE",
            category="章节完整性",
            severity=SEV_P1,
            title="第8章缺指标汇总表",
            location="第8章",
            expected="≥ 1 张汇总表",
            actual="0 张",
            suggestion="补第8章指标汇总表（各年指标 + 对标结论）",
        )
    ]


def check_tables_and_placeholders(blocks: Sequence[Block]) -> List[Finding]:
    findings: List[Finding] = []
    total_tables = sum(1 for b in blocks if b.kind == "tbl")
    if total_tables < MIN_TABLES:
        findings.append(
            Finding(
                code="V3.STRUCT.TABLE_COUNT",
                category="章节完整性",
                severity=SEV_P1,
                title=f"表格数量偏少（{total_tables} 张）",
                expected=f"≥ {MIN_TABLES} 张",
                actual=f"{total_tables} 张",
                suggestion="核对表1~表3、第2章建筑表、第5章指标表是否齐备",
            )
        )

    full_text = "\n".join(b.text for b in blocks if b.kind == "p")
    # 表格单元格文本并入占位符扫描（表内【待补充】也必须 P0 阻塞；
    # 原实现只扫段落，审计基本信息表/设备表内的占位符必然漏检）
    _cell_texts: List[str] = []
    for b in blocks:
        if b.kind != "tbl":
            continue
        try:
            for row in b.obj.rows:
                for cell in row.cells:
                    _cell_texts.append(cell.text)
        except Exception:
            continue
    if _cell_texts:
        full_text = full_text + "\n" + "\n".join(_cell_texts)
    for token, label in PLACEHOLDERS:
        occurrences = full_text.count(token)
        if occurrences:
            findings.append(
                Finding(
                    code="V3.STRUCT.PLACEHOLDER",
                    category="章节完整性",
                    severity=SEV_P0,
                    title=f"残留{label} {occurrences} 处",
                    detail=f"标记='{token}'",
                    expected="0 处",
                    actual=f"{occurrences} 处",
                    suggestion="补齐对应数据或删除占位段落后重新生成",
                )
            )

    missing_tables = [name for name in REQUIRED_TABLES if name not in full_text]
    if missing_tables:
        findings.append(
            Finding(
                code="V3.STRUCT.REQUIRED_TABLE",
                category="章节完整性",
                severity=SEV_P1,
                title=f"缺必备表：{'、'.join(missing_tables)}",
                location="审计基本信息",
                suggestion="补齐审计机构信息表/审计组人员名单/配合人员名单",
            )
        )
    return findings


# ================================================================
#  格式规范
# ================================================================

@dataclass
class Violation:
    """格式违规累加器（内部可变，仅用于统计，不对外暴露）。"""

    label: str
    expected: str
    samples: List[str] = field(default_factory=list)
    count: int = 0

    def add(self, sample: str) -> None:
        self.count += 1
        if len(self.samples) < FORMAT_SAMPLE_LIMIT:
            self.samples.append(sample)


def _east_asia_font(run: Any) -> Optional[str]:
    from docx.oxml.ns import qn

    properties = run._element.rPr
    if properties is None or properties.rFonts is None:
        return None
    return properties.rFonts.get(qn("w:eastAsia"))


def _paragraph_font(paragraph: Any) -> Tuple[Optional[str], Optional[float], Optional[bool]]:
    """取首个非空 run 的显式字体设置。None 表示继承样式，不判违规。"""
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        name = _east_asia_font(run) or run.font.name
        size = run.font.size.pt if run.font.size is not None else None
        return name, size, run.font.bold
    return None, None, None


def _classify(text: str) -> str:
    if _chapter_number(text):
        return "h1"
    if RE_H3.match(text):
        return "h3"
    if RE_H2.match(text):
        return "h2"
    return "body" if len(text) >= BODY_MIN_CHARS else "skip"


def check_format(blocks: Sequence[Block]) -> List[Finding]:
    specs = {"h1": FMT_H1, "h2": FMT_H2, "h3": FMT_H3, "body": FMT_BODY}
    violations: Dict[str, Violation] = {}

    def record(key: str, label: str, expected: str, sample: str) -> None:
        violations.setdefault(key, Violation(label, expected)).add(sample)

    for block in blocks:
        if block.kind != "p" or not block.text:
            continue
        kind = _classify(block.text)
        if kind == "skip":
            continue
        spec = specs[kind]
        excerpt = block.text[:24]
        name, size, bold = _paragraph_font(block.obj)

        if name is not None and name != spec["font"]:
            record(f"{kind}.font", f"{kind} 中文字体", str(spec["font"]), f"{excerpt} → {name}")
        if size is not None and abs(size - float(spec["size"])) > 0.51:
            record(f"{kind}.size", f"{kind} 字号", f"{spec['size']}pt", f"{excerpt} → {size}pt")
        if spec.get("bold") and bold is False:
            record(f"{kind}.bold", f"{kind} 加粗", "加粗", excerpt)

        paragraph_format = block.obj.paragraph_format
        alignment = paragraph_format.alignment
        if alignment is not None and int(alignment) != int(spec["align"]):
            record(
                f"{kind}.align",
                f"{kind} 对齐",
                str(spec["align"]),
                f"{excerpt} → {int(alignment)}",
            )

        if kind == "body":
            spacing = paragraph_format.line_spacing
            if spacing is not None and abs(float(spacing) - FMT_BODY["line_spacing"]) > 0.05:
                record("body.spacing", "正文行距", "1.5 倍", f"{excerpt} → {spacing}")
            indent = paragraph_format.first_line_indent
            if indent is not None and abs(indent.pt - FMT_BODY["indent_pt"]) > 2.0:
                record(
                    "body.indent",
                    "正文首行缩进",
                    "2 字符（24pt）",
                    f"{excerpt} → {indent.pt:.1f}pt",
                )

    bad_rows = 0
    total_rows = 0
    for block in blocks:
        if block.kind != "tbl":
            continue
        for row in block.obj.rows:
            total_rows += 1
            height = row.height
            if height is None or abs(int(height) - TABLE_ROW_HEIGHT_EMU) > TABLE_ROW_HEIGHT_TOL:
                bad_rows += 1
    if total_rows and bad_rows:
        record(
            "table.row_height",
            "表格行高",
            "1.01cm 固定行高",
            f"{bad_rows}/{total_rows} 行未设定或不符",
        )

    return [
        Finding(
            code=f"V3.FORMAT.{key.upper().replace('.', '_')}",
            category="格式规范",
            severity=SEV_P2,
            title=f"{violation.label}不符规范（{violation.count} 处）",
            detail="; ".join(violation.samples),
            expected=violation.expected,
            actual=f"{violation.count} 处偏离",
            suggestion="对齐 report_generator.FormatSpec 后重新生成，勿手工改样式",
        )
        for key, violation in sorted(violations.items())
    ]


# ================================================================
#  报告定位
# ================================================================

def locate_report(
    project: str, report: Optional[str] = None, output_dir: Optional[str] = None
) -> Tuple[Optional[Path], str]:
    if report:
        path = Path(report).expanduser()
        return (path, "") if path.is_file() else (None, f"报告文件不存在: {path}")

    search_dirs = [project_dir(project, output_dir), Path.cwd()]
    candidates: List[Path] = []
    for directory in search_dirs:
        if directory.is_dir():
            candidates += [
                p for p in directory.glob("*.docx") if not p.name.startswith("~$")
            ]
    if not candidates:
        return None, (
            "未找到 .docx 报告。用 --report <路径> 显式指定，"
            f"或把报告放到 {search_dirs[0]}"
        )
    preferred = [p for p in candidates if "能源审计报告" in p.name] or candidates
    return max(preferred, key=lambda p: p.stat().st_mtime), ""


# ================================================================
#  主流程
# ================================================================

def run(
    project: str,
    *,
    report: Optional[str] = None,
    output_dir: Optional[str] = None,
) -> ReviewResult:
    try:
        from docx import Document
    except ImportError as exc:
        return build_result(MODE, project, [], error=f"python-docx 不可用: {exc}")

    path, error = locate_report(project, report, output_dir)
    if error or path is None:
        return build_result(MODE, project, [], error=error or "报告定位失败")

    try:
        document = Document(str(path))
        blocks = parse_blocks(document)
    except Exception as exc:  # noqa: BLE001 — docx 解析异常形态多样，统一降级
        return build_result(MODE, project, [], error=f"报告解析失败 ({path.name}): {exc}")

    raw = read_json(project_data_path(project))
    base = (raw or {}).get("base") or {}
    truth_area = safe_float(base.get("building_area"))
    years = sorted(
        int(safe_float(row.get("year")))
        for row in ((raw or {}).get("energy_yearly") or [])
        if safe_float(row.get("year")) > 0
    )

    findings: List[Finding] = []
    findings += check_chapters(blocks)
    findings += check_tables_and_placeholders(blocks)
    findings += check_province_rules(blocks)
    findings += check_chapter6_h3(blocks)
    findings += check_chapter8_summary(blocks)
    findings += check_area_consistency(extract_areas(blocks), truth_area)
    findings += check_energy_consistency(blocks)
    if years:
        findings += check_year_coverage(blocks, years)
    findings += check_format(blocks)

    result = build_result(
        MODE,
        project,
        findings,
        inputs={
            "报告": str(path),
            "data.json": "已读取" if raw else "缺失（跳过与数据源比对）",
        },
        extra={
            "paragraphs": sum(1 for b in blocks if b.kind == "p"),
            "tables": sum(1 for b in blocks if b.kind == "tbl"),
            "chapters": sorted({b.chapter for b in blocks if b.chapter}),
            "audit_years": years,
        },
    )

    save_dir = project_dir(project, output_dir)
    review_path = Path(write_json(save_dir / "report_review.json", result.to_dict()))
    result = with_artifacts(result, {"report_review.json": str(review_path)})
    result = with_artifacts(
        result,
        {
            "report_review.txt": write_text(
                save_dir / "report_review.txt", result.render_text()
            )
        },
    )
    write_json(review_path, result.to_dict())
    return result
