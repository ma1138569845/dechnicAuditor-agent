"""Knowledge Base management API for Hermes Web UI.

This module provides a unified knowledge-base (KB) management surface backed by:
- Disk storage under HERMES_KNOWLEDGE_ROOT (per-KB subdirectories)
- SQLite metadata (knowledge_bases, folders, documents, chunks, jobs)
- Qdrant vector collections for RAG
- Energy-audit specific parsing/chunking via energy_audit_importer

Phase 1 (MVP) scope:
- KB CRUD with a default "energy_audit" KB that preserves legacy behaviour
- Persistent folder hierarchy
- Document upload / list / preview / delete
- Background vectorization with real job tracking
- Chunk-level inspection (Phase 1.5)
- Rebuild / reindex (Phase 1.5)

The legacy path-based endpoints in routes.py continue to operate against the
"default" KB so existing clients keep working.
"""
from __future__ import annotations

import atexit
import base64
import hashlib
import json
import logging
import mimetypes
import os
import re
import shutil
import sqlite3
import sys
import threading
import time
import traceback
import uuid
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Shared resolver: config.yaml knowledge_base: → env → defaults.
from rag.config import (
    deepseek_api_base as _deepseek_api_base_fn,
    deepseek_api_key as _deepseek_api_key_fn,
    knowledge_root as _knowledge_root_fn,
    qdrant_grpc_port as _qdrant_grpc_port_fn,
    qdrant_host as _qdrant_host_fn,
    summary_model as _summary_model_fn,
    summary_provider as _summary_provider_fn,
    wiki_vault as _wiki_vault_fn,
)

try:
    from hermes_constants import get_hermes_home as _get_hermes_home
except ImportError:
    from pathlib import Path as _Path

    def _get_hermes_home() -> _Path:
        return _Path.home() / ".hermes"

# Module-level aliases kept so tests can monkeypatch them. Values are a snapshot
# at import; _qdrant_client() / LLM helpers re-read rag.config unless patched.
_RAG_DATA_ROOT = _knowledge_root_fn()
_DEFAULT_KNOWLEDGE_ROOT = _RAG_DATA_ROOT
_QDRANT_HOST = _qdrant_host_fn()
_QDRANT_PORT = _qdrant_grpc_port_fn()
_DEFAULT_WIKI_VAULT = _wiki_vault_fn()
_DASHSCOPE_API_KEY = os.environ.get("DASHSCOPE_API_KEY", "")
_DEEPSEEK_API_KEY = _deepseek_api_key_fn()
_DEEPSEEK_API_BASE = _deepseek_api_base_fn()
_SUMMARY_PROVIDER = _summary_provider_fn()
_SUMMARY_MODEL = _summary_model_fn()

# Legacy alias kept for backwards compatibility
_KNOWLEDGE_BASE_ROOT = _DEFAULT_KNOWLEDGE_ROOT

# System knowledge bases that are created by default and cannot be deleted.
_SYSTEM_KB_IDS = frozenset({
    "energy_audit_reports",
    "energy_quota_standards",
    "energy_audit_technical_guidelines",
})

_DEFAULT_KBS = [
    {
        "id": "energy_audit_reports",
        "name": "能源审计报告",
        "description": "存储能源审计报告文档，支持语义检索和RAG问答。",
        "kb_type": "energy_audit",
        "qdrant_collection": "energy_audit_reports",
        "embedding_model": "dashscope/text-embedding-v3",
        "chunking_config": {"strategy": "three_level", "max_chunk_size": 512, "overlap": 64},
        "indexing_strategy": {"vector": True, "summary": True, "graph": True, "wiki": True},
    },
    {
        "id": "energy_quota_standards",
        "name": "能源定额标准",
        "description": "存储能源消耗定额标准文档，支持指标对比和合规检查。",
        "kb_type": "energy_audit",
        "qdrant_collection": "energy_quota_standards",
        "embedding_model": "dashscope/text-embedding-v3",
        "chunking_config": {"strategy": "three_level", "max_chunk_size": 512, "overlap": 64},
        "indexing_strategy": {"vector": True, "summary": True, "graph": True, "wiki": True},
    },
    {
        "id": "energy_audit_technical_guidelines",
        "name": "能源审计技术指南",
        "description": "存储能源审计技术规范和指南文档，支持审计流程查询。",
        "kb_type": "energy_audit",
        "qdrant_collection": "energy_audit_technical_guidelines",
        "embedding_model": "dashscope/text-embedding-v3",
        "chunking_config": {"strategy": "three_level", "max_chunk_size": 512, "overlap": 64},
        "indexing_strategy": {"vector": True, "summary": True, "graph": True, "wiki": True},
    },
]

# Supported file extensions for KB documents
_SUPPORTED_EXTS = {
    ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".pptx", ".ppt",
    ".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm",
    ".epub", ".mhtml",
    ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp",
    ".mp3", ".wav", ".m4a", ".flac", ".ogg",
}

# Extensions that are safe to decode as UTF-8 text for generic chunking.
_TEXT_EXTS = {".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm"}

# ── SQLite metadata store ──
_DB_PATH = _DEFAULT_KNOWLEDGE_ROOT / ".knowledge_meta.db"
_db_local = threading.local()
_db_init_lock = threading.Lock()
_db_global_conn: sqlite3.Connection | None = None  # main thread fallback
_stale_jobs_cleaned = False  # one-time cleanup guard for orphaned jobs


def _ensure_default_kb() -> None:
    """Create the 3 default knowledge bases if they do not exist.

    Each KB maps 1:1:
      - Remote Qdrant collection: {kb_id}
      - Local directory:          rag/data/{kb_id}/
    """
    db = _get_db()
    now = _now_iso()
    for kb_def in _DEFAULT_KBS:
        kb_id = kb_def["id"]
        row = db.execute("SELECT 1 FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
        if row:
            continue
        kb_root = _RAG_DATA_ROOT / kb_id
        kb_root.mkdir(parents=True, exist_ok=True)
        db.execute(
            """
            INSERT INTO knowledge_bases
            (id, name, description, kb_type, root_path, qdrant_collection,
             embedding_model, chunking_config, indexing_strategy, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                kb_id,
                kb_def["name"],
                kb_def["description"],
                kb_def["kb_type"],
                str(kb_root),
                kb_def["qdrant_collection"],
                kb_def["embedding_model"],
                json.dumps(kb_def["chunking_config"]),
                json.dumps(kb_def["indexing_strategy"]),
                now,
                now,
            ),
        )
    db.commit()


def _init_schema(conn: sqlite3.Connection) -> None:
    """Create all KB metadata tables."""
    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS knowledge_bases (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            description TEXT,
            kb_type TEXT NOT NULL,
            root_path TEXT NOT NULL,
            qdrant_collection TEXT NOT NULL,
            embedding_model TEXT,
            chunking_config TEXT,
            indexing_strategy TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS knowledge_folders (
            id TEXT PRIMARY KEY,
            kb_id TEXT NOT NULL REFERENCES knowledge_bases(id) ON DELETE CASCADE,
            parent_id TEXT REFERENCES knowledge_folders(id) ON DELETE CASCADE,
            name TEXT NOT NULL,
            path TEXT NOT NULL,
            depth INTEGER DEFAULT 0,
            sort_order INTEGER DEFAULT 0,
            created_at TEXT NOT NULL,
            UNIQUE(kb_id, path)
        );
        CREATE INDEX IF NOT EXISTS idx_kf_kb_parent ON knowledge_folders(kb_id, parent_id);
        CREATE INDEX IF NOT EXISTS idx_kf_kb_path ON knowledge_folders(kb_id, path);

        CREATE TABLE IF NOT EXISTS knowledge_documents (
            id TEXT PRIMARY KEY,
            kb_id TEXT NOT NULL REFERENCES knowledge_bases(id) ON DELETE CASCADE,
            folder_id TEXT REFERENCES knowledge_folders(id) ON DELETE SET NULL,
            title TEXT,
            file_name TEXT NOT NULL,
            file_type TEXT,
            file_size INTEGER,
            file_path TEXT NOT NULL,
            file_hash TEXT,
            source TEXT DEFAULT 'upload',
            metadata TEXT,
            parse_status TEXT DEFAULT 'pending',
            summary_status TEXT DEFAULT 'none',
            summary_text TEXT,
            chunk_count INTEGER DEFAULT 0,
            vector_count INTEGER DEFAULT 0,
            error_message TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            processed_at TEXT,
            UNIQUE(kb_id, file_path)
        );
        CREATE INDEX IF NOT EXISTS idx_kd_kb_folder ON knowledge_documents(kb_id, folder_id);
        CREATE INDEX IF NOT EXISTS idx_kd_kb_status ON knowledge_documents(kb_id, parse_status);
        CREATE INDEX IF NOT EXISTS idx_kd_kb_hash ON knowledge_documents(kb_id, file_hash);

        CREATE TABLE IF NOT EXISTS knowledge_chunks (
            id TEXT PRIMARY KEY,
            doc_id TEXT NOT NULL REFERENCES knowledge_documents(id) ON DELETE CASCADE,
            kb_id TEXT NOT NULL,
            chunk_index INTEGER,
            chunk_type TEXT,
            content TEXT NOT NULL,
            content_hash TEXT,
            char_count INTEGER,
            is_enabled INTEGER DEFAULT 1,
            parent_chunk_id TEXT,
            prev_chunk_id TEXT,
            next_chunk_id TEXT,
            metadata TEXT,
            qdrant_point_id TEXT,
            created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_kc_doc ON knowledge_chunks(doc_id);
        CREATE INDEX IF NOT EXISTS idx_kc_kb ON knowledge_chunks(kb_id);

        CREATE TABLE IF NOT EXISTS vectorization_jobs (
            id TEXT PRIMARY KEY,
            kb_id TEXT NOT NULL,
            doc_id TEXT NOT NULL REFERENCES knowledge_documents(id) ON DELETE CASCADE,
            status TEXT DEFAULT 'pending',
            progress INTEGER DEFAULT 0,
            chunks_total INTEGER DEFAULT 0,
            chunks_done INTEGER DEFAULT 0,
            error TEXT,
            started_at TEXT,
            completed_at TEXT,
            created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_vj_kb_status ON vectorization_jobs(kb_id, status);
        CREATE INDEX IF NOT EXISTS idx_vj_doc ON vectorization_jobs(doc_id);
        CREATE UNIQUE INDEX IF NOT EXISTS idx_vj_processing_doc ON vectorization_jobs(doc_id) WHERE status = 'processing';

        CREATE TABLE IF NOT EXISTS knowledge_entities (
            id TEXT PRIMARY KEY,
            kb_id TEXT NOT NULL REFERENCES knowledge_bases(id) ON DELETE CASCADE,
            doc_id TEXT REFERENCES knowledge_documents(id) ON DELETE CASCADE,
            name TEXT NOT NULL,
            entity_type TEXT,
            description TEXT,
            metadata TEXT,
            qdrant_point_id TEXT,
            created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_ke_kb ON knowledge_entities(kb_id);
        CREATE INDEX IF NOT EXISTS idx_ke_name ON knowledge_entities(kb_id, name);

        CREATE TABLE IF NOT EXISTS knowledge_relationships (
            id TEXT PRIMARY KEY,
            kb_id TEXT NOT NULL REFERENCES knowledge_bases(id) ON DELETE CASCADE,
            doc_id TEXT REFERENCES knowledge_documents(id) ON DELETE CASCADE,
            source_entity_id TEXT NOT NULL REFERENCES knowledge_entities(id) ON DELETE CASCADE,
            target_entity_id TEXT NOT NULL REFERENCES knowledge_entities(id) ON DELETE CASCADE,
            relation_type TEXT,
            description TEXT,
            metadata TEXT,
            created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_kr_kb ON knowledge_relationships(kb_id);
        CREATE INDEX IF NOT EXISTS idx_kr_src ON knowledge_relationships(source_entity_id);
        CREATE INDEX IF NOT EXISTS idx_kr_tgt ON knowledge_relationships(target_entity_id);

        CREATE TABLE IF NOT EXISTS knowledge_wiki_pages (
            id TEXT PRIMARY KEY,
            kb_id TEXT NOT NULL REFERENCES knowledge_bases(id) ON DELETE CASCADE,
            doc_id TEXT REFERENCES knowledge_documents(id) ON DELETE CASCADE,
            folder_id TEXT REFERENCES knowledge_folders(id) ON DELETE CASCADE,
            title TEXT NOT NULL,
            slug TEXT,
            content TEXT NOT NULL,
            status TEXT DEFAULT 'completed',
            review_status TEXT DEFAULT 'pending',
            source_hash TEXT,
            qdrant_point_id TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_kw_kb ON knowledge_wiki_pages(kb_id);
        CREATE INDEX IF NOT EXISTS idx_kw_title ON knowledge_wiki_pages(kb_id, title);

        CREATE TABLE IF NOT EXISTS knowledge_curation_jobs (
            id TEXT PRIMARY KEY,
            kb_id TEXT NOT NULL REFERENCES knowledge_bases(id) ON DELETE CASCADE,
            folder_id TEXT REFERENCES knowledge_folders(id) ON DELETE SET NULL,
            job_type TEXT NOT NULL DEFAULT 'curation',
            status TEXT NOT NULL DEFAULT 'pending',
            input_pages TEXT,
            output_pages TEXT,
            error_message TEXT,
            started_at TEXT,
            completed_at TEXT,
            created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_kcj_kb ON knowledge_curation_jobs(kb_id);
        CREATE INDEX IF NOT EXISTS idx_kcj_status ON knowledge_curation_jobs(kb_id, status);

        -- Legacy table kept for backwards compatibility during transition
        CREATE TABLE IF NOT EXISTS document_state (
            rel_path TEXT PRIMARY KEY,
            status TEXT DEFAULT 'pending',
            progress INTEGER DEFAULT 0,
            chunks_done INTEGER DEFAULT 0,
            chunks_total INTEGER DEFAULT 0,
            error TEXT DEFAULT '',
            file_hash TEXT DEFAULT '',
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE INDEX IF NOT EXISTS idx_doc_status ON document_state(status);
        """
    )
    conn.commit()


def _migrate_schema(conn: sqlite3.Connection) -> None:
    """Apply lightweight schema migrations for existing databases.

    `CREATE TABLE IF NOT EXISTS` only handles missing tables; this function
    adds columns/indexes that were introduced after the initial schema.
    """
    try:
        conn.execute("ALTER TABLE knowledge_wiki_pages ADD COLUMN folder_id TEXT REFERENCES knowledge_folders(id) ON DELETE CASCADE")
    except sqlite3.OperationalError:
        pass  # column already exists
    try:
        conn.execute("CREATE INDEX idx_kw_folder ON knowledge_wiki_pages(kb_id, folder_id)")
    except sqlite3.OperationalError:
        pass  # index already exists
    try:
        conn.execute("""
            CREATE TABLE knowledge_curation_jobs (
                id TEXT PRIMARY KEY,
                kb_id TEXT NOT NULL REFERENCES knowledge_bases(id) ON DELETE CASCADE,
                folder_id TEXT REFERENCES knowledge_folders(id) ON DELETE SET NULL,
                job_type TEXT NOT NULL DEFAULT 'curation',
                status TEXT NOT NULL DEFAULT 'pending',
                input_pages TEXT,
                output_pages TEXT,
                error_message TEXT,
                started_at TEXT,
                completed_at TEXT,
                created_at TEXT NOT NULL
            )
        """)
        conn.execute("CREATE INDEX idx_kcj_kb ON knowledge_curation_jobs(kb_id)")
        conn.execute("CREATE INDEX idx_kcj_status ON knowledge_curation_jobs(kb_id, status)")
    except sqlite3.OperationalError:
        pass  # table/index already exists
    try:
        conn.execute("ALTER TABLE knowledge_curation_jobs ADD COLUMN job_type TEXT NOT NULL DEFAULT 'curation'")
    except sqlite3.OperationalError:
        pass  # column already exists
    try:
        conn.execute("ALTER TABLE knowledge_wiki_pages ADD COLUMN review_status TEXT DEFAULT 'pending'")
    except sqlite3.OperationalError:
        pass  # column already exists
    try:
        conn.execute("CREATE INDEX idx_kw_review ON knowledge_wiki_pages(kb_id, review_status)")
    except sqlite3.OperationalError:
        pass  # index already exists
    try:
        conn.execute("ALTER TABLE knowledge_wiki_pages ADD COLUMN source_hash TEXT")
    except sqlite3.OperationalError:
        pass  # column already exists
    try:
        conn.execute("ALTER TABLE knowledge_wiki_pages ADD COLUMN quality_score REAL")
    except sqlite3.OperationalError:
        pass  # column already exists
    try:
        conn.execute("ALTER TABLE knowledge_wiki_pages ADD COLUMN quality_report TEXT")
    except sqlite3.OperationalError:
        pass  # column already exists
    conn.commit()

    # One-time cleanup: any jobs that were still running when the server last
    # exited are now orphaned; mark them failed so the UI does not show stale
    # running jobs. Guarded by a process-level flag so we don't mark new jobs.
    global _stale_jobs_cleaned
    if not _stale_jobs_cleaned:
        _stale_jobs_cleaned = True
        conn.execute(
            "UPDATE knowledge_curation_jobs SET status = ?, error_message = ?, completed_at = ? WHERE status IN (?, ?)",
            ("failed", "Server restarted while job was running", _now_iso(), "pending", "running"),
        )
        # Orphaned vectorization jobs must also be cleared: the partial unique
        # index on (doc_id) WHERE status='processing' would otherwise block all
        # future vectorization attempts for those documents.
        conn.execute(
            "UPDATE vectorization_jobs SET status = ?, error = ?, completed_at = ? WHERE status IN (?, ?)",
            ("failed", "Server restarted while job was running", _now_iso(), "pending", "processing"),
        )
        conn.commit()


def _get_db() -> sqlite3.Connection:
    """Return a SQLite connection for the current thread."""
    global _db_global_conn
    # Main thread uses the global connection for backwards compatibility;
    # worker threads get their own connection to avoid SQLite threading issues.
    if threading.current_thread() is threading.main_thread():
        if _db_global_conn is None:
            with _db_init_lock:
                if _db_global_conn is None:
                    _DEFAULT_KNOWLEDGE_ROOT.mkdir(parents=True, exist_ok=True)
                    _db_global_conn = sqlite3.connect(str(_DB_PATH), check_same_thread=False)
                    _db_global_conn.row_factory = sqlite3.Row
                    _db_global_conn.execute("PRAGMA journal_mode=WAL")
                    _db_global_conn.execute("PRAGMA busy_timeout=10000")
                    _db_global_conn.execute("PRAGMA foreign_keys=ON")
                    _init_schema(_db_global_conn)
                    _migrate_schema(_db_global_conn)
                    _ensure_default_kb()
        return _db_global_conn

    conn = getattr(_db_local, "conn", None)
    if conn is None:
        with _db_init_lock:
            _DEFAULT_KNOWLEDGE_ROOT.mkdir(parents=True, exist_ok=True)
            conn = sqlite3.connect(str(_DB_PATH), check_same_thread=False)
            conn.row_factory = sqlite3.Row
            conn.execute("PRAGMA journal_mode=WAL")
            conn.execute("PRAGMA busy_timeout=10000")
            conn.execute("PRAGMA foreign_keys=ON")
            _init_schema(conn)
            _migrate_schema(conn)
        _db_local.conn = conn
    return conn


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _new_id() -> str:
    return uuid.uuid4().hex


def _file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _compute_doc_source_hash(doc_id: str) -> str:
    """Return a stable source hash for a document.

    Prefers the persisted file_hash; falls back to a hash of the extracted
    preview text so that moved/deleted source files still participate in
    incremental updates.
    """
    doc = get_knowledge_document(doc_id)
    file_hash = doc.get("file_hash")
    if file_hash:
        return str(file_hash)
    try:
        preview = get_knowledge_file_preview_v2(doc_id, max_chars=12_000)
        content = preview.get("content", "")
    except Exception:
        content = ""
    h = hashlib.sha256(content.encode("utf-8"))
    return h.hexdigest()


def _compute_folder_wiki_source_hash(kb_id: str, folder_id: str | None, recursive: bool = True) -> str:
    """Return a stable hash representing the source documents for a folder wiki."""
    doc_ids = _get_folder_doc_ids(kb_id, folder_id, recursive=recursive)
    parts: list[str] = []
    for doc_id in doc_ids:
        parts.append(f"{doc_id}:{_compute_doc_source_hash(doc_id)}")
    parts.sort()
    h = hashlib.sha256()
    h.update("\n".join(parts).encode("utf-8"))
    return h.hexdigest()


def _get_kb_curation_config(kb: dict) -> dict:
    """Return curation/quality configuration for a KB, with sensible defaults."""
    strategy = kb.get("indexing_strategy") or {}
    threshold = strategy.get("auto_approve_threshold")
    # Backwards compatibility with the older key name.
    if threshold is None:
        threshold = strategy.get("wiki_quality_threshold", 100)
    return {
        "auto_approve_threshold": float(threshold),
        "wiki_temperature": float(strategy.get("wiki_temperature", 0.4)),
        "wiki_max_tokens": int(strategy.get("wiki_max_tokens", 1500)),
        "curator_temperature": float(strategy.get("curator_temperature", 0.2)),
        "curator_top_k": int(strategy.get("curator_top_k", 5)),
        "curation_auto_approve": bool(strategy.get("curation_auto_approve", True)),
    }


def _to_float(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _parse_quality_json(raw: str | None) -> dict:
    """Parse LLM quality evaluation JSON into a normalized dict."""
    default = {"score": 0.0, "report": "评估失败", "details": {}}
    if not raw:
        return default
    text = raw.strip()
    if "```" in text:
        start = text.find("```json")
        if start == -1:
            start = text.find("```")
        end = text.find("```", start + 3)
        if end != -1:
            text = text[start:end].strip("`").strip()
            if text.lower().startswith("json"):
                text = text[4:].strip()
    try:
        data = json.loads(text)
    except Exception:
        return default
    if not isinstance(data, dict):
        return default
    details = {
        "completeness": _to_float(data.get("completeness")),
        "accuracy": _to_float(data.get("accuracy")),
        "structure": _to_float(data.get("structure")),
        "readability": _to_float(data.get("readability")),
    }
    overall = _to_float(data.get("overall"))
    if overall is None:
        vals = [v for v in details.values() if v is not None]
        overall = sum(vals) / len(vals) if vals else 0.0
    for k, v in details.items():
        if v is None:
            details[k] = 0.0
    report = data.get("suggestions") or data.get("report") or ""
    return {"score": overall, "report": str(report), "details": details}


def _evaluate_wiki_quality(wiki_id: str, title: str, content: str, source_text: str = "", temperature: float = 0.2, max_tokens: int = 600) -> dict:
    """Use an LLM to score a generated wiki page and return a quality report."""
    if not content or not content.strip():
        return {"score": 0.0, "report": "内容为空", "details": {}}
    prompt = (
        "你是一位技术文档质量评估专家。请对下面生成的 Wiki 页面进行质量评估。\n\n"
        "评分维度（每项 0-100）：\n"
        "- completeness: 内容是否完整覆盖源材料关键信息\n"
        "- accuracy: 信息是否准确，无幻觉\n"
        "- structure: 结构是否清晰、标题层级合理\n"
        "- readability: 语言是否流畅、易于阅读\n\n"
        "请严格按 JSON 格式返回，不要包含任何其他内容：\n"
        "{\n"
        '  "completeness": 75,\n'
        '  "accuracy": 80,\n'
        '  "structure": 70,\n'
        '  "readability": 85,\n'
        '  "overall": 78,\n'
        '  "suggestions": "简要改进建议"\n'
        "}\n\n"
        f"标题: {title}\n"
        f"源材料摘要:\n{source_text[:4000]}\n\n"
        f"生成的 Wiki 内容:\n{content[:8000]}\n"
    )
    raw = _call_llm(
        [{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=max_tokens,
        task="wiki_quality_eval",
    )
    return _parse_quality_json(raw)


def _file_type_from_name(name: str) -> str:
    return Path(name).suffix.lower().lstrip(".") or "unknown"


def _json_loads(value: str | None) -> Any:
    if not value:
        return {}
    try:
        return json.loads(value)
    except Exception:
        return {}


def _json_dumps(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, default=str)


def _slugify(text: str) -> str:
    """Create a URL-friendly slug from arbitrary text, preserving CJK characters."""
    text = text.strip().lower()
    # Keep CJK characters, letters, digits, spaces, and dashes; remove other punctuation.
    text = re.sub(r"[^\u4e00-\u9fff\w\s-]", "", text)
    text = re.sub(r"[-\s]+", "-", text)
    return text.strip("-") or "wiki"


# ── path security helpers ──

def _resolve_kb_root(kb_id: str) -> Path:
    db = _get_db()
    row = db.execute("SELECT root_path FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
    if not row:
        raise ValueError(f"Knowledge base not found: {kb_id}")
    return Path(row["root_path"]).resolve()


def _get_kb_collection_name(kb_id: str) -> str:
    """Return the Qdrant collection name for a knowledge base."""
    db = _get_db()
    row = db.execute(
        "SELECT qdrant_collection FROM knowledge_bases WHERE id = ?", (kb_id,)
    ).fetchone()
    if not row:
        raise ValueError(f"Knowledge base not found: {kb_id}")
    return row["qdrant_collection"]


def _resolve_path(kb_id: str, rel_path: str, must_exist: bool = False) -> Path:
    """Resolve a relative path inside a KB root, preventing traversal."""
    root = _resolve_kb_root(kb_id)
    normalized = os.path.normpath(rel_path).lstrip("\\/")
    if normalized.startswith("..") or os.path.isabs(normalized):
        raise ValueError("Invalid path (traversal attempt)")
    target = (root / normalized).resolve()
    # Path.is_relative_to compares path components, so it is immune to the
    # case/separator mismatches that can fool a raw str.startswith check on
    # Windows (e.g. "C:\KB-root2" passing a "C:\KB-root" prefix test).
    if not target.is_relative_to(root):
        raise ValueError("Invalid path (outside root)")
    if must_exist and not target.exists():
        raise FileNotFoundError(f"Path not found: {normalized}")
    return target


def _rel_path(kb_id: str, target: Path) -> str:
    root = _resolve_kb_root(kb_id)
    try:
        rel = str(target.relative_to(root)).replace("\\", "/")
        return "" if rel == "." else rel
    except ValueError:
        return str(target).replace("\\", "/")


def _is_hidden(name: str) -> bool:
    return name.startswith(".") or name.startswith("~$")


# ═══════════════════════════════════════════════════════════════════════════════
# Knowledge Base CRUD
# ═══════════════════════════════════════════════════════════════════════════════


def create_knowledge_base(
    name: str,
    kb_type: str = "energy_audit",
    description: str = "",
    root_path: str | None = None,
    qdrant_collection: str | None = None,
    embedding_model: str = "dashscope/text-embedding-v3",
    chunking_config: dict | None = None,
    indexing_strategy: dict | None = None,
) -> dict:
    """Create a new knowledge base."""
    kb_id = _new_id()
    if root_path:
        kb_root = Path(root_path).resolve()
    else:
        kb_root = _DEFAULT_KNOWLEDGE_ROOT / kb_id
    kb_root.mkdir(parents=True, exist_ok=True)

    if not qdrant_collection:
        qdrant_collection = f"kb_{kb_id}"

    now = _now_iso()
    db = _get_db()
    db.execute(
        """
        INSERT INTO knowledge_bases
        (id, name, description, kb_type, root_path, qdrant_collection,
         embedding_model, chunking_config, indexing_strategy, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            kb_id,
            name,
            description,
            kb_type,
            str(kb_root),
            qdrant_collection,
            embedding_model,
            _json_dumps(chunking_config or {}),
            _json_dumps(indexing_strategy or {"vector": True, "summary": True, "graph": False, "wiki": False}),
            now,
            now,
        ),
    )
    db.commit()
    return {
        "id": kb_id,
        "name": name,
        "description": description,
        "kb_type": kb_type,
        "root_path": str(kb_root),
        "qdrant_collection": qdrant_collection,
        "created_at": now,
    }


def list_knowledge_bases() -> list[dict]:
    db = _get_db()
    rows = db.execute("SELECT * FROM knowledge_bases ORDER BY created_at DESC").fetchall()
    result = []
    for row in rows:
        stats = _get_kb_stats(row["id"])
        result.append({
            "id": row["id"],
            "name": row["name"],
            "description": row["description"],
            "kb_type": row["kb_type"],
            "root_path": row["root_path"],
            "qdrant_collection": row["qdrant_collection"],
            "embedding_model": row["embedding_model"],
            "chunking_config": _json_loads(row["chunking_config"]),
            "indexing_strategy": _json_loads(row["indexing_strategy"]),
            "created_at": row["created_at"],
            "updated_at": row["updated_at"],
            "is_system": row["id"] in _SYSTEM_KB_IDS,
            "stats": stats,
        })
    return result


def get_knowledge_base(kb_id: str) -> dict:
    db = _get_db()
    row = db.execute("SELECT * FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
    if not row:
        raise ValueError(f"Knowledge base not found: {kb_id}")
    return {
        "id": row["id"],
        "name": row["name"],
        "description": row["description"],
        "kb_type": row["kb_type"],
        "root_path": row["root_path"],
        "qdrant_collection": row["qdrant_collection"],
        "embedding_model": row["embedding_model"],
        "chunking_config": _json_loads(row["chunking_config"]),
        "indexing_strategy": _json_loads(row["indexing_strategy"]),
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
        "is_system": row["id"] in _SYSTEM_KB_IDS,
        "stats": _get_kb_stats(kb_id),
    }


def update_knowledge_base(kb_id: str, **kwargs) -> dict:
    allowed = {"name", "description", "chunking_config", "indexing_strategy"}
    updates = {k: v for k, v in kwargs.items() if k in allowed}
    if not updates:
        return get_knowledge_base(kb_id)

    sets = ", ".join(f"{k} = ?" for k in updates)
    values = list(updates.values())
    if "chunking_config" in updates:
        values[list(updates).index("chunking_config")] = _json_dumps(updates["chunking_config"])
    if "indexing_strategy" in updates:
        values[list(updates).index("indexing_strategy")] = _json_dumps(updates["indexing_strategy"])
    values.append(_now_iso())
    values.append(kb_id)

    db = _get_db()
    db.execute(
        f"UPDATE knowledge_bases SET {sets}, updated_at = ? WHERE id = ?",
        values,
    )
    db.commit()
    return get_knowledge_base(kb_id)


def delete_knowledge_base(kb_id: str, delete_files: bool = False) -> dict:
    if kb_id in _SYSTEM_KB_IDS:
        raise ValueError("Cannot delete a system knowledge base")
    kb = get_knowledge_base(kb_id)
    root = Path(kb["root_path"])

    db = _get_db()
    db.execute("DELETE FROM knowledge_bases WHERE id = ?", (kb_id,))
    db.commit()

    # Drop the KB's Qdrant collections (chunks / entities / wiki)
    for coll in (kb["qdrant_collection"], _kb_entity_collection(kb), _kb_wiki_collection(kb)):
        _qdrant_drop_collection(coll)

    if delete_files and root.exists():
        shutil.rmtree(root, ignore_errors=True)

    return {"id": kb_id, "deleted": True}


def _get_kb_stats(kb_id: str, folder_path: str | None = None) -> dict:
    db = _get_db()
    scope = "kb_id = ?"
    args: list[Any] = [kb_id]
    folder_path = (folder_path or "").strip("/")
    if folder_path:
        scope += " AND file_path LIKE ? ESCAPE '\\'"
        args.append(_escape_like(folder_path) + "/%")
    total = db.execute(
        f"SELECT COUNT(*) FROM knowledge_documents WHERE {scope}", args
    ).fetchone()[0]
    completed = db.execute(
        f"SELECT COUNT(*) FROM knowledge_documents WHERE {scope} AND parse_status = 'completed'", args
    ).fetchone()[0]
    processing = db.execute(
        f"SELECT COUNT(*) FROM knowledge_documents WHERE {scope} AND parse_status = 'processing'", args
    ).fetchone()[0]
    failed = db.execute(
        f"SELECT COUNT(*) FROM knowledge_documents WHERE {scope} AND parse_status = 'failed'", args
    ).fetchone()[0]
    total_size = db.execute(
        f"SELECT COALESCE(SUM(file_size), 0) FROM knowledge_documents WHERE {scope}", args
    ).fetchone()[0]
    # Orphans: DB rows whose backing file no longer exists on disk. Only
    # computed for the full-KB scope (folder-scoped chips don't need it).
    orphaned = 0
    if not folder_path:
        orphaned = len(list_orphaned_documents(kb_id))
    return {
        "total_documents": total,
        "completed": completed,
        "processing": processing,
        "failed": failed,
        "total_size": total_size,
        "orphaned": orphaned,
    }


def _escape_like(value: str) -> str:
    """Escape LIKE wildcards so file paths containing % or _ match literally."""
    return value.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")


# ═══════════════════════════════════════════════════════════════════════════════
# Folder CRUD (persistent hierarchy)
# ═══════════════════════════════════════════════════════════════════════════════


def list_knowledge_folders(kb_id: str, parent_id: str | None = None) -> list[dict]:
    db = _get_db()
    if parent_id:
        rows = db.execute(
            "SELECT * FROM knowledge_folders WHERE kb_id = ? AND parent_id = ? ORDER BY sort_order, name",
            (kb_id, parent_id),
        ).fetchall()
    else:
        rows = db.execute(
            "SELECT * FROM knowledge_folders WHERE kb_id = ? AND parent_id IS NULL ORDER BY sort_order, name",
            (kb_id,),
        ).fetchall()
    return [_folder_row_to_dict(row) for row in rows]


def get_folder_tree(kb_id: str, parent_id: str | None = None) -> list[dict]:
    """Return nested folder tree with documents populated."""
    folders = list_knowledge_folders(kb_id, parent_id)
    for folder in folders:
        folder["children"] = get_folder_tree(kb_id, folder["id"])
        folder["documents"] = _list_docs_in_folder(kb_id, folder["id"])
    return folders


def create_knowledge_folder_v2(kb_id: str, name: str, parent_id: str | None = None) -> dict:
    """Create a folder in metadata and on disk."""
    db = _get_db()
    parent_path = ""
    parent_depth = -1
    if parent_id:
        parent = db.execute(
            "SELECT * FROM knowledge_folders WHERE id = ? AND kb_id = ?",
            (parent_id, kb_id),
        ).fetchone()
        if not parent:
            raise ValueError("Parent folder not found")
        parent_path = parent["path"]
        parent_depth = parent["depth"]

    folder_id = _new_id()
    path = f"{parent_path}/{name}".lstrip("/")
    depth = parent_depth + 1
    target = _resolve_path(kb_id, path)
    target.mkdir(parents=True, exist_ok=True)

    now = _now_iso()
    try:
        db.execute(
            """
            INSERT INTO knowledge_folders (id, kb_id, parent_id, name, path, depth, sort_order, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (folder_id, kb_id, parent_id, name, path, depth, 0, now),
        )
        db.commit()
    except sqlite3.IntegrityError:
        raise ValueError(f"Folder already exists: {path}")

    return _folder_row_to_dict(
        db.execute("SELECT * FROM knowledge_folders WHERE id = ?", (folder_id,)).fetchone()
    )


def update_knowledge_folder(kb_id: str, folder_id: str, name: str | None = None, parent_id: str | None = None) -> dict:
    """Rename or move a folder and recursively update descendants."""
    db = _get_db()
    folder = db.execute(
        "SELECT * FROM knowledge_folders WHERE id = ? AND kb_id = ?", (folder_id, kb_id)
    ).fetchone()
    if not folder:
        raise ValueError("Folder not found")

    old_path = folder["path"]
    old_root = _resolve_path(kb_id, old_path)

    # Determine new parent path and depth
    new_parent_path = ""
    new_depth = 0
    if parent_id:
        if parent_id == folder_id:
            raise ValueError("Cannot move a folder into itself")
            # TODO: also prevent moving into own descendant
        parent = db.execute(
            "SELECT * FROM knowledge_folders WHERE id = ? AND kb_id = ?", (parent_id, kb_id)
        ).fetchone()
        if not parent:
            raise ValueError("Parent folder not found")
        # Prevent moving into own descendant
        if _is_descendant(kb_id, folder_id, parent_id):
            raise ValueError("Cannot move a folder into its own descendant")
        new_parent_path = parent["path"]
        new_depth = parent["depth"] + 1

    new_name = name or folder["name"]
    new_path = f"{new_parent_path}/{new_name}".lstrip("/")

    if new_path == old_path:
        return _folder_row_to_dict(folder)

    new_root = _resolve_path(kb_id, new_path)
    if new_root.exists():
        raise ValueError(f"Folder already exists: {new_path}")

    # Move on disk
    old_root.rename(new_root)

    # Update folder and all descendants in DB
    db.execute(
        "UPDATE knowledge_folders SET name = ?, parent_id = ?, path = ?, depth = ? WHERE id = ?",
        (new_name, parent_id, new_path, new_depth, folder_id),
    )

    # Recursively update descendants
    descendants = db.execute(
        "SELECT id, path, depth FROM knowledge_folders WHERE kb_id = ? AND path LIKE ? AND id != ?",
        (kb_id, old_path + "/%", folder_id),
    ).fetchall()
    for desc in descendants:
        rel = desc["path"][len(old_path) + 1 :]
        desc_new_path = f"{new_path}/{rel}"
        depth_delta = new_depth - folder["depth"]
        db.execute(
            "UPDATE knowledge_folders SET path = ?, depth = ? WHERE id = ?",
            (desc_new_path, desc["depth"] + depth_delta, desc["id"]),
        )
        # Update document file_path references
        docs = db.execute(
            "SELECT id, file_path FROM knowledge_documents WHERE folder_id = ?",
            (desc["id"],),
        ).fetchall()
        for doc in docs:
            new_doc_path = doc["file_path"].replace(desc["path"], desc_new_path, 1)
            db.execute(
                "UPDATE knowledge_documents SET file_path = ? WHERE id = ?",
                (new_doc_path, doc["id"]),
            )

    # Update own documents
    own_docs = db.execute(
        "SELECT id, file_path FROM knowledge_documents WHERE folder_id = ?", (folder_id,)
    ).fetchall()
    for doc in own_docs:
        new_doc_path = doc["file_path"].replace(old_path, new_path, 1)
        db.execute(
            "UPDATE knowledge_documents SET file_path = ? WHERE id = ?",
            (new_doc_path, doc["id"]),
        )

    db.commit()
    return _folder_row_to_dict(
        db.execute("SELECT * FROM knowledge_folders WHERE id = ?", (folder_id,)).fetchone()
    )


def delete_knowledge_folder_v2(kb_id: str, folder_id: str, recursive: bool = False) -> dict:
    db = _get_db()
    folder = db.execute(
        "SELECT * FROM knowledge_folders WHERE id = ? AND kb_id = ?", (folder_id, kb_id)
    ).fetchone()
    if not folder:
        raise ValueError("Folder not found")

    has_children = db.execute(
        "SELECT 1 FROM knowledge_folders WHERE parent_id = ? LIMIT 1", (folder_id,)
    ).fetchone()
    has_docs = db.execute(
        "SELECT 1 FROM knowledge_documents WHERE folder_id = ? LIMIT 1", (folder_id,)
    ).fetchone()

    if (has_children or has_docs) and not recursive:
        raise ValueError("Folder not empty; use recursive=true")

    # Collect all folder ids under this tree (including self) before deleting anything.
    folder_rows = db.execute(
        "SELECT id FROM knowledge_folders WHERE kb_id = ? AND (id = ? OR path LIKE ?)",
        (kb_id, folder_id, folder["path"] + "/%"),
    ).fetchall()
    folder_ids = [r["id"] for r in folder_rows]

    # Delete all documents in these folders first, so their vectors/files are purged
    # before the folder rows disappear (documents.folder_id is ON DELETE SET NULL).
    doc_rows = db.execute(
        f"SELECT id FROM knowledge_documents WHERE kb_id = ? AND folder_id IN ({','.join('?' * len(folder_ids))})",
        (kb_id, *folder_ids),
    ).fetchall()

    # Delete documents first (SQLite + vectors + file cleanup) while their
    # file paths are still known. If the process crashes after this loop,
    # leftover empty folders can be rescanned and removed later; deleting the
    # disk tree before the metadata would leave ghost document rows instead.
    for row in doc_rows:
        try:
            delete_knowledge_document(row["id"])
        except Exception:
            logger.exception("Failed to delete document %s while deleting folder %s", row["id"], folder_id)

    target = _resolve_path(kb_id, folder["path"])
    if target.exists():
        shutil.rmtree(target, ignore_errors=True)

    # Delete descendant folders then the folder itself
    db.execute("DELETE FROM knowledge_folders WHERE kb_id = ? AND path LIKE ?", (kb_id, folder["path"] + "/%"))
    db.execute("DELETE FROM knowledge_folders WHERE id = ?", (folder_id,))
    db.commit()
    return {"id": folder_id, "deleted": True}


def _folder_row_to_dict(row: sqlite3.Row) -> dict:
    return {
        "id": row["id"],
        "kb_id": row["kb_id"],
        "parent_id": row["parent_id"],
        "name": row["name"],
        "path": row["path"],
        "depth": row["depth"],
        "sort_order": row["sort_order"],
        "created_at": row["created_at"],
    }


def _is_descendant(kb_id: str, ancestor_id: str, candidate_id: str) -> bool:
    db = _get_db()
    current = candidate_id
    while current:
        folder = db.execute(
            "SELECT parent_id FROM knowledge_folders WHERE id = ? AND kb_id = ?",
            (current, kb_id),
        ).fetchone()
        if not folder:
            return False
        if folder["parent_id"] == ancestor_id:
            return True
        current = folder["parent_id"]
    return False


def _list_docs_in_folder(kb_id: str, folder_id: str) -> list[dict]:
    db = _get_db()
    rows = db.execute(
        "SELECT * FROM knowledge_documents WHERE kb_id = ? AND folder_id = ? ORDER BY file_name",
        (kb_id, folder_id),
    ).fetchall()
    return [_doc_row_to_dict(row) for row in rows]


def _curation_job_row_to_dict(row: sqlite3.Row) -> dict:
    return {
        "id": row["id"],
        "kb_id": row["kb_id"],
        "folder_id": row["folder_id"],
        "job_type": row["job_type"] or "curation",
        "status": row["status"],
        "input_pages": _json_loads(row["input_pages"]),
        "output_pages": _json_loads(row["output_pages"]),
        "error_message": row["error_message"],
        "started_at": row["started_at"],
        "completed_at": row["completed_at"],
        "created_at": row["created_at"],
    }


def _doc_row_to_dict(row: sqlite3.Row) -> dict:
    modified = ""
    if row["updated_at"]:
        try:
            modified = row["updated_at"][:10].replace("T", " ")
        except Exception:
            modified = row["updated_at"]
    return {
        "id": row["id"],
        "kb_id": row["kb_id"],
        "folder_id": row["folder_id"],
        "title": row["title"],
        "file_name": row["file_name"],
        # UI aliases expected by panels.js doc list / preview
        "name": row["file_name"],
        "size": row["file_size"],
        "file_type": row["file_type"],
        "file_size": row["file_size"],
        "file_path": row["file_path"],
        # UI alias: legacy call sites expect doc.path (file path, not doc id)
        "path": row["file_path"],
        "file_hash": row["file_hash"],
        "source": row["source"],
        "metadata": _json_loads(row["metadata"]),
        "parse_status": row["parse_status"],
        "parseStatus": row["parse_status"],
        "vect": "done" if row["parse_status"] == "completed" else row["parse_status"],
        "summary_status": row["summary_status"],
        "summary_text": row["summary_text"],
        "chunk_count": row["chunk_count"],
        "vector_count": row["vector_count"],
        "error_message": row["error_message"],
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
        "processed_at": row["processed_at"],
        "modified": modified,
    }



# ═══════════════════════════════════════════════════════════════════════════════
# Document management
# ═══════════════════════════════════════════════════════════════════════════════


def _sync_doc_from_disk(kb_id: str, folder_id: str | None, file_path: str) -> dict:
    """Create or update a document row from an existing disk file."""
    db = _get_db()
    abs_path = _resolve_path(kb_id, file_path, must_exist=True)
    rel = _rel_path(kb_id, abs_path)
    existing = db.execute(
        "SELECT id FROM knowledge_documents WHERE kb_id = ? AND file_path = ?",
        (kb_id, rel),
    ).fetchone()

    file_hash = _file_hash(abs_path)
    now = _now_iso()
    stat = abs_path.stat()
    file_name = abs_path.name
    file_type = _file_type_from_name(file_name)
    file_size = stat.st_size

    if existing:
        doc_id = existing["id"]
        db.execute(
            """
            UPDATE knowledge_documents
            SET folder_id = ?, file_name = ?, file_type = ?, file_size = ?,
                file_hash = ?, updated_at = ?
            WHERE id = ?
            """,
            (folder_id, file_name, file_type, file_size, file_hash, now, doc_id),
        )
    else:
        doc_id = _new_id()
        db.execute(
            """
            INSERT INTO knowledge_documents
            (id, kb_id, folder_id, title, file_name, file_type, file_size, file_path,
             file_hash, source, metadata, parse_status, summary_status, summary_text,
             chunk_count, vector_count, error_message, created_at, updated_at, processed_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                doc_id, kb_id, folder_id, file_name, file_name, file_type, file_size, rel,
                file_hash, "upload", "{}", "pending", "none", "", 0, 0, "",
                now, now, None,
            ),
        )
    db.commit()
    return _doc_row_to_dict(
        db.execute("SELECT * FROM knowledge_documents WHERE id = ?", (doc_id,)).fetchone()
    )


def list_knowledge_documents(kb_id: str, params: dict | None = None) -> dict:
    params = params or {}
    page = max(1, int(params.get("page", 1) or 1))
    page_size = max(1, min(100, int(params.get("page_size", 20) or 20)))
    keyword = (params.get("keyword") or "").lower().strip()
    file_type = (params.get("file_type") or "").lower().strip()
    parse_status = (params.get("parse_status") or "").lower().strip()
    folder_id = params.get("folder_id") or None
    sort_key = (params.get("sort_key") or "date").lower().strip()
    sort_order = (params.get("sort_order") or "desc").lower().strip()

    db = _get_db()
    where = ["d.kb_id = ?"]
    args: list[Any] = [kb_id]
    # folder_path prefix filter (folder and its subfolders) — anchored on the
    # stored file_path, robust against broken folder_id linkage in historical
    # rows. The legacy folder_id exact match is kept for old callers.
    folder_path = (params.get("folder_path") or "").strip("/")
    if folder_path:
        where.append("d.file_path LIKE ? ESCAPE '\\'")
        args.append(_escape_like(folder_path) + "/%")
    elif folder_id:
        where.append("folder_id = ?")
        args.append(folder_id)
    if keyword:
        where.append("(LOWER(file_name) LIKE ? OR LOWER(title) LIKE ?)")
        args.extend([f"%{keyword}%", f"%{keyword}%"])
    if file_type:
        where.append("file_type = ?")
        args.append(file_type)
    if parse_status:
        where.append("parse_status = ?")
        args.append(parse_status)

    order_col = {"name": "d.file_name COLLATE NOCASE", "date": "d.updated_at"}.get(sort_key, "d.updated_at")
    direction = "ASC" if sort_order == "asc" else "DESC"

    where_sql = " AND ".join(where)
    total = db.execute(
        f"SELECT COUNT(*) FROM knowledge_documents d WHERE {where_sql}", args
    ).fetchone()[0]

    rows = db.execute(
        f"""
        SELECT d.*, f.path as folder_path
        FROM knowledge_documents d
        LEFT JOIN knowledge_folders f ON d.folder_id = f.id
        WHERE {where_sql}
        ORDER BY {order_col} {direction}
        LIMIT ? OFFSET ?
        """,
        args + [page_size, (page - 1) * page_size],
    ).fetchall()

    data = []
    for row in rows:
        item = _doc_row_to_dict(row)
        item["folder_path"] = row["folder_path"]
        data.append(item)

    return {"data": data, "total": total, "page": page, "page_size": page_size}


def get_knowledge_document(doc_id: str) -> dict:
    db = _get_db()
    row = db.execute("SELECT * FROM knowledge_documents WHERE id = ?", (doc_id,)).fetchone()
    if not row:
        raise ValueError("Document not found")
    return _doc_row_to_dict(row)


def upload_knowledge_file_v2(
    kb_id: str,
    folder_id: str | None,
    raw_body: bytes,
    content_type: str,
    source: str = "upload",
) -> dict:
    """Multipart upload into a KB folder with metadata tracking."""
    boundary = None
    for part in content_type.split(";"):
        part = part.strip()
        if part.startswith("boundary="):
            boundary = part[len("boundary="):].strip('"')
    if not boundary:
        raise ValueError("Missing multipart boundary")

    b_boundary = boundary.encode("utf-8")
    # Split only on delimiters at line starts (RFC 2046): the first delimiter
    # opens the body, subsequent ones are preceded by CRLF. A naive
    # body.split(b"--" + boundary) would silently corrupt binary files (e.g.
    # PDF) whose content happens to contain the boundary byte sequence.
    delimiter = b"--" + b_boundary
    parts: list[bytes] = []
    pos = raw_body.find(delimiter)
    if pos == -1:
        raise ValueError("Malformed multipart body: delimiter not found")
    pos += len(delimiter)
    while True:
        if raw_body[pos : pos + 2] == b"--":  # closing delimiter
            break
        if raw_body[pos : pos + 2] == b"\r\n":
            pos += 2
        next_idx = raw_body.find(b"\r\n" + delimiter, pos)
        if next_idx == -1:
            parts.append(raw_body[pos:])
            break
        parts.append(raw_body[pos:next_idx])
        pos = next_idx + 2 + len(delimiter)

    filename = None
    file_bytes = b""
    for part in parts:
        if b"Content-Disposition" not in part or b"filename=" not in part:
            continue
        header_end = part.find(b"\r\n\r\n")
        if header_end == -1:
            continue  # malformed part without a header/body separator
        headers = part[:header_end].decode("utf-8", errors="replace")
        # Line-start splitting above already excludes the CRLF preceding the
        # next delimiter — do NOT strip a trailing CRLF here, it may belong to
        # the uploaded file itself.
        file_bytes = part[header_end + 4 :]
        for h in headers.split("\r\n"):
            if "filename=" in h:
                fn_start = h.find('filename="')
                if fn_start == -1:
                    continue  # e.g. RFC 5987 filename*= — unsupported, skip
                fn_start += len('filename="')
                fn_end = h.find('"', fn_start)
                if fn_end == -1:
                    continue
                filename = h[fn_start:fn_end]
        break

    if not filename:
        raise ValueError("No file in upload")

    safe_name = os.path.basename(filename)
    if not safe_name or safe_name.startswith("."):
        raise ValueError("Invalid filename")

    # Determine folder path
    folder_path = ""
    if folder_id:
        db = _get_db()
        folder = db.execute(
            "SELECT path FROM knowledge_folders WHERE id = ? AND kb_id = ?",
            (folder_id, kb_id),
        ).fetchone()
        if folder:
            folder_path = folder["path"]

    dest_dir = _resolve_path(kb_id, folder_path) if folder_path else _resolve_kb_root(kb_id)
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / safe_name

    # Dedup by hash
    dest.write_bytes(file_bytes)
    file_hash = _file_hash(dest)
    rel = _rel_path(kb_id, dest)

    db = _get_db()
    dup = db.execute(
        "SELECT id, file_path FROM knowledge_documents WHERE kb_id = ? AND file_hash = ?",
        (kb_id, file_hash),
    ).fetchone()
    if dup:
        # Only remove the just-written duplicate if it landed at a different path.
        # Re-uploading the same file to the same path should keep the disk copy.
        if dup["file_path"] != rel:
            try:
                dest.unlink()
            except OSError:
                logger.debug("Could not unlink duplicate upload at %s (may be locked)", dest)
        existing = _doc_row_to_dict(
            db.execute("SELECT * FROM knowledge_documents WHERE id = ?", (dup["id"],)).fetchone()
        )
        return {"duplicate": True, "existing": existing}

    doc = _sync_doc_from_disk(kb_id, folder_id, rel)
    return {
        "duplicate": False,
        "document": doc,
        "mime": mimetypes.guess_type(safe_name)[0] or "application/octet-stream",
    }


def delete_knowledge_document(doc_id: str) -> dict:
    doc = get_knowledge_document(doc_id)
    kb_id = doc["kb_id"]
    db = _get_db()

    # Delete file on disk. Ignore common transient errors (file locked by another
    # process, permission denied, already deleted) so metadata cleanup proceeds.
    try:
        target = _resolve_path(kb_id, doc["file_path"], must_exist=True)
        if target.is_file():
            target.unlink()
    except (FileNotFoundError, ValueError, PermissionError, OSError):
        pass

    # Purge the document's vectors from Qdrant (chunks, summary, entities, wiki)
    _purge_doc_vectors(kb_id, doc)

    db.execute("DELETE FROM knowledge_documents WHERE id = ?", (doc_id,))
    db.commit()
    _delete_vect_status(doc_id)
    return {"id": doc_id, "deleted": True}


def batch_knowledge_status_v2(doc_ids: list[str]) -> dict:
    db = _get_db()
    results = []
    for doc_id in doc_ids:
        row = db.execute(
            """
            SELECT d.id, d.parse_status, d.chunk_count, d.vector_count, d.error_message,
                   j.progress, j.chunks_done, j.chunks_total
            FROM knowledge_documents d
            LEFT JOIN vectorization_jobs j ON d.id = j.doc_id AND j.status = 'processing'
            WHERE d.id = ?
            """,
            (doc_id,),
        ).fetchone()
        if row:
            results.append({
                "id": row["id"],
                "parse_status": row["parse_status"],
                "parseStatus": row["parse_status"],
                "vect": "done" if row["parse_status"] == "completed" else row["parse_status"],
                "chunk_count": row["chunk_count"],
                "vector_count": row["vector_count"],
                "progress": row["progress"] or 0,
                "chunks_done": row["chunks_done"] or 0,
                "chunks_total": row["chunks_total"] or 0,
                "error": row["error_message"] or "",
            })
        else:
            results.append({
                "id": doc_id,
                "parse_status": "unknown",
                "parseStatus": "unknown",
                "vect": "pending",
                "progress": 0,
                "error": "",
            })
    return {"data": results}


def batch_delete_knowledge_documents(kb_id: str, doc_ids: list[str]) -> dict:
    """Delete many documents from a KB and clean up their vectors/files."""
    db = _get_db()
    deleted = 0
    failed = 0
    for doc_id in doc_ids:
        row = db.execute(
            "SELECT id FROM knowledge_documents WHERE id = ? AND kb_id = ?",
            (doc_id, kb_id),
        ).fetchone()
        if not row:
            failed += 1
            continue
        try:
            delete_knowledge_document(doc_id)
            deleted += 1
        except Exception:
            logger.exception("Batch delete failed for doc %s", doc_id)
            failed += 1
    return {"deleted": deleted, "failed": failed}


# ═══════════════════════════════════════════════════════════════════════════════
# Preview & summary
# ═══════════════════════════════════════════════════════════════════════════════


def _init_docreader():
    """Lazy-init the WeKnora docreader parser registry."""
    global _DOCREADER_AVAILABLE, _DOCREADER_REGISTRY
    if _DOCREADER_REGISTRY is not None:
        return
    try:
        # Env override first; fall back to the legacy sibling-directory guess.
        weknora_root = os.environ.get("HERMES_WEKNORA_PATH") or str(
            Path(__file__).resolve().parents[3].parent / "dc_agent" / "WeKnora-main" / "WeKnora-main"
        )
        if weknora_root not in sys.path:
            sys.path.insert(0, weknora_root)
        from docreader.parser.registry import registry as _reg
        _DOCREADER_REGISTRY = _reg
        _DOCREADER_AVAILABLE = True
    except Exception:
        _DOCREADER_AVAILABLE = False
        logger.debug("WeKnora docreader not available; preview will fall back to built-in parsers")


_DOCREADER_AVAILABLE = False
_DOCREADER_REGISTRY = None


def _docreader_extract(file_path: Path) -> str:
    _init_docreader()
    ext = file_path.suffix.lower().lstrip(".")
    if _DOCREADER_AVAILABLE and _DOCREADER_REGISTRY:
        try:
            parser_class = _DOCREADER_REGISTRY.get_parser_class("builtin", ext)
            parser = parser_class()
            with open(file_path, "rb") as f:
                content = f.read()
            doc = parser.parse_into_text(content)
            return doc.content if hasattr(doc, "content") else str(doc)
        except Exception:
            pass

    if ext == "epub":
        try:
            import ebooklib
            from ebooklib import epub
            book = epub.read_epub(str(file_path))
            texts = []
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                import html as _html
                texts.append(_html.unescape(item.get_content().decode("utf-8", errors="replace")))
            return "\n\n".join(texts)
        except ImportError:
            raise ValueError("Install ebooklib to preview EPUB files")
    elif ext == "mhtml":
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                raw = f.read()
            import re
            parts = re.split(r'------=|Content-Type: text/html|Content-Transfer-Encoding:', raw)
            for p in parts:
                if p.strip().startswith("<"):
                    import html as _html
                    clean = re.sub(r'<[^>]+>', ' ', p)
                    return _html.unescape(clean.strip())[:120_000]
            return raw[:120_000]
        except Exception:
            raise ValueError("Cannot parse MHTML file")
    else:
        raise ValueError(f"No parser available for .{ext} files")


def get_knowledge_file_preview_v2(doc_id: str, max_chars: int = 120_000) -> dict:
    doc = get_knowledge_document(doc_id)
    target = _resolve_path(doc["kb_id"], doc["file_path"], must_exist=True)
    suffix = target.suffix.lower()
    content = ""

    docreader_formats = {
        "epub", "mhtml", "pptx", "ppt", "png", "jpg", "jpeg", "gif", "bmp",
        "tiff", "webp", "mp3", "wav", "m4a", "flac", "ogg",
    }
    if suffix in docreader_formats:
        try:
            content = _docreader_extract(target)[:max_chars]
        except Exception:
            content = f"(docreader unavailable — cannot preview .{suffix})"
    elif suffix == ".pdf":
        try:
            import pymupdf
            doc_obj = pymupdf.open(str(target))
            parts = []
            total_chars = 0
            for page in doc_obj:
                text = page.get_text()
                if total_chars + len(text) > max_chars:
                    parts.append(text[: max_chars - total_chars])
                    break
                parts.append(text)
                total_chars += len(text)
            doc_obj.close()
            content = "\n".join(parts)
        except ImportError:
            content = "(pymupdf not installed — cannot preview PDF)"
        except Exception:
            content = "(failed to extract PDF text)"
    elif suffix == ".docx":
        try:
            from docx import Document
            doc_obj = Document(str(target))
            parts = [para.text for para in doc_obj.paragraphs]
            content = "\n".join(parts)[:max_chars]
        except ImportError:
            content = "(python-docx not installed — cannot preview DOCX)"
        except Exception:
            content = "(failed to extract DOCX text)"
    elif suffix in (".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm"):
        try:
            content = target.read_text(encoding="utf-8", errors="replace")[:max_chars]
        except Exception:
            content = "(failed to read file)"
    else:
        content = f"(preview not available for .{suffix} files)"

    lines = content.count("\n") + 1 if content else 0
    return {
        "id": doc_id,
        "path": doc["file_path"],
        "file_name": doc["file_name"],
        "content": content,
        "lines": lines,
        "size": target.stat().st_size,
        "summary": doc.get("summary_text") or "",
    }


def _call_llm_with_fallback(
    messages: list[dict],
    temperature: float = 0.3,
    max_tokens: int = 1000,
    task: str = "kb_task",
) -> str | None:
    """Call LLM via Hermes auxiliary client; fall back to direct DeepSeek.

    All LLM calls in the KB pipeline (summary, graph extraction, wiki generation,
    GraphRAG answer, quality evaluation) share this single fallback chain.
    """
    # Primary: Hermes auxiliary client (unified config + credential pool).
    try:
        from agent.auxiliary_client import call_llm

        from rag.config import summary_model, summary_provider

        resp = call_llm(
            task=task,
            provider=summary_provider(),
            model=summary_model(),
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
        )
        if resp and resp.choices:
            return resp.choices[0].message.content.strip()
    except Exception:
        logger.exception("Hermes auxiliary LLM call failed; falling back to direct DeepSeek")

    # Fallback: direct DeepSeek-compatible call.
    from rag.config import deepseek_api_base, deepseek_api_key, summary_model

    api_key = deepseek_api_key()
    if not api_key:
        return None
    try:
        from openai import OpenAI

        client = OpenAI(api_key=api_key, base_url=deepseek_api_base())
        resp = client.chat.completions.create(
            model=summary_model(),
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return resp.choices[0].message.content.strip() if resp.choices else None
    except Exception:
        logger.exception("Direct LLM call failed (task=%s)", task)
        return None


def _llm_summarize(text: str, max_input_chars: int = 10_000) -> str | None:
    """Summarize text into a concise 3-5 sentence Chinese summary."""
    if len(text) > max_input_chars:
        text = text[:max_input_chars] + "\n\n[...content truncated...]"

    messages = [
        {
            "role": "system",
            "content": (
                "You are a helpful assistant that summarizes documents concisely "
                "in Chinese. Output only the summary, no extra explanation."
            ),
        },
        {
            "role": "user",
            "content": f"请用 3-5 句话总结以下文档的核心内容：\n\n{text}",
        },
    ]
    return _call_llm_with_fallback(
        messages, temperature=0.3, max_tokens=500, task="kb_summary"
    )


def _call_llm(
    messages: list[dict],
    temperature: float = 0.3,
    max_tokens: int = 1000,
    task: str = "kb_task",
) -> str | None:
    """Call LLM via the shared fallback chain (convenience alias)."""
    return _call_llm_with_fallback(
        messages, temperature=temperature, max_tokens=max_tokens, task=task
    )


def _kb_entity_collection(kb: dict) -> str:
    """Qdrant collection name for a KB's entity embeddings."""
    return f"{kb['qdrant_collection']}_entities"


def _kb_wiki_collection(kb: dict) -> str:
    """Qdrant collection name for a KB's wiki page embeddings."""
    return f"{kb['qdrant_collection']}_wiki"


def _ensure_qdrant_collection(collection_name: str) -> None:
    """Create a cosine 1024-dim Qdrant collection if it does not exist."""
    from qdrant_client.models import VectorParams, Distance

    client = _qdrant_client()
    try:
        client.get_collection(collection_name)
    except Exception:
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=1024, distance=Distance.COSINE),
        )


def _qdrant_upsert(collection_name: str, points: list) -> None:
    """Upsert points into a Qdrant collection."""
    if not points:
        return
    _qdrant_client().upsert(collection_name=collection_name, points=points)


_qdrant_client_cache = None
_qdrant_client_key = None

def _qdrant_client():
    """Return a cached QdrantClient, rebuilt when host/port/key change."""
    global _qdrant_client_cache, _qdrant_client_key
    from rag.config import qdrant_api_key, qdrant_client_kwargs, qdrant_grpc_port, qdrant_host

    key = (qdrant_host(), qdrant_grpc_port(), qdrant_api_key())
    if _qdrant_client_cache is not None and _qdrant_client_key == key:
        return _qdrant_client_cache
    from qdrant_client import QdrantClient
    _qdrant_client_key = key
    _qdrant_client_cache = QdrantClient(**qdrant_client_kwargs())
    return _qdrant_client_cache


def _qdrant_delete_by_filter(collection_name: str, conditions: list) -> None:
    """Delete points matching filter conditions. Best-effort (collection may not exist)."""
    from qdrant_client.models import Filter
    try:
        _qdrant_client().delete(collection_name=collection_name, points_selector=Filter(must=conditions))
    except Exception:
        logger.warning("Qdrant delete-by-filter failed on %s", collection_name, exc_info=True)


def _qdrant_delete_points(collection_name: str, point_ids: list[str]) -> None:
    """Delete explicit point ids. Best-effort."""
    from qdrant_client.models import PointIdsList
    if not point_ids:
        return
    try:
        _qdrant_client().delete(collection_name=collection_name, points_selector=PointIdsList(points=point_ids))
    except Exception:
        logger.warning("Qdrant delete-points failed on %s", collection_name, exc_info=True)


def _qdrant_drop_collection(name: str) -> None:
    """Drop an entire collection. Best-effort."""
    try:
        _qdrant_client().delete_collection(name)
    except Exception:
        logger.debug("Failed to drop Qdrant collection %s (may not exist)", name)


def _chunk_point_id(doc_id: str, chunk_index: int) -> str:
    """Deterministic Qdrant point id for a chunk slot (idempotent re-vectorize)."""
    return str(uuid.uuid5(uuid.NAMESPACE_URL, f"kb-chunk:{doc_id}:{chunk_index}"))


def _purge_doc_vectors(kb_id: str, doc: dict) -> None:
    """Delete ALL Qdrant points belonging to a document.

    Covers: chunk points (payload filename), summary points (payload doc_id),
    entity points and wiki page points (payload doc_id), across the KB's three
    collections. Called on document deletion and before re-vectorization.
    """
    from qdrant_client.models import FieldCondition, MatchValue

    kb = get_knowledge_base(kb_id)
    doc_id = doc["id"]
    chunk_coll = kb["qdrant_collection"]
    # Only filter by doc_id. Filtering by filename would accidentally delete
    # chunks from other documents with the same basename in different folders.
    _qdrant_delete_by_filter(chunk_coll, [FieldCondition(key="doc_id", match=MatchValue(value=doc_id))])
    _qdrant_delete_by_filter(_kb_entity_collection(kb), [FieldCondition(key="doc_id", match=MatchValue(value=doc_id))])
    _qdrant_delete_by_filter(_kb_wiki_collection(kb), [FieldCondition(key="doc_id", match=MatchValue(value=doc_id))])


def _get_embedding(texts: list[str]) -> list[list[float]]:
    """Embed a list of texts using DashScope text-embedding-v3."""
    from rag.embedding import embed_texts
    return embed_texts(texts)


def _extract_json_block(text: str) -> str:
    """Extract JSON from a markdown code block or raw string."""
    text = text.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:].strip()
    return text


def _llm_extract_graph(text: str, context: str = "") -> dict:
    """Extract entities and relationships from text for GraphRAG."""
    prompt = (
        "从以下文本中提取知识图谱实体和关系。输出严格合法的 JSON，不要包含任何解释。\n"
        "JSON 格式：\n"
        "{\"entities\": [{\"name\": \"实体名称\", \"type\": \"实体类型\", \"description\": \"简短描述\"}], "
        "\"relationships\": [{\"source\": \"源实体\", \"target\": \"目标实体\", \"type\": \"关系类型\", \"description\": \"简短描述\"}]}\n\n"
        f"上下文：{context}\n\n文本：\n{text[:6000]}"
    )
    messages = [
        {"role": "system", "content": "你是一个知识图谱抽取助手，只输出 JSON。"},
        {"role": "user", "content": prompt},
    ]
    raw = _call_llm(messages, temperature=0.1, max_tokens=1500, task="kb_graph")
    if not raw:
        return {"entities": [], "relationships": []}
    try:
        return json.loads(_extract_json_block(raw))
    except Exception:
        logger.exception("Graph extraction JSON parse failed: %s", raw[:200])
        return {"entities": [], "relationships": []}


def _llm_generate_wiki(text: str, title: str = "", temperature: float = 0.4, max_tokens: int = 1500) -> dict:
    """Generate a wiki-style page from text."""
    prompt = (
        "根据以下文本生成一篇维基百科风格的知识库条目。输出严格合法的 JSON，不要包含任何解释。\n"
        "JSON 格式：\n"
        "{\"title\": \"条目标题\", \"content\": \"条目正文（Markdown 格式，300-800 字）\"}\n\n"
        f"文档标题：{title}\n\n文本：\n{text[:8000]}"
    )
    messages = [
        {"role": "system", "content": "你是一个知识库维基条目生成助手，只输出 JSON。"},
        {"role": "user", "content": prompt},
    ]
    raw = _call_llm(messages, temperature=temperature, max_tokens=max_tokens, task="kb_wiki")
    if not raw:
        return {"title": title or "Wiki", "content": ""}
    try:
        return json.loads(_extract_json_block(raw))
    except Exception:
        logger.exception("Wiki generation JSON parse failed: %s", raw[:200])
        return {"title": title or "Wiki", "content": raw[:2000]}


def _llm_generate_folder_wiki(text: str, title: str = "", folder_path: str = "", temperature: float = 0.3, max_tokens: int = 2000) -> dict:
    """Generate a folder-level wiki page synthesizing multiple documents."""
    prompt = (
        "你是一个知识库架构师。以下是一个目录（可能包含子目录）中多篇文档的内容节选。"
        "请基于这些文档生成一篇主题性的维基百科风格知识库条目，整合共性内容、梳理结构、提炼核心概念，"
        "而不是简单罗列每篇文档。输出严格合法的 JSON，不要包含任何解释。\n"
        "JSON 格式：\n"
        "{\"title\": \"条目标题\", \"content\": \"条目正文（Markdown 格式，800-1500 字）\"}\n\n"
        f"目录路径：{folder_path or '根目录'}\n"
        f"建议标题：{title}\n\n"
        f"文档内容：\n\n{text[:24000]}"
    )
    messages = [
        {"role": "system", "content": "你是一个知识库主题整合专家，只输出 JSON。"},
        {"role": "user", "content": prompt},
    ]
    raw = _call_llm(messages, temperature=temperature, max_tokens=max_tokens, task="kb_folder_wiki")
    if not raw:
        return {"title": title or "Folder Wiki", "content": ""}
    try:
        return json.loads(_extract_json_block(raw))
    except Exception:
        logger.exception("Folder wiki generation JSON parse failed: %s", raw[:200])
        return {"title": title or "Folder Wiki", "content": raw[:2000]}


def get_document_file(doc_id: str) -> tuple[Path, str, str]:
    """Resolve a document's on-disk file for raw binary streaming.

    Returns (absolute_path, mime_type, file_name). Used by the adaptive
    per-type preview (PDF iframe, <img>, <audio>, docx/xlsx renderers).
    """
    doc = get_knowledge_document(doc_id)
    target = _resolve_path(doc["kb_id"], doc["file_path"], must_exist=True)
    mime = mimetypes.guess_type(doc["file_name"])[0] or "application/octet-stream"
    return target, mime, doc["file_name"]


_IMAGE_SUFFIXES = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".svg", ".webp"}
_FILE_PAYLOAD_MAX_BYTES = 15 * 1024 * 1024


def get_document_file_payload(doc_id: str, max_bytes: int = _FILE_PAYLOAD_MAX_BYTES) -> dict:
    """JSON envelope for desktop IPC, which cannot stream FileResponse bytes."""
    path, mime, name = get_document_file(doc_id)
    size = path.stat().st_size
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        kind = "pdf"
    elif suffix in _IMAGE_SUFFIXES:
        kind = "image"
    else:
        kind = "binary"
    cap = max(_FILE_PAYLOAD_MAX_BYTES if max_bytes <= 0 else min(max_bytes, 50 * 1024 * 1024), 1)
    if size > cap:
        return {
            "filename": name,
            "kind": kind,
            "mime": mime,
            "size": size,
            "too_large": True,
        }
    return {
        "data": base64.b64encode(path.read_bytes()).decode("ascii"),
        "filename": name,
        "kind": kind,
        "mime": mime,
        "size": size,
        "too_large": False,
    }


_pipeline_locks: dict[tuple[str, str], threading.Lock] = {}
_pipeline_locks_guard = threading.Lock()


def _pipeline_lock_for(kind: str, key: str) -> threading.Lock:
    with _pipeline_locks_guard:
        lock = _pipeline_locks.get((kind, key))
        if lock is None:
            lock = threading.Lock()
            _pipeline_locks[(kind, key)] = lock
        return lock


def _run_pipeline_job(kind: str, key: str, fn, *args, **kwargs) -> None:
    """Run fn under a per-(kind, key) lock. Skip if that job is already in flight."""
    lock = _pipeline_lock_for(kind, key)
    if not lock.acquire(blocking=False):
        logger.info("%s already running for %s; skip duplicate", kind, key)
        return
    try:
        fn(*args, **kwargs)
    except Exception:
        logger.exception("%s failed for %s", kind, key)
    finally:
        lock.release()


def _require_kb_folder(kb_id: str, folder_id: str | None) -> None:
    get_knowledge_base(kb_id)
    if not folder_id:
        return
    row = _get_db().execute(
        "SELECT id FROM knowledge_folders WHERE id = ? AND kb_id = ?",
        (folder_id, kb_id),
    ).fetchone()
    if not row:
        raise ValueError("Folder not found")


def generate_document_summary(doc_id: str) -> dict:
    """Generate a document summary using LLM if available."""
    doc = get_knowledge_document(doc_id)
    db = _get_db()
    db.execute(
        "UPDATE knowledge_documents SET summary_status = ?, updated_at = ? WHERE id = ?",
        ("processing", _now_iso(), doc_id),
    )
    db.commit()

    preview = get_knowledge_file_preview_v2(doc_id, max_chars=10_000)
    content = preview["content"]
    if len(content) < 100:
        db.execute(
            "UPDATE knowledge_documents SET summary_text = ?, summary_status = ?, updated_at = ? WHERE id = ?",
            ("", "skipped", _now_iso(), doc_id),
        )
        db.commit()
        return {"id": doc_id, "summary": "", "status": "skipped"}

    summary = _llm_summarize(content)
    if summary:
        status = "completed"
    else:
        summary = f"【摘要占位】该文档为 {doc['file_name']}，共 {preview['lines']} 行，{len(content)} 字符。"
        status = "completed_no_llm"

    db.execute(
        "UPDATE knowledge_documents SET summary_text = ?, summary_status = ?, updated_at = ? WHERE id = ?",
        (summary, status, _now_iso(), doc_id),
    )
    db.commit()

    # WeKnora-style: the summary itself is indexed so doc-level semantic
    # search can match whole-document meaning, not just chunk fragments.
    if status == "completed" and summary.strip():
        try:
            _embed_doc_summary(doc, summary)
        except Exception:
            logger.exception("Failed to index summary for doc %s", doc_id)

    return {"id": doc_id, "summary": summary, "status": status}


def _run_summary_job(doc_id: str) -> None:
    _run_pipeline_job("summary", doc_id, generate_document_summary, doc_id)


def start_summary_build(doc_id: str) -> dict:
    """Enqueue document summarization and return immediately."""
    get_knowledge_document(doc_id)
    _get_postprocess_executor().submit(_run_summary_job, doc_id)
    return {"id": doc_id, "status": "processing"}


def _embed_doc_summary(doc: dict, summary: str) -> None:
    """Embed a document summary into the KB's chunk collection as a doc-level
    point. Uses a deterministic UUID so re-summarization overwrites in place."""
    from qdrant_client.models import PointStruct

    kb = get_knowledge_base(doc["kb_id"])
    collection = kb["qdrant_collection"]
    _ensure_qdrant_collection(collection)
    vector = _get_embedding([summary])[0]
    point_id = str(uuid.uuid5(uuid.NAMESPACE_URL, f"kb-summary:{doc['id']}"))
    _qdrant_upsert(
        collection,
        [
            PointStruct(
                id=point_id,
                vector=vector,
                payload={
                    "filename": doc["file_name"],
                    "file_path": doc["file_path"],
                    "file_type": Path(doc["file_name"]).suffix.lower().lstrip("."),
                    "text": summary,
                    "chapter": "文档摘要",
                    "source": "summary",
                    "doc_id": doc["id"],
                },
            )
        ],
    )


def build_document_graph(doc_id: str) -> dict:
    """Extract and store a knowledge graph from a document's chunks."""
    doc = get_knowledge_document(doc_id)
    kb_id = doc["kb_id"]
    db = _get_db()

    # Remove old graph data for this doc
    db.execute("DELETE FROM knowledge_relationships WHERE doc_id = ?", (doc_id,))
    db.execute("DELETE FROM knowledge_entities WHERE doc_id = ?", (doc_id,))
    db.commit()

    chunks = db.execute(
        "SELECT id, content FROM knowledge_chunks WHERE doc_id = ? AND is_enabled = 1 ORDER BY chunk_index",
        (doc_id,),
    ).fetchall()

    entity_map: dict[str, str] = {}
    rel_count = 0
    for chunk in chunks:
        extracted = _llm_extract_graph(chunk["content"], context=doc["file_name"])
        for ent in extracted.get("entities", []):
            name = (ent.get("name") or "").strip()
            if not name or name in entity_map:
                continue
            entity_id = _new_id()
            entity_map[name] = entity_id
            db.execute(
                """
                INSERT INTO knowledge_entities
                (id, kb_id, doc_id, name, entity_type, description, metadata, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    entity_id,
                    kb_id,
                    doc_id,
                    name,
                    ent.get("type", ""),
                    ent.get("description", ""),
                    _json_dumps(ent),
                    _now_iso(),
                ),
            )
        for rel in extracted.get("relationships", []):
            src = (rel.get("source") or "").strip()
            tgt = (rel.get("target") or "").strip()
            if not src or not tgt:
                continue
            # Ensure both entities exist (create stubs if missing)
            for name in (src, tgt):
                if name not in entity_map:
                    entity_id = _new_id()
                    entity_map[name] = entity_id
                    db.execute(
                        """
                        INSERT INTO knowledge_entities
                        (id, kb_id, doc_id, name, entity_type, description, metadata, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (entity_id, kb_id, doc_id, name, "", "", "{}", _now_iso()),
                    )
            db.execute(
                """
                INSERT INTO knowledge_relationships
                (id, kb_id, doc_id, source_entity_id, target_entity_id, relation_type, description, metadata, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    _new_id(),
                    kb_id,
                    doc_id,
                    entity_map[src],
                    entity_map[tgt],
                    rel.get("type", ""),
                    rel.get("description", ""),
                    _json_dumps(rel),
                    _now_iso(),
                ),
            )
            rel_count += 1
        db.commit()

    # Embed entities and upsert to Qdrant for semantic graph search.
    # SQLite graph rows are already committed; a down Qdrant must not
    # roll the extraction back (or 500 a waiting HTTP client).
    try:
        _embed_kb_entities(kb_id, doc_id)
    except Exception:
        logger.exception("Failed to embed graph entities for doc %s", doc_id)

    return {"id": doc_id, "entities": len(entity_map), "relationships": rel_count}


def _run_graph_job(doc_id: str) -> None:
    _run_pipeline_job("graph", doc_id, build_document_graph, doc_id)


def start_graph_build(doc_id: str) -> dict:
    """Enqueue graph extraction and return immediately.

    Per-chunk LLM calls routinely exceed the desktop's 180s pipeline
    timeout when run on the HTTP request thread.
    """
    get_knowledge_document(doc_id)
    _get_postprocess_executor().submit(_run_graph_job, doc_id)
    return {"id": doc_id, "status": "processing"}


def _embed_kb_entities(kb_id: str, doc_id: str) -> None:
    """Generate embeddings for entities of a document and store them in Qdrant."""
    from qdrant_client.models import FieldCondition, MatchValue, PointStruct

    kb = get_knowledge_base(kb_id)
    collection = _kb_entity_collection(kb)
    _ensure_qdrant_collection(collection)

    doc = get_knowledge_document(doc_id)

    # Purge stale entity vectors for this document before re-embedding.
    _qdrant_delete_by_filter(
        collection, [FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]
    )

    db = _get_db()
    rows = db.execute(
        "SELECT id, name, entity_type, description FROM knowledge_entities WHERE doc_id = ?",
        (doc_id,),
    ).fetchall()
    if not rows:
        return

    texts = [f"{r['name']} {r['entity_type']} {r['description']}".strip() for r in rows]
    vectors = _get_embedding(texts)

    points = []
    for idx, (row, vector) in enumerate(zip(rows, vectors)):
        point_id = str(uuid.uuid4())
        points.append(
            PointStruct(
                id=point_id,
                vector=vector,
                payload={
                    "entity_id": row["id"],
                    "name": row["name"],
                    "entity_type": row["entity_type"],
                    "description": row["description"],
                    "text": texts[idx],
                    "doc_id": doc_id,
                    "file_path": doc["file_path"],
                    "file_type": Path(doc["file_name"]).suffix.lower().lstrip("."),
                    "source": "entity",
                },
            )
        )
        db.execute(
            "UPDATE knowledge_entities SET qdrant_point_id = ? WHERE id = ?",
            (point_id, row["id"]),
        )
    _qdrant_upsert(collection, points)
    db.commit()


def generate_document_wiki(doc_id: str, curate: bool = False) -> dict:
    """Generate and store a wiki-style page for a document.

    Args:
        curate: when True, after exporting to the llm-wiki vault also launch a
            headless agent running the llm-wiki skill to curate the generated
            page (cross-references, contradiction checks, entity merging).
    """
    doc = get_knowledge_document(doc_id)
    kb_id = doc["kb_id"]

    # Incremental update: skip if an up-to-date wiki already exists.
    source_hash = _compute_doc_source_hash(doc_id)
    db = _get_db()
    existing = db.execute(
        """
        SELECT id, title, slug, content, status, source_hash
        FROM knowledge_wiki_pages
        WHERE doc_id = ? AND status = 'completed'
        ORDER BY updated_at DESC
        LIMIT 1
        """,
        (doc_id,),
    ).fetchone()
    if existing and existing["source_hash"] == source_hash:
        return {
            "id": doc_id,
            "wiki_id": existing["id"],
            "title": existing["title"],
            "skipped": True,
            "reason": "source unchanged",
        }

    preview = get_knowledge_file_preview_v2(doc_id, max_chars=12_000)
    content = preview["content"]

    db.execute("DELETE FROM knowledge_wiki_pages WHERE doc_id = ?", (doc_id,))
    db.commit()

    kb = get_knowledge_base(kb_id)
    cfg = _get_kb_curation_config(kb)
    wiki = _llm_generate_wiki(
        content,
        title=doc.get("title") or doc["file_name"],
        temperature=cfg["wiki_temperature"],
        max_tokens=cfg["wiki_max_tokens"],
    )
    wiki_title = wiki.get("title", doc["file_name"])
    wiki_content = wiki.get("content", "")
    wiki_id = _new_id()
    now = _now_iso()

    # Quality evaluation and threshold-based review gating.
    quality = _evaluate_wiki_quality(
        wiki_id, wiki_title, wiki_content, source_text=content,
        temperature=cfg["curator_temperature"],
    )
    review_status = "approved" if quality["score"] >= cfg["auto_approve_threshold"] else "pending"

    db.execute(
        """
        INSERT INTO knowledge_wiki_pages
        (id, kb_id, doc_id, title, slug, content, status, review_status, source_hash, quality_score, quality_report, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            wiki_id,
            kb_id,
            doc_id,
            wiki_title,
            _slugify(wiki_title),
            wiki_content,
            "completed",
            review_status,
            source_hash,
            quality["score"],
            _json_dumps(quality),
            now,
            now,
        ),
    )
    db.commit()

    # Embed wiki page and store in Qdrant
    _embed_kb_wiki_page(kb_id, wiki_id, wiki_title, wiki_content, doc_id)

    # One-way export into the llm-wiki vault (generated/ subtree)
    folder_path = _get_folder_path(kb_id, doc.get("folder_id"))
    vault_info = _export_wiki_to_vault(
        kb,
        wiki_id=wiki_id,
        title=wiki_title,
        slug=_slugify(wiki_title),
        content=wiki_content,
        doc=doc,
        folder_path=folder_path,
    )

    result = {
        "id": doc_id,
        "wiki_id": wiki_id,
        "title": wiki_title,
        "review_status": review_status,
        "quality_score": quality["score"],
    }
    if vault_info:
        result["vault_path"] = vault_info["vault_path"]

    # Optional: synchronously trigger llm-wiki curation (cross-referencing,
    # contradiction checks, entity-page merge) via a headless agent.
    if curate and vault_info:
        _start_wiki_curation(kb_id, [vault_info["vault_path"]])
        result["curation"] = "started"

    return result


def _run_doc_wiki_job(doc_id: str, curate: bool = False) -> None:
    _run_pipeline_job("wiki", doc_id, generate_document_wiki, doc_id, curate)


def start_wiki_build(doc_id: str, curate: bool = False) -> dict:
    """Enqueue document wiki generation and return immediately."""
    get_knowledge_document(doc_id)
    _get_postprocess_executor().submit(_run_doc_wiki_job, doc_id, curate)
    return {"id": doc_id, "status": "processing"}


def generate_folder_wiki(kb_id: str, folder_id: str | None, title: str = "", curate: bool = False) -> dict:
    """Generate and store a folder-level wiki page synthesizing its documents.

    Args:
        folder_id: target folder; None means the KB root folder.
        title: optional override title; defaults to folder name or KB name.
        curate: when True, trigger llm-wiki curation after export.
    """
    kb = get_knowledge_base(kb_id)
    folder_path = _get_folder_path(kb_id, folder_id)
    default_title = folder_path.split("/")[-1] if folder_path else kb["name"]
    combined, used_docs = _collect_folder_contents(kb_id, folder_id, max_total_chars=24_000)
    if not combined:
        raise ValueError("Folder has no parseable documents")

    db = _get_db()

    # Incremental update: skip if source documents have not changed.
    source_hash = _compute_folder_wiki_source_hash(kb_id, folder_id, recursive=True)
    existing = db.execute(
        """
        SELECT id, title, slug, content, status, source_hash
        FROM knowledge_wiki_pages
        WHERE kb_id = ? AND folder_id IS ? AND doc_id IS NULL AND status = 'completed'
        ORDER BY updated_at DESC
        LIMIT 1
        """,
        (kb_id, folder_id),
    ).fetchone()
    if existing and existing["source_hash"] == source_hash:
        return {
            "kb_id": kb_id,
            "folder_id": folder_id,
            "wiki_id": existing["id"],
            "title": existing["title"],
            "doc_count": len(used_docs),
            "skipped": True,
            "reason": "source unchanged",
        }

    # Remove previous folder wiki for this folder (if any)
    if folder_id:
        db.execute("DELETE FROM knowledge_wiki_pages WHERE kb_id = ? AND folder_id = ?", (kb_id, folder_id))
    else:
        db.execute(
            "DELETE FROM knowledge_wiki_pages WHERE kb_id = ? AND folder_id IS NULL AND doc_id IS NULL",
            (kb_id,),
        )
    db.commit()

    cfg = _get_kb_curation_config(kb)
    wiki = _llm_generate_folder_wiki(
        combined,
        title=title or default_title,
        folder_path=folder_path,
        temperature=cfg["wiki_temperature"],
        max_tokens=cfg["wiki_max_tokens"],
    )
    wiki_title = wiki.get("title") or title or default_title
    wiki_slug = _slugify(wiki_title) or "folder-wiki"
    wiki_content = wiki.get("content", "")
    wiki_id = _new_id()
    now = _now_iso()
    source_hash = _compute_folder_wiki_source_hash(kb_id, folder_id, recursive=True)

    quality = _evaluate_wiki_quality(
        wiki_id, wiki_title, wiki_content, source_text=combined,
        temperature=cfg["curator_temperature"],
    )
    review_status = "approved" if quality["score"] >= cfg["auto_approve_threshold"] else "pending"

    db.execute(
        """
        INSERT INTO knowledge_wiki_pages
        (id, kb_id, folder_id, title, slug, content, status, review_status, source_hash, quality_score, quality_report, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (wiki_id, kb_id, folder_id, wiki_title, wiki_slug, wiki_content, "completed", review_status, source_hash, quality["score"], _json_dumps(quality), now, now),
    )
    db.commit()

    # Embed folder wiki page and store in Qdrant
    _embed_kb_wiki_page(kb_id, wiki_id, wiki_title, wiki_content, doc_id=None)

    # One-way export into the llm-wiki vault (generated/ subtree)
    vault_info = _export_wiki_to_vault(
        kb,
        wiki_id=wiki_id,
        title=wiki_title,
        slug=wiki_slug,
        content=wiki_content,
        doc=None,
        folder_id=folder_id,
        folder_path=folder_path,
    )

    result = {
        "kb_id": kb_id,
        "folder_id": folder_id,
        "wiki_id": wiki_id,
        "title": wiki_title,
        "doc_count": len(used_docs),
        "review_status": review_status,
        "quality_score": quality["score"],
    }
    if vault_info:
        result["vault_path"] = vault_info["vault_path"]

    if curate and vault_info:
        _start_wiki_curation(kb_id, [vault_info["vault_path"]])
        result["curation"] = "started"

    return result


def _collect_hierarchical_folder_contents(
    kb_id: str,
    folder_id: str | None,
    child_wikis: list[dict],
    max_total_chars: int = 24_000,
) -> tuple[str, list[dict], int]:
    """Collect direct documents + child wiki summaries for hierarchical synthesis.

    Returns (combined_text, used_doc_items, used_child_count).
    """
    doc_ids = _get_folder_doc_ids(kb_id, folder_id, recursive=False)
    chunks: list[dict] = []

    # Child folder wiki summaries come first — they are structured sub-overviews.
    for child in child_wikis:
        summary = (child.get("content") or "").strip()[:2_000]
        if summary:
            chunks.append({
                "type": "child_wiki",
                "title": child.get("title", ""),
                "text": f"## 子目录：{child.get('title', '')}\n\n{summary}",
            })

    for doc_id in doc_ids:
        doc = get_knowledge_document(doc_id)
        try:
            preview = get_knowledge_file_preview_v2(doc_id, max_chars=2_000)
            text = preview.get("content", "").strip()
            if text:
                chunks.append({
                    "type": "doc",
                    "doc_id": doc_id,
                    "file_name": doc["file_name"],
                    "text": f"## 文档：{doc['file_name']}\n\n{text}",
                })
        except Exception:
            logger.exception("Failed to preview doc %s for hierarchical wiki", doc_id)

    if not chunks:
        return "", [], 0

    parts: list[str] = []
    total = 0
    used_docs: list[dict] = []
    used_child_count = 0
    for item in chunks:
        if total >= max_total_chars:
            break
        room = max_total_chars - total
        snippet = item["text"][:room]
        parts.append(snippet)
        total += len(snippet)
        if item["type"] == "doc":
            used_docs.append(item)
        else:
            used_child_count += 1

    return "\n\n".join(parts), used_docs, used_child_count


def _generate_single_folder_wiki(
    kb: dict,
    folder_id: str | None,
    folder_wiki_map: dict[str | None, dict],
    curate: bool = False,
) -> dict | None:
    """Generate one folder wiki from direct docs + child folder wikis.

    Returns the generated result dict or None if there is no content.
    """
    kb_id = kb["id"]
    db = _get_db()
    child_rows = db.execute(
        "SELECT id FROM knowledge_folders WHERE kb_id = ? AND parent_id = ?",
        (kb_id, folder_id),
    ).fetchall()
    child_ids = [r["id"] for r in child_rows]
    child_wikis = [folder_wiki_map[cid] for cid in child_ids if cid in folder_wiki_map]

    combined, used_docs, used_child_count = _collect_hierarchical_folder_contents(
        kb_id, folder_id, child_wikis, max_total_chars=24_000
    )
    if not combined:
        return None

    folder_path = _get_folder_path(kb_id, folder_id)
    default_title = folder_path.split("/")[-1] if folder_path else kb["name"]

    # Source hash: direct docs + child wiki ids.
    direct_doc_ids = _get_folder_doc_ids(kb_id, folder_id, recursive=False)
    source_parts: list[str] = []
    for doc_id in direct_doc_ids:
        source_parts.append(f"{doc_id}:{_compute_doc_source_hash(doc_id)}")
    source_parts.extend([cw.get("wiki_id", "") for cw in child_wikis])
    source_parts.sort()
    h = hashlib.sha256()
    h.update("\n".join(source_parts).encode("utf-8"))
    source_hash = h.hexdigest()

    # Incremental update: skip if an up-to-date folder wiki already exists.
    existing = db.execute(
        """
        SELECT id, title, slug, content, status, source_hash
        FROM knowledge_wiki_pages
        WHERE kb_id = ? AND folder_id IS ? AND doc_id IS NULL AND status = 'completed'
        ORDER BY updated_at DESC
        LIMIT 1
        """,
        (kb_id, folder_id),
    ).fetchone()
    if existing and existing["source_hash"] == source_hash:
        folder_wiki_map[folder_id] = {
            "wiki_id": existing["id"],
            "title": existing["title"],
            "slug": existing["slug"],
            "content": existing["content"],
            "vault_path": None,
        }
        return {
            "kb_id": kb_id,
            "folder_id": folder_id,
            "wiki_id": existing["id"],
            "title": existing["title"],
            "doc_count": len(used_docs),
            "child_wiki_count": used_child_count,
            "skipped": True,
            "reason": "source unchanged",
        }

    # Remove previous folder wiki for this folder.
    if folder_id:
        db.execute("DELETE FROM knowledge_wiki_pages WHERE kb_id = ? AND folder_id = ?", (kb_id, folder_id))
    else:
        db.execute(
            "DELETE FROM knowledge_wiki_pages WHERE kb_id = ? AND folder_id IS NULL AND doc_id IS NULL",
            (kb_id,),
        )
    db.commit()

    cfg = _get_kb_curation_config(kb)
    wiki = _llm_generate_folder_wiki(
        combined,
        title=default_title,
        folder_path=folder_path,
        temperature=cfg["wiki_temperature"],
        max_tokens=cfg["wiki_max_tokens"],
    )
    wiki_title = wiki.get("title") or default_title
    wiki_slug = _slugify(wiki_title) or "folder-wiki"
    wiki_content = wiki.get("content", "")
    wiki_id = _new_id()
    now = _now_iso()

    quality = _evaluate_wiki_quality(
        wiki_id, wiki_title, wiki_content, source_text=combined,
        temperature=cfg["curator_temperature"],
    )
    review_status = "approved" if quality["score"] >= cfg["auto_approve_threshold"] else "pending"

    db.execute(
        """
        INSERT INTO knowledge_wiki_pages
        (id, kb_id, folder_id, title, slug, content, status, review_status, source_hash, quality_score, quality_report, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (wiki_id, kb_id, folder_id, wiki_title, wiki_slug, wiki_content, "completed", review_status, source_hash, quality["score"], _json_dumps(quality), now, now),
    )
    db.commit()

    _embed_kb_wiki_page(kb_id, wiki_id, wiki_title, wiki_content, doc_id=None)

    vault_info = _export_wiki_to_vault(
        kb,
        wiki_id=wiki_id,
        title=wiki_title,
        slug=wiki_slug,
        content=wiki_content,
        doc=None,
        folder_id=folder_id,
        folder_path=folder_path,
    )

    result = {
        "kb_id": kb_id,
        "folder_id": folder_id,
        "wiki_id": wiki_id,
        "title": wiki_title,
        "doc_count": len(used_docs),
        "child_wiki_count": used_child_count,
        "review_status": review_status,
        "quality_score": quality["score"],
    }
    if vault_info:
        result["vault_path"] = vault_info["vault_path"]

    folder_wiki_map[folder_id] = {
        "wiki_id": wiki_id,
        "title": wiki_title,
        "slug": wiki_slug,
        "content": wiki_content,
        "vault_path": vault_info["vault_path"] if vault_info else None,
    }

    if curate and vault_info:
        _start_wiki_curation(kb_id, [vault_info["vault_path"]], folder_id=folder_id)
        result["curation"] = "started"

    return result


def _run_folder_wiki_job(
    kb_id: str, folder_id: str | None, title: str = "", curate: bool = False
) -> None:
    key = f"{kb_id}:{folder_id or 'root'}"
    _run_pipeline_job("folder_wiki", key, generate_folder_wiki, kb_id, folder_id, title, curate)


def start_folder_wiki_build(
    kb_id: str, folder_id: str | None, title: str = "", curate: bool = False
) -> dict:
    """Enqueue folder wiki generation and return immediately."""
    _require_kb_folder(kb_id, folder_id)
    _get_postprocess_executor().submit(_run_folder_wiki_job, kb_id, folder_id, title, curate)
    return {"kb_id": kb_id, "folder_id": folder_id, "status": "processing"}


def generate_hierarchical_folder_wiki(
    kb_id: str,
    folder_id: str | None = None,
    curate: bool = False,
) -> dict:
    """Generate folder wikis bottom-up for a folder subtree (P2).

    For each folder (processed leaves-first), synthesize its direct documents
    plus summaries of already-generated child folder wikis. When invoked at the
    KB root, also generates a top-level KB wiki from root-level documents and
    root folder wiki summaries.

    Args:
        folder_id: target folder; None means the KB root.
        curate: trigger llm-wiki curation for each generated folder wiki.

    Returns:
        {"kb_id": str, "folder_id": str|None, "generated": [...], "total": int}
    """
    kb = get_knowledge_base(kb_id)

    if folder_id:
        all_folder_ids = _get_folder_descendant_ids(kb_id, folder_id)
    else:
        all_folder_ids = _get_folder_descendant_ids(kb_id, None)

    # BFS -> reversed gives leaves before parents.
    seen: set[str | None] = set()
    ordered: list[str | None] = []
    for fid in reversed(all_folder_ids):
        if fid not in seen:
            seen.add(fid)
            ordered.append(fid)

    generated: list[dict] = []
    folder_wiki_map: dict[str | None, dict] = {}

    for fid in ordered:
        result = _generate_single_folder_wiki(kb, fid, folder_wiki_map, curate=curate)
        if result:
            generated.append(result)

    # At KB root, also synthesize a top-level wiki from root docs + root folder summaries.
    if folder_id is None:
        root_result = _generate_single_folder_wiki(kb, None, folder_wiki_map, curate=curate)
        if root_result:
            generated.append(root_result)

    return {
        "kb_id": kb_id,
        "folder_id": folder_id,
        "generated": generated,
        "total": len(generated),
    }


def _run_hierarchical_wiki_job(kb_id: str, folder_id: str | None, curate: bool = False) -> None:
    key = f"{kb_id}:{folder_id or 'root'}"
    _run_pipeline_job(
        "folder_wiki",
        key,
        generate_hierarchical_folder_wiki,
        kb_id,
        folder_id,
        curate,
    )


def start_hierarchical_wiki_build(
    kb_id: str, folder_id: str | None = None, curate: bool = False
) -> dict:
    """Enqueue hierarchical folder-wiki generation and return immediately."""
    _require_kb_folder(kb_id, folder_id)
    _get_postprocess_executor().submit(_run_hierarchical_wiki_job, kb_id, folder_id, curate)
    return {"kb_id": kb_id, "folder_id": folder_id, "status": "processing"}


def _embed_kb_wiki_page(kb_id: str, wiki_id: str, title: str, content: str, doc_id: str | None = None) -> None:
    """Generate embedding for a wiki page and store it in Qdrant."""
    from qdrant_client.models import PointStruct

    kb = get_knowledge_base(kb_id)
    collection = _kb_wiki_collection(kb)
    _ensure_qdrant_collection(collection)

    text = f"{title}\n\n{content}".strip()
    if not text:
        return

    db = _get_db()
    # Remove the previous embedding for this wiki page to avoid orphan vectors.
    row = db.execute(
        "SELECT qdrant_point_id FROM knowledge_wiki_pages WHERE id = ?", (wiki_id,)
    ).fetchone()
    old_point_id = row["qdrant_point_id"] if row else None
    if old_point_id:
        _qdrant_delete_points(collection, [old_point_id])

    vector = _get_embedding([text])[0]
    point_id = str(uuid.uuid4())
    payload = {
        "wiki_id": wiki_id,
        "title": title,
        "text": text,
        "source": "wiki",
    }
    if doc_id:
        doc = get_knowledge_document(doc_id)
        payload["doc_id"] = doc_id
        payload["file_path"] = doc["file_path"]
        payload["file_type"] = Path(doc["file_name"]).suffix.lower().lstrip(".")
    _qdrant_upsert(
        collection,
        [PointStruct(id=point_id, vector=vector, payload=payload)],
    )
    db.execute("UPDATE knowledge_wiki_pages SET qdrant_point_id = ? WHERE id = ?", (point_id, wiki_id))
    db.commit()


def _get_folder_path(kb_id: str, folder_id: str | None) -> str:
    """Return the stored path for a folder, or '' for root."""
    if not folder_id:
        return ""
    db = _get_db()
    row = db.execute("SELECT path FROM knowledge_folders WHERE id = ? AND kb_id = ?", (folder_id, kb_id)).fetchone()
    return (row["path"] or "") if row else ""


def _get_folder_descendant_ids(kb_id: str, folder_id: str | None) -> list[str]:
    """Return [folder_id, ...all descendant folder ids] for a KB folder tree."""
    db = _get_db()
    if not folder_id:
        rows = db.execute("SELECT id FROM knowledge_folders WHERE kb_id = ? AND parent_id IS NULL", (kb_id,)).fetchall()
    else:
        rows = db.execute("SELECT id FROM knowledge_folders WHERE id = ? AND kb_id = ?", (folder_id, kb_id)).fetchall()
    result: list[str] = []
    queue = [r["id"] for r in rows]
    while queue:
        fid = queue.pop(0)
        result.append(fid)
        children = db.execute(
            "SELECT id FROM knowledge_folders WHERE kb_id = ? AND parent_id = ?",
            (kb_id, fid),
        ).fetchall()
        queue.extend([c["id"] for c in children])
    return result


def _get_folder_doc_ids(kb_id: str, folder_id: str | None, recursive: bool = True) -> list[str]:
    """Return document ids inside a folder (optionally recursively)."""
    db = _get_db()
    if not recursive or not folder_id:
        folder_ids: list[str | None] = [folder_id] if folder_id else [None]
    else:
        folder_ids = _get_folder_descendant_ids(kb_id, folder_id)
    doc_ids: list[str] = []
    for fid in folder_ids:
        if fid is None:
            rows = db.execute(
                "SELECT id FROM knowledge_documents WHERE kb_id = ? AND folder_id IS NULL AND parse_status = 'completed'",
                (kb_id,),
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT id FROM knowledge_documents WHERE kb_id = ? AND folder_id = ? AND parse_status = 'completed'",
                (kb_id, fid),
            ).fetchall()
        doc_ids.extend([r["id"] for r in rows])
    return doc_ids


def _collect_folder_contents(kb_id: str, folder_id: str | None, max_total_chars: int = 24_000) -> tuple[str, list[dict]]:
    """Collect preview text from documents in a folder for folder-level wiki generation.

    Returns (combined_text, used_doc_items).
    """
    doc_ids = _get_folder_doc_ids(kb_id, folder_id, recursive=True)
    chunks: list[dict] = []
    for doc_id in doc_ids:
        doc = get_knowledge_document(doc_id)
        try:
            preview = get_knowledge_file_preview_v2(doc_id, max_chars=4_000)
            text = preview.get("content", "").strip()
            if text:
                chunks.append({"doc_id": doc_id, "file_name": doc["file_name"], "text": text})
        except Exception:
            logger.exception("Failed to preview doc %s for folder wiki", doc_id)
    if not chunks:
        return "", []

    parts: list[str] = []
    total = 0
    used: list[dict] = []
    for item in chunks:
        if total >= max_total_chars:
            break
        header = f"## 文档：{item['file_name']}\n\n"
        room = max_total_chars - total - len(header)
        if room <= 0:
            break
        snippet = item["text"][:room]
        parts.append(header + snippet)
        total += len(header) + len(snippet)
        used.append(item)
    return "\n\n".join(parts), used


# ── llm-wiki vault integration ────────────────────────────────────────────────

def _resolve_wiki_vault(kb: dict) -> Path | None:
    """Resolve the llm-wiki vault root for a KB.

    Priority: indexing_strategy.wiki_vault → HERMES_WIKI_VAULT env → default.
    Set wiki_vault to "off"/false to disable export for this KB.

    The resolved path is constrained to be under the Hermes home directory to
    prevent path-traversal writes via environment variables or config files.
    """
    strategy = kb.get("indexing_strategy") or {}
    vault = strategy.get("wiki_vault")
    if vault is None or vault is True:
        vault = _DEFAULT_WIKI_VAULT
    if vault is False or str(vault).strip().lower() in ("", "off", "none", "false", "disabled"):
        return None

    vault_path = Path(str(vault)).expanduser().resolve()
    hermes_home = _get_hermes_home().resolve()

    # Defensive: reject absolute paths outside Hermes home and any ".." components.
    try:
        vault_path.relative_to(hermes_home)
    except ValueError as exc:
        raise ValueError(
            f"Wiki vault path {vault_path} must be inside Hermes home {hermes_home}"
        ) from exc

    return vault_path


def _wiki_page_vault_path(kb: dict, slug: str, folder_id: str | None = None) -> Path | None:
    """Compute the expected vault path for an auto-generated wiki page.

    Mirrors the layout used by `_export_wiki_to_vault`.
    """
    vault = _resolve_wiki_vault(kb)
    if vault is None or not slug:
        return None
    kb_slug = _slugify(kb["name"]) or kb["id"]
    folder_path = _get_folder_path(kb["id"], folder_id)
    target_dir = vault / "generated" / kb_slug
    if folder_path:
        target_dir = target_dir / folder_path.replace("/", os.sep)
    return target_dir / f"{slug}.md"


def _export_wiki_to_vault(
    kb: dict,
    wiki_id: str,
    title: str,
    slug: str,
    content: str,
    doc: dict | None = None,
    folder_id: str | None = None,
    folder_path: str = "",
) -> dict | None:
    """Write a generated wiki page into the llm-wiki vault's generated/ subtree.

    Layout:
        {vault}/generated/{kb_slug}/{folder_path}/{slug}.md   — page with frontmatter
        {vault}/generated/{kb_slug}/_index.md                  — per-KB recursive catalog
        {vault}/log.md                                         — global chronological log

    This is a ONE-WAY sync (pipeline → vault). Human/agent curation of these
    pages happens later via the llm-wiki skill; curated pages live in the
    vault's main tree (entities/, concepts/, ...) and are never written back.
    """
    vault = _resolve_wiki_vault(kb)
    if vault is None:
        return None
    try:
        kb_slug = _slugify(kb["name"]) or kb["id"]
        gen_dir = vault / "generated" / kb_slug
        target_dir = gen_dir
        if folder_path:
            target_dir = gen_dir / folder_path.replace("/", os.sep)
        target_dir.mkdir(parents=True, exist_ok=True)

        today = datetime.now().strftime("%Y-%m-%d")
        page_path = target_dir / f"{slug}.md"
        fm_lines = [
            "---",
            f"title: {title}",
            f"created: {today}",
            f"updated: {today}",
            "type: wiki-auto",
            "source: auto",
            f"kb_id: {kb['id']}",
            f'kb_name: "{kb["name"]}"',
        ]
        if doc:
            fm_lines.append(f"doc_id: {doc['id']}")
            # Derive energy-audit classification tags from the filename so that
            # local wiki fallback search can filter by tags exactly via frontmatter.
            try:
                from tools.energy_audit.institution_classifier import classify_institution
                institution_category, specific_type = classify_institution(doc["file_name"])
                fm_lines.extend([
                    "audit_type: 公共机构",
                    f"institution_category: {institution_category}",
                    f"specific_type: {specific_type}",
                ])
            except Exception:
                pass
        if folder_id:
            fm_lines.append(f"folder_id: {folder_id}")
        fm_lines.extend([
            f"wiki_id: {wiki_id}",
            "tags: [auto-generated]",
        ])
        if doc:
            fm_lines.append(f"sources: [\"{doc['file_name']}\"]")
        else:
            fm_lines.append(f"sources: [\"folder:{folder_path or 'root'}\"]")
        fm_lines.extend([
            "confidence: medium",
            "---",
            "",
        ])
        page_path.write_text("\n".join(fm_lines) + content.strip() + "\n", encoding="utf-8")

        _rebuild_generated_index(vault, gen_dir, kb)
        rel_page = page_path.relative_to(gen_dir).as_posix()
        log_detail = ["- source: auto (WebUI KB pipeline)", f"- page: generated/{kb_slug}/{rel_page}"]
        if doc:
            log_detail.append(f"- doc: {doc['file_name']}")
        if folder_id:
            log_detail.append(f"- folder: {folder_path or 'root'}")
        _append_vault_log(vault, f"ingest | {title}", log_detail)
        return {"vault_path": str(page_path), "kb_slug": kb_slug}
    except Exception:
        logger.exception("Wiki vault export failed")
        return None


def _rebuild_generated_index(vault: Path, gen_dir: Path, kb: dict) -> None:
    """Regenerate the per-KB _index.md from all generated pages (including subfolders)."""
    pages = sorted(p for p in gen_dir.rglob("*.md") if p.name != "_index.md")
    lines = [
        f"# Generated Wiki Pages — {kb['name']}",
        "",
        "> Auto-generated by the WebUI KB pipeline. Curate into the main wiki",
        "> (entities/, concepts/, ...) via the llm-wiki skill.",
        "",
        f"> Total pages: {len(pages)}",
        "",
    ]
    for p in pages:
        page_title = p.stem
        try:
            head = p.read_text(encoding="utf-8")[:800]
            m = re.search(r"^title:\s*(.+)$", head, re.MULTILINE)
            if m:
                page_title = m.group(1).strip().strip('"').strip("'")
        except Exception:
            pass
        rel = p.relative_to(gen_dir).with_suffix("").as_posix()
        lines.append(f"- [[{rel}]] — {page_title}")
    (gen_dir / "_index.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def _append_vault_log(vault: Path, action: str, detail_lines: list[str]) -> None:
    """Append an entry to the vault's log.md (create with header if missing)."""
    log_path = vault / "log.md"
    if not log_path.exists():
        log_path.parent.mkdir(parents=True, exist_ok=True)
        log_path.write_text(
            "# Wiki Log\n\n> Chronological record of all wiki actions. Append-only.\n\n",
            encoding="utf-8",
        )
    today = datetime.now().strftime("%Y-%m-%d")
    entry = f"## [{today}] {action}\n" + "\n".join(detail_lines) + "\n\n"
    with log_path.open("a", encoding="utf-8") as f:
        f.write(entry)


def _start_wiki_curation(kb_id: str, page_paths: list[str], folder_id: str | None = None) -> str:
    """Launch llm-wiki curation in a background daemon thread.

    Creates a persistent job record so callers can poll for status.
    Returns the job id.
    """
    job_id = _new_id()
    db = _get_db()
    db.execute(
        """
        INSERT INTO knowledge_curation_jobs
        (id, kb_id, folder_id, job_type, status, input_pages, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (job_id, kb_id, folder_id, "curation", "pending", _json_dumps(page_paths), _now_iso()),
    )
    db.commit()
    thread = threading.Thread(
        target=_run_wiki_curation,
        args=(kb_id, page_paths, job_id),
        kwargs={"folder_id": folder_id},
        daemon=True,
    )
    thread.start()
    return job_id


def _run_wiki_curation(kb_id: str, page_paths: list[str], job_id: str, folder_id: str | None = None) -> None:
    """Run a headless agent with the llm-wiki skill to curate generated pages.

    Mirrors hermes_cli/oneshot.py::_run_agent: config + provider resolution,
    AIAgent with file toolsets, skill content injected as the prompt.
    Best-effort — failures are logged, never raised.
    """
    db = _get_db()
    db.execute(
        "UPDATE knowledge_curation_jobs SET status = ?, started_at = ? WHERE id = ?",
        ("running", _now_iso(), job_id),
    )
    db.commit()
    try:
        kb = get_knowledge_base(kb_id)
        vault = _resolve_wiki_vault(kb)
        if vault is None or not page_paths:
            db.execute(
                "UPDATE knowledge_curation_jobs SET status = ?, completed_at = ? WHERE id = ?",
                ("completed", _now_iso(), job_id),
            )
            db.commit()
            return

        from agent.skill_commands import build_skill_invocation_message
        from hermes_cli.config import load_config
        from hermes_cli.runtime_provider import resolve_runtime_provider
        from hermes_cli.tools_config import _get_platform_tools
        from hermes_cli.fallback_config import get_fallback_chain
        from run_agent import AIAgent

        try:
            from hermes_state import SessionDB
            session_db = SessionDB()
        except Exception:
            session_db = None

        pages_list = "\n".join(f"- {p}" for p in page_paths)
        instruction = (
            f"The WebUI KB pipeline just generated new wiki page(s) under:\n{pages_list}\n\n"
            f"The wiki vault is at: {vault}\n"
            "Follow the skill's orientation steps (read SCHEMA.md, index.md, recent log.md — "
            "create SCHEMA.md/index.md customized to this wiki's domain if they don't exist yet), "
            "then curate these generated pages: merge their content into entity/concept pages "
            "per the Page Thresholds, add [[wikilinks]] cross-references (both directions where "
            "appropriate), note any contradictions with existing pages per the Update Policy, "
            "and update index.md and log.md. The generated/ pages themselves are auto-managed — "
            "do not edit or delete them; your curated output goes into the main tree."
        )
        prompt = build_skill_invocation_message("/llm-wiki", user_instruction=instruction)
        if not prompt:
            logger.warning("llm-wiki skill not found; skipping curation")
            db.execute(
                "UPDATE knowledge_curation_jobs SET status = ?, completed_at = ?, error_message = ? WHERE id = ?",
                ("completed", _now_iso(), "llm-wiki skill not found", job_id),
            )
            db.commit()
            return

        cfg = load_config()
        model_cfg = cfg.get("model") or {}
        cfg_model = model_cfg if isinstance(model_cfg, str) else (model_cfg.get("default") or model_cfg.get("model") or "")
        runtime = resolve_runtime_provider(requested=None, target_model=cfg_model or None)

        toolsets_list = sorted(_get_platform_tools(cfg, "cli"))
        agent = AIAgent(
            api_key=runtime.get("api_key"),
            base_url=runtime.get("base_url"),
            provider=runtime.get("provider"),
            api_mode=runtime.get("api_mode"),
            model=cfg_model,
            enabled_toolsets=toolsets_list,
            quiet_mode=True,
            platform="webui",
            session_db=session_db,
            credential_pool=runtime.get("credential_pool"),
            fallback_model=get_fallback_chain(cfg) or None,
        )
        agent.suppress_status_output = True
        agent.stream_delta_callback = None
        agent.tool_gen_callback = None

        result = agent.run_conversation(prompt)
        final = result.get("final_response") or ""
        logger.info("llm-wiki curation finished for KB %s: %s", kb_id, final[:500])

        _append_vault_log(
            vault,
            f"curate | {len(page_paths)} generated page(s)",
            [f"- pages: {', '.join(Path(p).name for p in page_paths)}", f"- result: {final[:300]}"],
        )
        db.execute(
            "UPDATE knowledge_curation_jobs SET status = ?, completed_at = ?, output_pages = ? WHERE id = ?",
            (
                "completed",
                _now_iso(),
                _json_dumps({"input_pages": page_paths, "summary": final[:1000]}),
                job_id,
            ),
        )
        db.commit()
    except Exception as e:
        logger.exception("llm-wiki curation failed for KB %s", kb_id)
        try:
            db.execute(
                "UPDATE knowledge_curation_jobs SET status = ?, completed_at = ?, error_message = ? WHERE id = ?",
                ("failed", _now_iso(), str(e)[:2000], job_id),
            )
            db.commit()
        except Exception:
            pass


def list_curation_jobs(kb_id: str, status: str | None = None, limit: int = 50) -> list[dict]:
    """Return recent curation jobs for a KB, newest first."""
    db = _get_db()
    sql = "SELECT * FROM knowledge_curation_jobs WHERE kb_id = ?"
    params: list = [kb_id]
    if status:
        sql += " AND status = ?"
        params.append(status)
    sql += " ORDER BY created_at DESC LIMIT ?"
    params.append(limit)
    rows = db.execute(sql, params).fetchall()
    return [_curation_job_row_to_dict(r) for r in rows]


def get_curation_job(job_id: str) -> dict | None:
    """Return a single curation job by id."""
    db = _get_db()
    row = db.execute("SELECT * FROM knowledge_curation_jobs WHERE id = ?", (job_id,)).fetchone()
    return _curation_job_row_to_dict(row) if row else None


def start_global_curation(
    kb_id: str,
    folder_id: str | None = None,
    page_ids: list[str] | None = None,
    review_status: str | None = None,
) -> dict:
    """Launch a curation job over selected or all wiki pages in a KB.

    Args:
        folder_id: curate all wiki pages under this folder (recursively if
            the folder has sub-folders with their own wiki pages).
        page_ids: explicit list of wiki page ids to curate. Overrides folder_id.
        review_status: if set (e.g. "approved"), only curate pages with this
            review status. Used by P5 to let approved content gate curation.

    Returns:
        {"job_id": str, "page_count": int}
    """
    kb = get_knowledge_base(kb_id)
    vault = _resolve_wiki_vault(kb)
    if vault is None:
        raise ValueError("Knowledge base has no wiki vault configured")

    db = _get_db()
    rows: list[sqlite3.Row] = []
    params: list = [kb_id]
    status_sql = ""
    if review_status:
        status_sql = " AND review_status = ?"
        params.append(review_status)

    if page_ids:
        if not page_ids:
            raise ValueError("page_ids is empty")
        placeholders = ",".join("?" * len(page_ids))
        rows = db.execute(
            f"""
            SELECT id, title, slug, folder_id
            FROM knowledge_wiki_pages
            WHERE kb_id = ? AND id IN ({placeholders}){status_sql}
            """,
            (kb_id, *page_ids, *([review_status] if review_status else [])),
        ).fetchall()
    elif folder_id is not None:
        rows = db.execute(
            f"""
            SELECT id, title, slug, folder_id
            FROM knowledge_wiki_pages
            WHERE kb_id = ? AND folder_id = ?{status_sql}
            """,
            (kb_id, folder_id, *([review_status] if review_status else [])),
        ).fetchall()
    else:
        rows = db.execute(
            f"""
            SELECT id, title, slug, folder_id
            FROM knowledge_wiki_pages
            WHERE kb_id = ?{status_sql}
            """,
            params,
        ).fetchall()

    page_paths: list[str] = []
    for row in rows:
        path = _wiki_page_vault_path(kb, row["slug"], row["folder_id"])
        if path:
            page_paths.append(str(path))

    if not page_paths:
        raise ValueError("No wiki pages found for curation")

    job_id = _start_wiki_curation(kb_id, page_paths, folder_id=folder_id)
    return {"job_id": job_id, "page_count": len(page_paths)}


def _start_bulk_wiki_job(
    kb_id: str,
    doc_ids: list[str],
    folder_id: str | None = None,
) -> str:
    """Launch bulk wiki generation in a background daemon thread."""
    job_id = _new_id()
    db = _get_db()
    db.execute(
        """
        INSERT INTO knowledge_curation_jobs
        (id, kb_id, folder_id, job_type, status, input_pages, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (job_id, kb_id, folder_id, "bulk_wiki", "pending", _json_dumps(doc_ids), _now_iso()),
    )
    db.commit()
    thread = threading.Thread(
        target=_run_bulk_wiki_job,
        args=(kb_id, doc_ids, job_id),
        kwargs={"folder_id": folder_id},
        daemon=True,
    )
    thread.start()
    return job_id


def _run_bulk_wiki_job(
    kb_id: str,
    doc_ids: list[str],
    job_id: str,
    folder_id: str | None = None,
) -> None:
    """Generate wiki pages for a list of documents in the background."""
    db = _get_db()
    db.execute(
        "UPDATE knowledge_curation_jobs SET status = ?, started_at = ? WHERE id = ?",
        ("running", _now_iso(), job_id),
    )
    db.commit()

    results: list[dict] = []
    errors: list[str] = []
    completed = 0
    failed = 0
    skipped = 0
    try:
        for i, doc_id in enumerate(doc_ids):
            try:
                result = generate_document_wiki(doc_id, curate=False)
                results.append({
                    "doc_id": doc_id,
                    "wiki_id": result.get("wiki_id"),
                    "title": result.get("title"),
                    "skipped": result.get("skipped", False),
                    "reason": result.get("reason", ""),
                })
                if result.get("skipped"):
                    skipped += 1
                else:
                    completed += 1
            except Exception as e:
                logger.exception("bulk wiki generation failed for doc %s", doc_id)
                errors.append(f"{doc_id}: {e}")
                failed += 1
            # Update progress after each doc so the UI can poll meaningfully.
            db.execute(
                "UPDATE knowledge_curation_jobs SET output_pages = ? WHERE id = ?",
                (_json_dumps({"completed": completed, "skipped": skipped, "failed": failed, "total": len(doc_ids), "errors": errors[:10]}), job_id),
            )
            db.commit()

        db.execute(
            "UPDATE knowledge_curation_jobs SET status = ?, completed_at = ?, output_pages = ? WHERE id = ?",
            (
                "completed",
                _now_iso(),
                _json_dumps({"completed": completed, "skipped": skipped, "failed": failed, "total": len(doc_ids), "errors": errors[:20], "results": results}),
                job_id,
            ),
        )
        db.commit()
    except Exception as e:
        logger.exception("bulk wiki generation failed for KB %s", kb_id)
        try:
            db.execute(
                "UPDATE knowledge_curation_jobs SET status = ?, completed_at = ?, error_message = ? WHERE id = ?",
                ("failed", _now_iso(), str(e)[:2000], job_id),
            )
            db.commit()
        except Exception:
            pass


def start_bulk_wiki_generation(
    kb_id: str,
    folder_id: str | None = None,
    doc_ids: list[str] | None = None,
) -> dict:
    """Launch bulk wiki generation for selected docs or all docs in a folder/KB.

    Args:
        folder_id: generate wiki for all completed documents in this folder.
        doc_ids: explicit list of document ids. Overrides folder_id.

    Returns:
        {"job_id": str, "doc_count": int}
    """
    get_knowledge_base(kb_id)  # validate
    db = _get_db()
    ids: list[str] = []
    if doc_ids:
        ids = [d for d in doc_ids if d]
    elif folder_id is not None:
        rows = db.execute(
            "SELECT id FROM knowledge_documents WHERE kb_id = ? AND folder_id = ? AND parse_status = ?",
            (kb_id, folder_id, "completed"),
        ).fetchall()
        ids = [r["id"] for r in rows]
    else:
        rows = db.execute(
            "SELECT id FROM knowledge_documents WHERE kb_id = ? AND parse_status = ?",
            (kb_id, "completed"),
        ).fetchall()
        ids = [r["id"] for r in rows]

    if not ids:
        raise ValueError("No completed documents found for wiki generation")

    job_id = _start_bulk_wiki_job(kb_id, ids, folder_id=folder_id)
    return {"job_id": job_id, "doc_count": len(ids)}


# ═══════════════════════════════════════════════════════════════════════════════
# Vectorization worker
# ═══════════════════════════════════════════════════════════════════════════════


def _get_executor() -> ThreadPoolExecutor:
    global _vect_executor
    if _vect_executor is None:
        _vect_executor = ThreadPoolExecutor(max_workers=2, thread_name_prefix="kb-vect-")
    return _vect_executor


# Separate single-thread pool for LLM post-processing (summary / graph / wiki).
# These calls are slow and must not occupy the 2 vectorization workers.
_post_executor: ThreadPoolExecutor | None = None


def _get_postprocess_executor() -> ThreadPoolExecutor:
    global _post_executor
    if _post_executor is None:
        _post_executor = ThreadPoolExecutor(max_workers=1, thread_name_prefix="kb-post-")
    return _post_executor


_vect_executor: ThreadPoolExecutor | None = None


def _shutdown_kb_executor() -> None:
    """Allow the process to exit cleanly even if vectorization is in flight."""
    global _vect_executor, _post_executor
    if _vect_executor is not None:
        _vect_executor.shutdown(wait=False)
        _vect_executor = None
    if _post_executor is not None:
        _post_executor.shutdown(wait=False)
        _post_executor = None


atexit.register(_shutdown_kb_executor)


def _update_job(job_id: str, **kwargs):
    db = _get_db()
    sets = ", ".join(f"{k} = ?" for k in kwargs)
    db.execute(
        f"UPDATE vectorization_jobs SET {sets} WHERE id = ?",
        list(kwargs.values()) + [job_id],
    )
    db.commit()


def _update_doc_status(doc_id: str, **kwargs):
    db = _get_db()
    sets = ", ".join(f"{k} = ?" for k in kwargs)
    db.execute(
        f"UPDATE knowledge_documents SET {sets}, updated_at = ? WHERE id = ?",
        list(kwargs.values()) + [_now_iso(), doc_id],
    )
    db.commit()


def _chunk_general_document(file_path: Path, max_size: int = 512, overlap: int = 64) -> list[dict]:
    """Simple generic chunker for non-PDF/DOCX energy-audit documents.

    Splits by paragraphs first, then by sentences if a paragraph is too long.
    """
    content = ""
    filename = file_path.name
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        # No binary fallback: reading a PDF as UTF-8 produces garbage chunks.
        # Let the exception propagate so the job is marked failed with the
        # real parser error instead of silently embedding noise.
        import pymupdf
        doc = pymupdf.open(str(file_path))
        content = "\n\n".join(page.get_text() for page in doc)
        doc.close()
    elif suffix == ".docx":
        from docx import Document
        doc = Document(str(file_path))
        content = "\n\n".join(para.text for para in doc.paragraphs)
    elif suffix in _TEXT_EXTS:
        content = file_path.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type for generic chunking: {suffix}")

    content = content.strip()
    if not content:
        return []

    # Split into paragraphs, then into chunks of approximately max_size chars
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]
    chunks = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 2 <= max_size:
            current = current + "\n\n" + para if current else para
        else:
            if current:
                chunks.append(current)
            # If single paragraph exceeds max_size, split by sentences
            if len(para) > max_size:
                sentences = re.split(r"(?<=[。！？.!?])\s+", para)
                current = ""
                for sent in sentences:
                    if len(current) + len(sent) + 1 <= max_size:
                        current = current + sent if current else sent
                    else:
                        if current:
                            chunks.append(current)
                        current = sent
            else:
                current = para
    if current:
        chunks.append(current)

    # Add overlap between consecutive chunks
    final_chunks = []
    for i, text in enumerate(chunks):
        prev_overlap = chunks[i - 1][-overlap:] if i > 0 and overlap > 0 else ""
        final_chunks.append({
            "id": _new_id(),
            "text": (prev_overlap + "\n" + text).strip() if prev_overlap else text,
            "chunk_type": "text",
            "chunk_index": i,
            "filename": filename,
            "prev_chunk_id": final_chunks[i - 1].get("id") if i > 0 else None,
        })
    # Fix next_chunk_id links
    for i, ch in enumerate(final_chunks):
        if i + 1 < len(final_chunks):
            ch["next_chunk_id"] = final_chunks[i + 1]["id"]
    return final_chunks


def _run_vectorization_job(job_id: str, kb_id: str, doc_id: str):
    """Background task: parse, chunk, embed and store a single document."""
    _update_job(job_id, status="processing", started_at=_now_iso())
    _update_doc_status(doc_id, parse_status="processing")

    try:
        doc = get_knowledge_document(doc_id)
        target = _resolve_path(kb_id, doc["file_path"], must_exist=True)
        kb = get_knowledge_base(kb_id)
        collection = kb["qdrant_collection"]
        kb_type = kb["kb_type"]
        chunking_config = kb.get("chunking_config") or {}

        if kb_type == "energy_audit":
            from rag.energy_audit_importer import extract_pdf_structure, build_chunks, embed_and_store
            structure = extract_pdf_structure(str(target))
            chunks = build_chunks(structure)
        else:
            # Generic document chunking
            max_size = int(chunking_config.get("max_chunk_size", 512))
            overlap = int(chunking_config.get("overlap", 64))
            chunks = _chunk_general_document(target, max_size=max_size, overlap=overlap)
            # Generic embedding store path
            from rag.energy_audit_importer import embed_and_store

        total = len(chunks)
        _update_job(job_id, chunks_total=total, chunks_done=0, progress=0)
        _update_doc_status(doc_id, chunk_count=total)

        # Idempotent re-vectorize: purge this doc's previous vectors and SQLite
        # chunks first, then re-embed with deterministic point ids (uuid5 per
        # doc+index) so repeated runs overwrite in place instead of accumulating
        # duplicates.
        _purge_doc_vectors(kb_id, doc)
        db = _get_db()
        db.execute("DELETE FROM knowledge_chunks WHERE doc_id = ?", (doc_id,))
        db.commit()
        for idx, ch in enumerate(chunks):
            ch["point_id"] = _chunk_point_id(doc_id, idx)
            # Ensure scope filters can be pushed down to Qdrant.
            ch.setdefault("file_path", doc["file_path"])
            ch.setdefault("file_type", Path(doc["file_name"]).suffix.lower().lstrip("."))

        def _progress_cb(done, total):
            _update_job(job_id, chunks_done=done, progress=int(round(done * 100 / total)) if total else 0)

        embed_and_store(chunks, collection, progress_callback=_progress_cb)

        # Persist chunks to SQLite for inspection. Commit in batches so the
        # write lock is not held for the entire insert sequence; this keeps
        # concurrent start-vectorization requests responsive.
        db = _get_db()
        for idx, ch in enumerate(chunks):
            chunk_id = _new_id()
            db.execute(
                """
                INSERT INTO knowledge_chunks
                (id, doc_id, kb_id, chunk_index, chunk_type, content, content_hash,
                 char_count, is_enabled, parent_chunk_id, prev_chunk_id, next_chunk_id,
                 metadata, qdrant_point_id, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    chunk_id,
                    doc_id,
                    kb_id,
                    idx,
                    ch.get("chunk_type", "text"),
                    ch.get("text", ""),
                    hashlib.sha256(ch.get("text", "").encode()).hexdigest(),
                    len(ch.get("text", "")),
                    1,
                    ch.get("parent_chunk_id"),
                    ch.get("prev_chunk_id"),
                    ch.get("next_chunk_id"),
                    _json_dumps({k: v for k, v in ch.items() if k not in ("text", "embedding", "point_id")}),
                    ch.get("point_id") or ch.get("id"),
                    _now_iso(),
                ),
            )
            if (idx + 1) % 100 == 0:
                db.commit()

        # Sync legacy document_state for old clients (mark vectorization done)
        _set_vect_status(doc["file_path"], status="completed", progress=100, chunks_done=total, chunks_total=total)

        _update_job(job_id, status="completed", progress=100, chunks_done=total, completed_at=_now_iso())
        _update_doc_status(doc_id, parse_status="completed", vector_count=total, processed_at=_now_iso())

        # Auto-generate summary / graph / wiki on the separate post-process
        # pool: these LLM calls are slow and must not block the vectorization
        # workers queued behind this document. The job is already marked
        # completed above, so clients see vectorization done while these run.
        indexing_strategy = kb.get("indexing_strategy") or {}
        if indexing_strategy.get("summary") or indexing_strategy.get("graph") or indexing_strategy.get("wiki"):
            _get_postprocess_executor().submit(_run_auto_postprocess, doc_id, indexing_strategy)

    except Exception as e:
        logger.exception("Vectorization failed for doc %s", doc_id)
        _update_job(job_id, status="failed", error=str(e), completed_at=_now_iso())
        _update_doc_status(doc_id, parse_status="failed", error_message=str(e))
        try:
            doc = get_knowledge_document(doc_id)
            _set_vect_status(doc["file_path"], status="failed", error=str(e))
        except Exception:
            pass


def _run_auto_postprocess(doc_id: str, indexing_strategy: dict) -> None:
    """Background: LLM post-processing (summary / graph / wiki) after a
    document's vectorization has completed. Runs on the dedicated post-process
    pool so slow LLM calls never occupy vectorization workers."""
    if indexing_strategy.get("summary"):
        _run_summary_job(doc_id)

    if indexing_strategy.get("graph"):
        _run_graph_job(doc_id)

    if indexing_strategy.get("wiki"):
        _run_doc_wiki_job(doc_id, curate=bool(indexing_strategy.get("wiki_curate")))


def start_vectorization_v2(doc_id: str) -> dict:
    doc = get_knowledge_document(doc_id)
    kb_id = doc["kb_id"]
    db = _get_db()

    # Use INSERT OR IGNORE guarded by a partial unique index on
    # (doc_id) WHERE status = 'processing' to make the "check then insert"
    # atomic and prevent concurrent requests from starting duplicate workers.
    # The row is inserted as 'processing' directly so the partial unique index
    # actually guards it; cursor.rowcount tells us whether WE won the race.
    job_id = _new_id()
    now = _now_iso()
    cursor = db.execute(
        """
        INSERT OR IGNORE INTO vectorization_jobs (id, kb_id, doc_id, status, progress, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (job_id, kb_id, doc_id, "processing", 0, now),
    )
    if cursor.rowcount == 0:
        # Another request already owns an active job for this document.
        db.rollback()
        row = db.execute(
            "SELECT id FROM vectorization_jobs WHERE doc_id = ? AND status = 'processing' LIMIT 1",
            (doc_id,),
        ).fetchone()
        return {"job_id": row["id"] if row else None, "doc_id": doc_id, "status": "processing"}

    # We own the new job. Mark the document as queued in the same transaction.
    db.execute(
        """
        UPDATE knowledge_documents
        SET parse_status = ?, vector_count = ?, chunk_count = ?, error_message = ?, updated_at = ?
        WHERE id = ?
        """,
        ("processing", 0, 0, "", now, doc_id),
    )
    db.commit()

    _get_executor().submit(_run_vectorization_job, job_id, kb_id, doc_id)
    return {"job_id": job_id, "doc_id": doc_id, "status": "processing"}


def get_vectorization_job(job_id: str) -> dict:
    db = _get_db()
    row = db.execute("SELECT * FROM vectorization_jobs WHERE id = ?", (job_id,)).fetchone()
    if not row:
        raise ValueError("Job not found")
    return {
        "id": row["id"],
        "kb_id": row["kb_id"],
        "doc_id": row["doc_id"],
        "status": row["status"],
        "progress": row["progress"],
        "chunks_done": row["chunks_done"],
        "chunks_total": row["chunks_total"],
        "error": row["error"],
        "started_at": row["started_at"],
        "completed_at": row["completed_at"],
        "created_at": row["created_at"],
    }


def list_vectorization_jobs(kb_id: str, status: str | None = None) -> list[dict]:
    db = _get_db()
    sql = "SELECT * FROM vectorization_jobs WHERE kb_id = ?"
    args = [kb_id]
    if status:
        sql += " AND status = ?"
        args.append(status)
    sql += " ORDER BY created_at DESC"
    rows = db.execute(sql, args).fetchall()
    return [get_vectorization_job(row["id"]) for row in rows]


def rebuild_knowledge_base(kb_id: str, targets: list[str] | None = None) -> dict:
    """Rebuild vector index (and optionally graph/wiki) for a KB."""
    kb = get_knowledge_base(kb_id)
    db = _get_db()
    rows = db.execute(
        "SELECT id FROM knowledge_documents WHERE kb_id = ?", (kb_id,)
    ).fetchall()

    job_ids = []
    for row in rows:
        doc_id = row["id"]
        _update_doc_status(doc_id, parse_status="pending", vector_count=0, chunk_count=0)
        # Delete existing chunks
        db.execute("DELETE FROM knowledge_chunks WHERE doc_id = ?", (doc_id,))
        db.commit()
        result = start_vectorization_v2(doc_id)
        job_ids.append(result["job_id"])

    return {"kb_id": kb_id, "rebuild_jobs": job_ids, "queued_documents": len(rows)}


# ═══════════════════════════════════════════════════════════════════════════════
# Chunk management
# ═══════════════════════════════════════════════════════════════════════════════


def list_document_chunks(doc_id: str) -> dict:
    db = _get_db()
    rows = db.execute(
        "SELECT * FROM knowledge_chunks WHERE doc_id = ? ORDER BY chunk_index",
        (doc_id,),
    ).fetchall()
    return {
        "doc_id": doc_id,
        "chunks": [
            {
                "id": row["id"],
                "chunk_index": row["chunk_index"],
                "chunk_type": row["chunk_type"],
                "content": row["content"],
                "char_count": row["char_count"],
                "is_enabled": bool(row["is_enabled"]),
                "parent_chunk_id": row["parent_chunk_id"],
                "prev_chunk_id": row["prev_chunk_id"],
                "next_chunk_id": row["next_chunk_id"],
                "metadata": _json_loads(row["metadata"]),
                "qdrant_point_id": row["qdrant_point_id"],
            }
            for row in rows
        ],
    }


def update_knowledge_chunk(chunk_id: str, content: str | None = None, is_enabled: bool | None = None) -> dict:
    db = _get_db()
    row = db.execute("SELECT * FROM knowledge_chunks WHERE id = ?", (chunk_id,)).fetchone()
    if not row:
        raise ValueError("Chunk not found")

    updates = {}
    if content is not None:
        updates["content"] = content
        updates["content_hash"] = hashlib.sha256(content.encode()).hexdigest()
        updates["char_count"] = len(content)
    if is_enabled is not None:
        updates["is_enabled"] = 1 if is_enabled else 0

    if not updates:
        return list_document_chunks(row["doc_id"])

    sets = ", ".join(f"{k} = ?" for k in updates)
    db.execute(f"UPDATE knowledge_chunks SET {sets} WHERE id = ?", list(updates.values()) + [chunk_id])
    db.commit()

    # Keep Qdrant in sync: disabling removes the vector (so disabled chunks
    # stop surfacing in search), enabling/editing re-embeds it.
    try:
        new_content = content if content is not None else row["content"]
        enabled = (row["is_enabled"] == 1) if is_enabled is None else bool(is_enabled)
        _sync_chunk_vector(row, new_content, enabled, content_changed=content is not None)
    except Exception:
        logger.exception("Failed to sync chunk vector %s", chunk_id)

    return list_document_chunks(row["doc_id"])


def _sync_chunk_vector(chunk_row, content: str, enabled: bool, content_changed: bool) -> None:
    """Sync one chunk's Qdrant point with its enable/edit state."""
    from qdrant_client.models import FieldCondition, MatchValue, PointStruct

    doc = get_knowledge_document(chunk_row["doc_id"])
    kb = get_knowledge_base(chunk_row["kb_id"])
    collection = kb["qdrant_collection"]
    stored_id = chunk_row["qdrant_point_id"]
    point_id = _chunk_point_id(chunk_row["doc_id"], chunk_row["chunk_index"] or 0)

    # Remove the stale point: by explicit id when we have one, plus a payload
    # filter fallback for legacy points whose payload carries the chunk's
    # internal id (general-KB chunker writes it as payload key "id").
    if not enabled or content_changed:
        if stored_id:
            _qdrant_delete_points(collection, [stored_id])
            _qdrant_delete_by_filter(collection, [FieldCondition(key="id", match=MatchValue(value=stored_id))])

    if enabled:
        vector = _get_embedding([content])[0]
        _qdrant_upsert(
            collection,
            [
                PointStruct(
                    id=point_id,
                    vector=vector,
                    payload={
                        "filename": doc["file_name"],
                        "file_path": doc["file_path"],
                        "file_type": Path(doc["file_name"]).suffix.lower().lstrip("."),
                        "text": content,
                        "chunk_type": chunk_row["chunk_type"] or "text",
                        "doc_id": chunk_row["doc_id"],
                        "id": point_id,
                    },
                )
            ],
        )
        if stored_id != point_id:
            db = _get_db()
            db.execute("UPDATE knowledge_chunks SET qdrant_point_id = ? WHERE id = ?", (point_id, chunk_row["id"]))
            db.commit()


def delete_knowledge_chunk(chunk_id: str) -> dict:
    db = _get_db()
    row = db.execute(
        "SELECT doc_id, kb_id, qdrant_point_id, chunk_index FROM knowledge_chunks WHERE id = ?",
        (chunk_id,),
    ).fetchone()
    if not row:
        raise ValueError("Chunk not found")
    doc_id = row["doc_id"]
    kb_id = row["kb_id"]
    point_id = row["qdrant_point_id"] or _chunk_point_id(doc_id, row["chunk_index"])

    db.execute("DELETE FROM knowledge_chunks WHERE id = ?", (chunk_id,))
    db.commit()

    kb = get_knowledge_base(kb_id)
    _qdrant_delete_points(kb["qdrant_collection"], [point_id])

    return list_document_chunks(doc_id)


# ═══════════════════════════════════════════════════════════════════════════════
# Search
# ═══════════════════════════════════════════════════════════════════════════════


def _get_folder_descendant_paths(kb_id: str, folder_id: str | None) -> set[str]:
    """Return all file paths under a folder (including subfolders)."""
    if not folder_id:
        return set()
    db = _get_db()
    folder = db.execute(
        "SELECT path FROM knowledge_folders WHERE id = ? AND kb_id = ?", (folder_id, kb_id)
    ).fetchone()
    if not folder:
        return set()
    base_path = folder["path"]
    rows = db.execute(
        "SELECT file_path FROM knowledge_documents WHERE kb_id = ? AND (folder_id = ? OR file_path LIKE ?)",
        (kb_id, folder_id, base_path + "/%"),
    ).fetchall()
    return {row["file_path"] for row in rows}


def _search_qdrant_for_kb(
    kb_id: str,
    query: str,
    top_k: int = 5,
    folder_id: str | None = None,
    file_type: str | None = None,
    doc_id: str | None = None,
    score_threshold: float | None = None,
) -> dict:
    """Search a KB's Qdrant collection with optional scope filters."""
    from qdrant_client.models import Filter, FieldCondition, MatchAny, MatchValue

    kb = get_knowledge_base(kb_id)
    collection = kb["qdrant_collection"]

    from rag.rag_search import _embed_query

    vector = _embed_query(query)

    conditions: list = []
    if folder_id:
        allowed_paths = _get_folder_descendant_paths(kb_id, folder_id)
        if allowed_paths:
            conditions.append(
                FieldCondition(
                    key="file_path",
                    match=MatchAny(any=list(allowed_paths)),
                )
            )
        else:
            # empty folder => no results
            return {"results": [], "source": "qdrant_vector", "count": 0}
    elif doc_id:
        doc = get_knowledge_document(doc_id)
        if doc and doc.get("file_path"):
            conditions.append(
                FieldCondition(
                    key="file_path",
                    match=MatchValue(value=doc["file_path"]),
                )
            )
        else:
            return {"results": [], "source": "qdrant_vector", "count": 0}

    if file_type:
        conditions.append(
            FieldCondition(
                key="file_type",
                match=MatchValue(value=file_type.lower().lstrip(".")),
            )
        )

    qdrant_filter = Filter(must=conditions) if conditions else None

    client = _qdrant_client()
    r = client.query_points(
        collection_name=collection,
        query=vector,
        query_filter=qdrant_filter,
        limit=top_k,
        score_threshold=score_threshold,
        with_payload=True,
    )

    results = []
    for hit in r.points:
        filename = hit.payload.get("filename", "")
        results.append({
            "score": hit.score,
            "filename": filename,
            "chapter": hit.payload.get("chapter", ""),
            "text": hit.payload.get("text", "")[:2000],
            "metadata": {
                "audit_type": hit.payload.get("audit_type", ""),
                "institution_category": hit.payload.get("institution_category", ""),
                "specific_type": hit.payload.get("specific_type", ""),
                "filename": filename,
            },
        })

    return {"results": results, "source": "qdrant_vector", "count": len(results)}


def _search_graph_for_kb(
    kb_id: str,
    query: str,
    top_k: int = 5,
    score_threshold: float | None = None,
) -> dict:
    """Vector GraphRAG search: semantic entity match + multi-hop traversal + LLM answer."""
    kb = get_knowledge_base(kb_id)
    collection = _kb_entity_collection(kb)

    from rag.rag_search import _embed_query

    vector = _embed_query(query)

    client = _qdrant_client()

    matched = client.query_points(
        collection_name=collection,
        query=vector,
        limit=top_k * 3,
        score_threshold=score_threshold,
        with_payload=True,
    )

    db = _get_db()
    seed_ids: set[str] = set()
    for hit in matched.points:
        eid = hit.payload.get("entity_id")
        if eid:
            seed_ids.add(eid)

    if not seed_ids:
        return {"results": [], "source": "graph", "count": 0}

    # Multi-hop traversal (max 3 hops) collecting ranked neighbors
    max_hops = 3
    nodes_by_id: dict[str, dict] = {}
    edges: list[dict] = []
    visited_pairs: set[tuple[str, str]] = set()
    frontier = set(seed_ids)
    for hop in range(max_hops):
        if not frontier:
            break
        rows = db.execute(
            f"""
            SELECT r.id, r.relation_type, r.description,
                   s.id AS source_id, s.name AS source_name, s.entity_type AS source_type, s.description AS source_desc,
                   t.id AS target_id, t.name AS target_name, t.entity_type AS target_type, t.description AS target_desc
            FROM knowledge_relationships r
            JOIN knowledge_entities s ON s.id = r.source_entity_id
            JOIN knowledge_entities t ON t.id = r.target_entity_id
            WHERE r.kb_id = ? AND (r.source_entity_id IN ({','.join('?' * len(frontier))})
                                  OR r.target_entity_id IN ({','.join('?' * len(frontier))}))
            LIMIT ?
            """,
            (kb_id, *tuple(frontier), *tuple(frontier), top_k * 8),
        ).fetchall()
        frontier = set()
        for row in rows:
            pair = tuple(sorted((row["source_id"], row["target_id"])))
            if pair in visited_pairs:
                continue
            visited_pairs.add(pair)
            edges.append({
                "source": row["source_name"],
                "source_type": row["source_type"],
                "target": row["target_name"],
                "target_type": row["target_type"],
                "relation": row["relation_type"],
                "description": row["description"],
                "hop": hop + 1,
            })
            for col, nid, nname, ntype, ndesc in [
                ("source", row["source_id"], row["source_name"], row["source_type"], row["source_desc"]),
                ("target", row["target_id"], row["target_name"], row["target_type"], row["target_desc"]),
            ]:
                if nid not in nodes_by_id:
                    nodes_by_id[nid] = {"id": nid, "name": nname, "type": ntype, "description": ndesc, "hop": hop + 1}
                    frontier.add(nid)

    # Include seed nodes too
    for eid in seed_ids:
        if eid not in nodes_by_id:
            row = db.execute(
                "SELECT id, name, entity_type, description FROM knowledge_entities WHERE id = ?",
                (eid,),
            ).fetchone()
            if row:
                nodes_by_id[eid] = {
                    "id": row["id"],
                    "name": row["name"],
                    "type": row["entity_type"],
                    "description": row["description"],
                    "hop": 0,
                }

    nodes = list(nodes_by_id.values())
    # Fetch relevant chunks connected to top nodes for grounded answer generation
    top_node_ids = [n["id"] for n in sorted(nodes, key=lambda x: (x["hop"], x["name"]))[:top_k * 2]]
    chunks = _fetch_chunks_for_entities(kb_id, top_node_ids, limit=top_k * 3)

    graph_context = _format_graph_context(nodes, edges, chunks)
    answer = _generate_graphrag_answer(query, graph_context)

    results = [{
        "score": 1.0,
        "text": answer,
        "query": query,
        "entities": nodes[:top_k * 2],
        "relationships": edges[:top_k * 3],
        "chunks": chunks[:top_k],
        "metadata": {"entity_count": len(nodes), "edge_count": len(edges), "seed_count": len(seed_ids)},
    }]
    return {"results": results, "source": "graph", "count": len(results)}


def _fetch_chunks_for_entities(kb_id: str, entity_ids: list[str], limit: int = 20) -> list[dict]:
    """Retrieve chunk content linked to entities through shared documents."""
    if not entity_ids:
        return []
    db = _get_db()
    # Map entities to their originating doc_ids
    rows = db.execute(
        f"""
        SELECT DISTINCT doc_id FROM knowledge_entities
        WHERE kb_id = ? AND id IN ({','.join('?' * len(entity_ids))})
        """,
        (kb_id, *entity_ids),
    ).fetchall()
    doc_ids = [r["doc_id"] for r in rows]
    if not doc_ids:
        return []
    c_rows = db.execute(
        f"""
        SELECT id, doc_id, content, chunk_index FROM knowledge_chunks
        WHERE doc_id IN ({','.join('?' * len(doc_ids))}) AND is_enabled = 1
        ORDER BY doc_id, chunk_index
        LIMIT ?
        """,
        (*doc_ids, limit),
    ).fetchall()
    return [
        {"chunk_id": r["id"], "doc_id": r["doc_id"], "text": r["content"][:1200], "index": r["chunk_index"]}
        for r in c_rows
    ]


def _format_graph_context(nodes: list[dict], edges: list[dict], chunks: list[dict]) -> str:
    lines = []
    lines.append("## 相关实体")
    for n in nodes:
        lines.append(f"- {n['name']}（{n['type']}）：{n.get('description', '')}")
    lines.append("\n## 关系")
    for e in edges:
        desc = e.get("description") or ""
        lines.append(f"- {e['source']} --[{e['relation']}]--> {e['target']}" + (f"：{desc}" if desc else ""))
    if chunks:
        lines.append("\n## 原始文本片段")
        for c in chunks[:10]:
            lines.append(f"- {c['text']}")
    return "\n".join(lines)


def _generate_graphrag_answer(query: str, graph_context: str) -> str:
    """Use LLM to synthesize a natural-language answer from graph context."""
    prompt = (
        "你是一位知识库问答助手。请根据下面的知识图谱信息（实体、关系、原始文本片段）回答用户问题。"
        "如果信息不足，请明确说明。回答要简洁、准确，并用中文。\n\n"
        f"用户问题：{query}\n\n"
        f"知识图谱上下文：\n{graph_context}\n\n"
        "答案："
    )
    answer = _call_llm(
        [
            {"role": "system", "content": "你是一位基于知识图谱回答问题的助手。"},
            {"role": "user", "content": prompt},
        ],
        temperature=0.3,
        max_tokens=1200,
        task="graphrag_answer",
    )
    return answer or "未能根据知识图谱生成答案。"


def _search_wiki_for_kb(
    kb_id: str,
    query: str,
    top_k: int = 5,
    score_threshold: float | None = None,
) -> dict:
    """Vector semantic search over generated wiki pages."""
    kb = get_knowledge_base(kb_id)
    collection = _kb_wiki_collection(kb)

    from rag.rag_search import _embed_query

    vector = _embed_query(query)

    client = _qdrant_client()

    r = client.query_points(
        collection_name=collection,
        query=vector,
        limit=top_k,
        score_threshold=score_threshold,
        with_payload=True,
    )

    results = []
    for hit in r.points:
        payload = hit.payload or {}
        results.append({
            "score": hit.score,
            "title": payload.get("title", ""),
            "wiki_id": payload.get("wiki_id", ""),
            "text": payload.get("text", "")[:2000],
        })
    return {"results": results, "source": "wiki", "count": len(results)}


def list_kb_entities(kb_id: str, top_k: int = 100) -> dict:
    """List entities extracted for a KB."""
    db = _get_db()
    rows = db.execute(
        "SELECT id, name, entity_type, description FROM knowledge_entities WHERE kb_id = ? ORDER BY name LIMIT ?",
        (kb_id, top_k),
    ).fetchall()
    return {
        "kb_id": kb_id,
        "entities": [{"id": r["id"], "name": r["name"], "type": r["entity_type"], "description": r["description"]} for r in rows],
    }


def list_kb_relationships(kb_id: str, top_k: int = 200) -> dict:
    """List relationships extracted for a KB."""
    db = _get_db()
    rows = db.execute(
        """
        SELECT r.id, r.relation_type, r.description,
               s.name as source_name, t.name as target_name
        FROM knowledge_relationships r
        JOIN knowledge_entities s ON s.id = r.source_entity_id
        JOIN knowledge_entities t ON t.id = r.target_entity_id
        WHERE r.kb_id = ?
        ORDER BY r.created_at DESC
        LIMIT ?
        """,
        (kb_id, top_k),
    ).fetchall()
    return {
        "kb_id": kb_id,
        "relationships": [
            {
                "id": r["id"],
                "source": r["source_name"],
                "target": r["target_name"],
                "relation": r["relation_type"],
                "description": r["description"],
            }
            for r in rows
        ],
    }


def list_kb_wiki_pages(kb_id: str, top_k: int = 100, review_status: str | None = None) -> dict:
    """List wiki pages for a KB, optionally filtered by review_status."""
    db = _get_db()
    sql = """
        SELECT w.id, w.title, w.slug, w.status, w.review_status, w.updated_at, w.folder_id, w.doc_id,
               w.quality_score, w.quality_report, f.path as folder_path
        FROM knowledge_wiki_pages w
        LEFT JOIN knowledge_folders f ON f.id = w.folder_id
        WHERE w.kb_id = ?
    """
    params: list = [kb_id]
    if review_status:
        sql += " AND w.review_status = ?"
        params.append(review_status)
    sql += " ORDER BY w.updated_at DESC LIMIT ?"
    params.append(top_k)
    rows = db.execute(sql, params).fetchall()
    return {
        "kb_id": kb_id,
        "pages": [
            {
                "id": r["id"],
                "title": r["title"],
                "slug": r["slug"],
                "status": r["status"],
                "review_status": r["review_status"] or "pending",
                "updated_at": r["updated_at"],
                "folder_id": r["folder_id"],
                "folder_path": r["folder_path"] or "",
                "doc_id": r["doc_id"],
                "source": "folder" if r["folder_id"] else ("doc" if r["doc_id"] else "unknown"),
                "quality_score": r["quality_score"] if r["quality_score"] is not None else None,
                "quality_report": _json_loads(r["quality_report"]),
            }
            for r in rows
        ],
    }


def get_wiki_page(wiki_id: str) -> dict:
    """Get a single wiki page by id."""
    db = _get_db()
    row = db.execute(
        """
        SELECT w.*, f.path as folder_path
        FROM knowledge_wiki_pages w
        LEFT JOIN knowledge_folders f ON f.id = w.folder_id
        WHERE w.id = ?
        """,
        (wiki_id,),
    ).fetchone()
    if not row:
        raise ValueError("Wiki page not found")
    return {
        "id": row["id"],
        "kb_id": row["kb_id"],
        "doc_id": row["doc_id"],
        "folder_id": row["folder_id"],
        "folder_path": row["folder_path"] or "",
        "title": row["title"],
        "slug": row["slug"],
        "content": row["content"],
        "status": row["status"],
        "review_status": row["review_status"] or "pending",
        "updated_at": row["updated_at"],
        "source": "folder" if row["folder_id"] else ("doc" if row["doc_id"] else "unknown"),
        "quality_score": row["quality_score"] if row["quality_score"] is not None else None,
        "quality_report": _json_loads(row["quality_report"]),
    }


def update_wiki_review_status(wiki_id: str, review_status: str) -> dict:
    """Update the manual review status of a wiki page.

    Allowed values: pending, approved, rejected.
    """
    if review_status not in ("pending", "approved", "rejected"):
        raise ValueError("review_status must be one of: pending, approved, rejected")
    db = _get_db()
    row = db.execute("SELECT id FROM knowledge_wiki_pages WHERE id = ?", (wiki_id,)).fetchone()
    if not row:
        raise ValueError("Wiki page not found")
    db.execute(
        "UPDATE knowledge_wiki_pages SET review_status = ?, updated_at = ? WHERE id = ?",
        (review_status, _now_iso(), wiki_id),
    )
    db.commit()
    return {"id": wiki_id, "review_status": review_status}


def evaluate_wiki_quality(wiki_id: str) -> dict:
    """Re-evaluate and store the quality score for a wiki page."""
    db = _get_db()
    row = db.execute(
        "SELECT kb_id, doc_id, folder_id, title, content FROM knowledge_wiki_pages WHERE id = ?",
        (wiki_id,),
    ).fetchone()
    if not row:
        raise ValueError("Wiki page not found")

    source_text = ""
    if row["doc_id"]:
        try:
            preview = get_knowledge_file_preview_v2(row["doc_id"], max_chars=12_000)
            source_text = preview.get("content", "")
        except Exception:
            pass

    kb = get_knowledge_base(row["kb_id"])
    cfg = _get_kb_curation_config(kb)
    quality = _evaluate_wiki_quality(
        wiki_id, row["title"], row["content"], source_text,
        temperature=cfg["curator_temperature"],
    )
    db.execute(
        "UPDATE knowledge_wiki_pages SET quality_score = ?, quality_report = ?, updated_at = ? WHERE id = ?",
        (quality["score"], _json_dumps(quality), _now_iso(), wiki_id),
    )
    db.commit()
    return {"id": wiki_id, "quality_score": quality["score"], "quality_report": quality}


def _run_wiki_quality_job(wiki_id: str) -> None:
    _run_pipeline_job("wiki_quality", wiki_id, evaluate_wiki_quality, wiki_id)


def start_wiki_quality_eval(wiki_id: str) -> dict:
    """Enqueue wiki quality evaluation and return immediately."""
    get_wiki_page(wiki_id)
    _get_postprocess_executor().submit(_run_wiki_quality_job, wiki_id)
    return {"id": wiki_id, "status": "processing"}


def search_knowledge_v2(
    kb_id: str,
    query: str,
    top_k: int = 5,
    mode: str = "vector",
    folder_id: str | None = None,
    file_type: str | None = None,
    doc_id: str | None = None,
    score_threshold: float | None = None,
) -> dict:
    kb = get_knowledge_base(kb_id)
    try:
        if mode == "vector":
            result = _search_qdrant_for_kb(
                kb_id, query, top_k, folder_id, file_type, doc_id,
                score_threshold=score_threshold,
            )
        elif mode == "graph":
            result = _search_graph_for_kb(kb_id, query, top_k, score_threshold=score_threshold)
        elif mode == "wiki":
            result = _search_wiki_for_kb(kb_id, query, top_k, score_threshold=score_threshold)
        elif mode == "graph_wiki":
            graph_result = _search_graph_for_kb(kb_id, query, top_k, score_threshold=score_threshold)
            wiki_result = _search_wiki_for_kb(kb_id, query, top_k, score_threshold=score_threshold)
            graph_item = (graph_result.get("results") or [{}])[0]
            wiki_results = wiki_result.get("results") or []
            # Use the graph RAG answer as the primary synthesized answer; if the
            # graph is empty, fall back to the top wiki snippets.
            combined_answer = graph_item.get("text", "")
            if not combined_answer and wiki_results:
                combined_answer = "\n\n".join(r.get("text", "") for r in wiki_results[:2])
            result = {
                "results": [
                    {
                        "type": "graph_wiki",
                        "score": 1.0,
                        "answer": combined_answer,
                        "graph": {
                            "entities": graph_item.get("entities", []),
                            "relationships": graph_item.get("relationships", []),
                        },
                        "wiki": {"results": wiki_results},
                    }
                ],
                "source": "graph+wiki",
                "count": len(wiki_results) + len(graph_item.get("entities", [])),
            }
        elif mode == "unified":
            wiki_result = _search_wiki_for_kb(kb_id, query, top_k, score_threshold=score_threshold)
            graph_result = _search_graph_for_kb(kb_id, query, top_k, score_threshold=score_threshold)
            result = {
                "results": [
                    {
                        "type": "graph",
                        "score": 1.0,
                        "text": (graph_result.get("results") or [{}])[0].get("text", ""),
                        "entities": (graph_result.get("results") or [{}])[0].get("entities", []),
                        "relationships": (graph_result.get("results") or [{}])[0].get("relationships", []),
                    },
                    *[
                        {"type": "wiki", "score": r.get("score", 0), "title": r.get("title", ""), "wiki_id": r.get("wiki_id", ""), "text": r.get("text", "")}
                        for r in (wiki_result.get("results") or [])
                    ],
                ],
                "source": "unified",
                "count": (wiki_result.get("count", 0) + 1),
            }
        else:
            return {
                "kb_id": kb_id,
                "query": query,
                "mode": mode,
                "results": [],
                "source": "error",
                "error": f"Unsupported search mode: {mode!r}. Supported modes: vector, graph, wiki, graph_wiki, unified.",
            }
        return {
            "kb_id": kb_id,
            "query": query,
            "mode": mode,
            "results": result.get("results", []),
            "source": result.get("source", "none"),
            "count": result.get("count", 0),
        }
    except ImportError as e:
        return {"kb_id": kb_id, "query": query, "mode": mode, "results": [], "source": "error", "error": str(e)}
    except Exception as e:
        logger.exception("Search failed")
        return {"kb_id": kb_id, "query": query, "mode": mode, "results": [], "source": "error", "error": str(e)}


# ═══════════════════════════════════════════════════════════════════════════════
# Backwards-compatible legacy API (operates on default KB)
# ═══════════════════════════════════════════════════════════════════════════════


def _default_kb_id() -> str:
    _ensure_default_kb()
    return _DEFAULT_KBS[0]["id"]


def list_knowledge_tree(folder: str = "", kb_id: str | None = None) -> dict:
    """Tree endpoint — the DB is the single source of truth.

    Disk is reconciled into the DB first (new/changed files appear; orphans
    are kept and flagged), then the tree is built purely from
    knowledge_folders + knowledge_documents. Orphaned documents (DB row
    without a disk file) are included with "orphaned": true so the UI can
    surface them for one-click cleanup.
    """
    kb_id = kb_id or _default_kb_id()
    root = _resolve_kb_root(kb_id)
    base = root / folder if folder else root
    base = base.resolve()
    if not base.exists():
        return {"entries": [], "root": str(root), "collection": _get_kb_collection_name(kb_id), "kb_id": kb_id}

    _sync_disk_to_db(kb_id)

    db = _get_db()
    folder_rows = db.execute(
        "SELECT id, path, name FROM knowledge_folders WHERE kb_id = ? ORDER BY depth, name",
        (kb_id,),
    ).fetchall()
    doc_rows = db.execute(
        "SELECT file_name, file_path, file_size, parse_status, updated_at FROM knowledge_documents WHERE kb_id = ?",
        (kb_id,),
    ).fetchall()

    # Build folder nodes keyed by relative path.
    nodes: dict[str, dict] = {}
    for f in folder_rows:
        if f["path"] == "":
            continue  # the root record is implicit, not a tree node
        nodes[f["path"]] = {
            "name": f["name"] or Path(f["path"]).name,
            "path": f["path"],
            "type": "folder",
            "children": [],
            "folder_id": f["id"],
        }

    tree: list[dict] = []
    for path, node in nodes.items():
        parent = path.rsplit("/", 1)[0] if "/" in path else ""
        if parent and parent in nodes:
            nodes[parent]["children"].append(node)
        else:
            tree.append(node)

    # Attach documents to their parent folder node (or the top level).
    for d in doc_rows:
        fp = d["file_path"] or ""
        parent = fp.rsplit("/", 1)[0] if "/" in fp else ""
        node = {
            "name": d["file_name"] or Path(fp).name,
            "path": fp,
            "type": "file",
            "size": d["file_size"] or 0,
            "modified": (d["updated_at"] or "")[:10],
            "vect": "done" if d["parse_status"] == "completed" else (d["parse_status"] or "pending"),
        }
        if not (root / fp).exists():
            node["orphaned"] = True
        if parent and parent in nodes:
            nodes[parent]["children"].append(node)
        else:
            tree.append(node)

    # Folders first, then files, name-sorted at every level.
    def _sort_level(items: list[dict]) -> list[dict]:
        return sorted(items, key=lambda n: (n["type"] == "file", (n["name"] or "").lower()))

    def _sort_recursive(items: list[dict]) -> list[dict]:
        for n in items:
            if n.get("children"):
                n["children"] = _sort_recursive(n["children"])
        return _sort_level(items)

    tree = _sort_recursive(tree)

    if folder:
        # Preserve the legacy subtree behaviour when a folder is requested.
        def _find(items: list[dict], target: str) -> dict | None:
            for n in items:
                if n["path"] == target:
                    return n
                if n.get("children"):
                    hit = _find(n["children"], target)
                    if hit:
                        return hit
            return None

        sub = _find(tree, folder.replace("\\", "/").strip("/"))
        tree = sub["children"] if sub else []

    return {
        "entries": tree,
        "root": str(root),
        "collection": _get_kb_collection_name(kb_id),
        "kb_id": kb_id,
    }


def create_knowledge_folder(path: str) -> dict:
    """Legacy folder creation — default KB, parent inferred from path."""
    kb_id = _default_kb_id()
    parent_path = str(Path(path).parent).replace("\\", "/")
    parent_id = None
    if parent_path and parent_path != ".":
        db = _get_db()
        row = db.execute(
            "SELECT id FROM knowledge_folders WHERE kb_id = ? AND path = ?",
            (kb_id, parent_path),
        ).fetchone()
        if row:
            parent_id = row["id"]
    result = create_knowledge_folder_v2(kb_id, Path(path).name, parent_id)
    return {"path": result["path"], "created": True}


def delete_knowledge_folder(path: str, recursive: bool = False) -> dict:
    kb_id = _default_kb_id()
    db = _get_db()
    row = db.execute(
        "SELECT id FROM knowledge_folders WHERE kb_id = ? AND path = ?", (kb_id, path)
    ).fetchone()
    if not row:
        # Fallback to raw disk delete if not in DB
        target = _resolve_path(kb_id, path, must_exist=True)
        if target.is_dir():
            if recursive:
                shutil.rmtree(target)
            else:
                target.rmdir()
        return {"path": path, "deleted": True}
    return delete_knowledge_folder_v2(kb_id, row["id"], recursive=recursive)


def upload_knowledge_file(folder: str, raw_body: bytes, content_type: str) -> dict:
    """Legacy upload endpoint — default KB."""
    kb_id = _default_kb_id()
    folder_id = None
    if folder:
        db = _get_db()
        row = db.execute(
            "SELECT id FROM knowledge_folders WHERE kb_id = ? AND path = ?",
            (kb_id, folder),
        ).fetchone()
        if row:
            folder_id = row["id"]
    result = upload_knowledge_file_v2(kb_id, folder_id, raw_body, content_type)
    if result.get("duplicate"):
        existing = result["existing"]
        return {
            "path": existing["file_path"],
            "name": existing["file_name"],
            "size": existing["file_size"],
            "mime": mimetypes.guess_type(existing["file_name"])[0] or "application/octet-stream",
            "duplicate": True,
        }
    doc = result["document"]
    return {
        "path": doc["file_path"],
        "name": doc["file_name"],
        "size": doc["file_size"],
        "mime": result["mime"],
    }


def delete_knowledge_file(path: str) -> dict:
    """Legacy file deletion — default KB."""
    kb_id = _default_kb_id()
    db = _get_db()
    row = db.execute(
        "SELECT id FROM knowledge_documents WHERE kb_id = ? AND file_path = ?",
        (kb_id, path),
    ).fetchone()
    if row:
        return delete_knowledge_document(row["id"])
    # Fallback to raw disk delete
    target = _resolve_path(kb_id, path, must_exist=True)
    target.unlink()
    _delete_vect_status(path)
    return {"path": path, "deleted": True}


def get_knowledge_file_preview(path: str) -> dict:
    """Legacy preview endpoint — default KB."""
    kb_id = _default_kb_id()
    db = _get_db()
    row = db.execute(
        "SELECT id FROM knowledge_documents WHERE kb_id = ? AND file_path = ?",
        (kb_id, path),
    ).fetchone()
    if row:
        return get_knowledge_file_preview_v2(row["id"])
    # Fallback
    target = _resolve_path(kb_id, path, must_exist=True)
    suffix = target.suffix.lower()
    content = ""
    if suffix in (".txt", ".md", ".csv", ".json"):
        content = target.read_text(encoding="utf-8", errors="replace")[:120_000]
    return {"path": path, "content": content, "lines": content.count("\n") + 1, "size": target.stat().st_size}


def start_vectorization(path: str) -> dict:
    """Legacy vectorization endpoint — default KB."""
    kb_id = _default_kb_id()
    db = _get_db()
    row = db.execute(
        "SELECT id FROM knowledge_documents WHERE kb_id = ? AND file_path = ?",
        (kb_id, path),
    ).fetchone()
    if not row:
        # Fallback: sync from disk then vectorize
        _sync_doc_from_disk(kb_id, None, path)
        row = db.execute(
            "SELECT id FROM knowledge_documents WHERE kb_id = ? AND file_path = ?",
            (kb_id, path),
        ).fetchone()
    if not row:
        return {"path": path, "status": "failed", "error": "Document not found"}
    result = start_vectorization_v2(row["id"])
    return {"path": path, "status": result["status"], "job_id": result["job_id"]}


def get_vectorization_status(path: str) -> dict:
    """Legacy status endpoint — default KB."""
    kb_id = _default_kb_id()
    db = _get_db()
    row = db.execute(
        "SELECT id, parse_status, chunk_count, vector_count, error_message FROM knowledge_documents WHERE kb_id = ? AND file_path = ?",
        (kb_id, path),
    ).fetchone()
    if row:
        return {
            "status": row["parse_status"],
            "progress": 100 if row["parse_status"] == "completed" else 0,
            "chunks_done": row["vector_count"],
            "chunks_total": row["chunk_count"],
            "error": row["error_message"] or "",
        }
    return _get_vect_status(path)


def search_knowledge(query: str, top_k: int = 5) -> dict:
    """Legacy search endpoint — default KB."""
    return search_knowledge_v2(_default_kb_id(), query, top_k=top_k)


def list_knowledge_docs(params: dict) -> dict:
    """Legacy paginated list — default KB."""
    return list_knowledge_documents(_default_kb_id(), params)


def batch_knowledge_status(paths: list[str]) -> dict:
    """Legacy batch status — default KB."""
    kb_id = _default_kb_id()
    db = _get_db()
    doc_ids = []
    for path in paths:
        row = db.execute(
            "SELECT id FROM knowledge_documents WHERE kb_id = ? AND file_path = ?",
            (kb_id, path),
        ).fetchone()
        if row:
            doc_ids.append(row["id"])
    if doc_ids:
        return batch_knowledge_status_v2(doc_ids)
    # Fallback to legacy document_state
    return {
        "data": [
            {
                "id": p,
                "parse_status": _get_vect_status(p).get("status", "pending"),
                "progress": _get_vect_status(p).get("progress", 0),
                "chunks_done": _get_vect_status(p).get("chunks_done", 0),
                "chunks_total": _get_vect_status(p).get("chunks_total", 0),
                "error": _get_vect_status(p).get("error", ""),
            }
            for p in paths
        ]
    }


# ═══════════════════════════════════════════════════════════════════════════════
# Helpers for legacy document_state table
# ═══════════════════════════════════════════════════════════════════════════════


def _get_vect_status(rel_path: str) -> dict:
    db = _get_db()
    row = db.execute(
        "SELECT status, progress, chunks_done, chunks_total, error FROM document_state WHERE rel_path = ?",
        (rel_path,),
    ).fetchone()
    if row:
        return {
            "status": row["status"],
            "progress": row["progress"],
            "chunks_done": row["chunks_done"],
            "chunks_total": row["chunks_total"],
            "error": row["error"] or "",
        }
    return {"status": "pending", "progress": 0, "chunks_done": 0, "chunks_total": 0, "error": ""}


def _set_vect_status(rel_path: str, **kwargs):
    db = _get_db()
    existing = db.execute("SELECT 1 FROM document_state WHERE rel_path = ?", (rel_path,)).fetchone()
    if existing:
        sets = ", ".join(f"{k} = ?" for k in kwargs)
        vals = list(kwargs.values()) + [rel_path]
        db.execute(f"UPDATE document_state SET {sets}, updated_at = datetime('now') WHERE rel_path = ?", vals)
    else:
        cols = ", ".join(kwargs.keys())
        placeholders = ", ".join("?" for _ in kwargs)
        db.execute(
            f"INSERT INTO document_state (rel_path, {cols}) VALUES (?, {placeholders})",
            (rel_path, *kwargs.values()),
        )
    db.commit()


def _delete_vect_status(key: str):
    db = _get_db()
    db.execute("DELETE FROM document_state WHERE rel_path = ?", (key,))
    db.commit()


# ═══════════════════════════════════════════════════════════════════════════════
# Disk → DB sync utilities
# ═══════════════════════════════════════════════════════════════════════════════


def _sync_disk_to_db(kb_id: str, delete_orphans: bool = False) -> dict:
    """Scan the KB root and ensure folders/documents are reflected in metadata.

    Returns {"inserted": n, "updated": n, "orphaned": n}. Orphaned rows (DB
    record without a disk file) are only deleted when delete_orphans=True;
    otherwise they are just counted so callers can surface them for explicit
    cleanup.
    """
    stats = {"inserted": 0, "updated": 0, "orphaned": 0}
    root = _resolve_kb_root(kb_id)
    db = _get_db()
    folder_map: dict[str, str] = {}
    disk_files: set[str] = set()

    # Exclude other KB roots when scanning default/shared roots
    other_roots = {
        Path(r["root_path"]).resolve()
        for r in db.execute("SELECT root_path FROM knowledge_bases WHERE id != ?", (kb_id,)).fetchall()
    }

    # Ensure root folder record exists
    root_id = _ensure_folder_record(kb_id, None, "")
    folder_map[""] = root_id

    for dirpath, dirnames, filenames in os.walk(root):
        current = Path(dirpath).resolve()
        rel = _rel_path(kb_id, current)

        # Skip directories that are other KB roots or inside them
        if current in other_roots or any(str(current).startswith(str(oroot) + os.sep) for oroot in other_roots):
            dirnames[:] = []
            continue

        # Filter hidden folders
        dirnames[:] = [d for d in dirnames if not _is_hidden(d)]
        # Also filter out other KB root directories at top level
        dirnames[:] = [d for d in dirnames if (current / d).resolve() not in other_roots]

        parent_rel = str(Path(rel).parent).replace("\\", "/") if rel else ""
        if parent_rel == ".":
            parent_rel = ""
        parent_id = folder_map.get(parent_rel)
        current_id = _ensure_folder_record(kb_id, parent_id, rel)
        folder_map[rel] = current_id

        for filename in filenames:
            if _is_hidden(filename):
                continue
            suffix = Path(filename).suffix.lower()
            if suffix not in _SUPPORTED_EXTS:
                continue
            file_rel = f"{rel}/{filename}".lstrip("/") if rel else filename
            disk_files.add(file_rel)
            abs_path = current / filename
            if not abs_path.exists():
                continue
            existing = db.execute(
                "SELECT id, file_hash, metadata FROM knowledge_documents WHERE kb_id = ? AND file_path = ?",
                (kb_id, file_rel),
            ).fetchone()
            st = abs_path.stat()
            now = _now_iso()
            if existing:
                # Fast path: skip the full-file SHA256 when mtime+size match the
                # values recorded at last sync. Hashing every file on every tree
                # request made panel loads O(total bytes) instead of O(stat).
                try:
                    prev_mtime = json.loads(existing["metadata"] or "{}").get("mtime")
                except (ValueError, TypeError):
                    prev_mtime = None
                if prev_mtime == st.st_mtime:
                    continue
                file_hash = _file_hash(abs_path)
                if existing["file_hash"] != file_hash:
                    db.execute(
                        "UPDATE knowledge_documents SET file_size = ?, file_hash = ?, metadata = ?, updated_at = ? WHERE id = ?",
                        (st.st_size, file_hash, _json_dumps({"mtime": st.st_mtime}), now, existing["id"]),
                    )
                    stats["updated"] += 1
                else:
                    # Content unchanged (same hash) — just record the new mtime
                    # so the fast path hits next time.
                    db.execute(
                        "UPDATE knowledge_documents SET metadata = ?, updated_at = ? WHERE id = ?",
                        (_json_dumps({"mtime": st.st_mtime}), now, existing["id"]),
                    )
            else:
                file_hash = _file_hash(abs_path)
                doc_id = _new_id()
                db.execute(
                    """
                    INSERT INTO knowledge_documents
                    (id, kb_id, folder_id, title, file_name, file_type, file_size, file_path,
                     file_hash, source, metadata, parse_status, summary_status, summary_text,
                     chunk_count, vector_count, error_message, created_at, updated_at, processed_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        doc_id, kb_id, current_id, filename, filename,
                        suffix.lstrip("."), st.st_size, file_rel,
                        file_hash, "upload", _json_dumps({"mtime": st.st_mtime}), "pending", "none", "",
                        0, 0, "", now, now, None,
                    ),
                )
                stats["inserted"] += 1

    # Cleanup metadata for files/folders removed from disk.
    existing_docs = db.execute(
        "SELECT id, file_path FROM knowledge_documents WHERE kb_id = ?", (kb_id,)
    ).fetchall()
    # Documents whose file no longer exists on disk are ORPHANS. They are
    # never auto-deleted here — the DB is the single source of truth and
    # orphan cleanup is an explicit, user-confirmed action (see
    # cleanup_orphaned_documents). We only count/report them.
    for row in existing_docs:
        if row["file_path"] not in disk_files:
            stats["orphaned"] += 1
            if delete_orphans:
                try:
                    delete_knowledge_document(row["id"])
                except Exception:
                    logger.exception("Failed to clean up removed document %s", row["id"])

    folder_paths = set(folder_map.keys())
    if folder_paths:
        db.execute(
            "DELETE FROM knowledge_folders WHERE kb_id = ? AND path != '' AND path NOT IN (%s)"
            % ",".join("?" * len(folder_paths)),
            (kb_id, *folder_paths),
        )

    db.commit()
    return stats


def reconcile_knowledge_base(kb_id: str) -> dict:
    """Explicit disk→DB reconcile: sync, repair folder_id linkage, report.

    The DB is the single source of truth for tree/list/stats; this endpoint
    brings it in line with disk and fixes historical rows whose folder_id is
    NULL or stale (uploads predating persistent folder records). Orphans are
    reported, never deleted — use cleanup_orphaned_documents for that.
    """
    sync_stats = _sync_disk_to_db(kb_id)
    repaired = _repair_folder_links(kb_id)
    return {
        "kb_id": kb_id,
        "inserted": sync_stats["inserted"],
        "updated": sync_stats["updated"],
        "orphaned": sync_stats["orphaned"],
        "folder_links_repaired": repaired,
    }


def _repair_folder_links(kb_id: str) -> int:
    """Recompute every document's folder_id from its file_path by longest
    prefix match against knowledge_folders. Returns number of rows changed."""
    db = _get_db()
    folders = db.execute(
        "SELECT id, path FROM knowledge_folders WHERE kb_id = ?", (kb_id,)
    ).fetchall()
    # Longest path first so the most specific folder wins.
    folder_list = sorted(folders, key=lambda f: len(f["path"]), reverse=True)
    root_id = next((f["id"] for f in folders if f["path"] == ""), None)

    docs = db.execute(
        "SELECT id, file_path, folder_id FROM knowledge_documents WHERE kb_id = ?", (kb_id,)
    ).fetchall()
    repaired = 0
    for doc in docs:
        fp = doc["file_path"] or ""
        parent = fp.rsplit("/", 1)[0] if "/" in fp else ""
        target_id = root_id
        for f in folder_list:
            if f["path"] and f["path"] == parent:
                target_id = f["id"]
                break
        if target_id != doc["folder_id"]:
            db.execute(
                "UPDATE knowledge_documents SET folder_id = ?, updated_at = ? WHERE id = ?",
                (target_id, _now_iso(), doc["id"]),
            )
            repaired += 1
    if repaired:
        db.commit()
    return repaired


def list_orphaned_documents(kb_id: str) -> list[dict]:
    """DB rows whose backing file no longer exists on disk."""
    root = _resolve_kb_root(kb_id)
    db = _get_db()
    rows = db.execute(
        "SELECT id, file_name, file_path, parse_status FROM knowledge_documents WHERE kb_id = ?",
        (kb_id,),
    ).fetchall()
    return [
        {"id": r["id"], "file_name": r["file_name"], "file_path": r["file_path"], "parse_status": r["parse_status"]}
        for r in rows
        if not (root / (r["file_path"] or "")).exists()
    ]


def cleanup_orphaned_documents(kb_id: str) -> dict:
    """One-click cleanup: delete DB rows (and their vectors) whose file is
    gone from disk. Returns the deleted ids for UI confirmation."""
    orphans = list_orphaned_documents(kb_id)
    deleted, failed = [], 0
    for doc in orphans:
        try:
            delete_knowledge_document(doc["id"])
            deleted.append(doc["id"])
        except Exception:
            logger.exception("Orphan cleanup failed for doc %s", doc["id"])
            failed += 1
    return {"kb_id": kb_id, "deleted": len(deleted), "failed": failed, "doc_ids": deleted}


def _ensure_folder_record(kb_id: str, parent_id: str | None, rel_path: str) -> str:
    db = _get_db()
    row = db.execute(
        "SELECT id FROM knowledge_folders WHERE kb_id = ? AND path = ?", (kb_id, rel_path)
    ).fetchone()
    if row:
        return row["id"]

    folder_id = _new_id()
    name = Path(rel_path).name if rel_path else ""
    depth = rel_path.count("/") if rel_path else 0
    now = _now_iso()
    db.execute(
        """
        INSERT INTO knowledge_folders (id, kb_id, parent_id, name, path, depth, sort_order, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (folder_id, kb_id, parent_id, name, rel_path, depth, 0, now),
    )
    return folder_id


# Migration alias: keep old _rel_path signature for any external callers
# (new code should use _rel_path(kb_id, target))
def _legacy_rel_path(target: Path) -> str:
    return _rel_path(_default_kb_id(), target)
